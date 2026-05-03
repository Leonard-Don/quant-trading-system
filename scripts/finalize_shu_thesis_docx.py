from __future__ import annotations

import json
import re
import subprocess
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from docxcompose.composer import Composer
from pypdf import PdfReader, PdfWriter


def resolve_existing_path(description: str, *candidates: Path) -> Path:
    for candidate in candidates:
        if candidate.exists():
            return candidate
    searched = "\n".join(f"- {candidate}" for candidate in candidates)
    raise FileNotFoundError(f"Failed to locate {description}. Checked:\n{searched}")


def find_first_existing_path(*candidates: Path) -> Path | None:
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOC_OUTPUT_DIR = PROJECT_ROOT / "output" / "doc"
TMP_DOC_DIR = PROJECT_ROOT / "tmp" / "docs"
SCREENSHOT_DIR = PROJECT_ROOT / "docs" / "screenshots"
DOWNLOAD_DIR = Path.home() / "Downloads"
HEATMAP_HISTORY_PATH = PROJECT_ROOT / "data" / "industry" / "heatmap_history.json"
THESIS_HEATMAP_SNAPSHOT_PATH = PROJECT_ROOT / "docs" / "thesis_assets" / "frozen_heatmap_snapshot.json"
THESIS_LEADER_CASE_PATH = PROJECT_ROOT / "docs" / "thesis_assets" / "frozen_leader_case.json"

DOC_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究.docx"
TEMPLATE_PATH = resolve_existing_path(
    "official template DOCX",
    DOWNLOAD_DIR / "论文模版" / "上海大学本科毕业论文（设计）撰写格式模板.docx",
    DOWNLOAD_DIR / "上海大学本科毕业论文（设计）撰写格式模板.docx",
)
TEMPLATE_PDF_PATH = resolve_existing_path(
    "official template PDF",
    DOWNLOAD_DIR / "论文模版" / "上海大学本科毕业论文（设计）撰写格式模板.pdf",
    DOWNLOAD_DIR / "上海大学本科毕业论文（设计）撰写格式模板.pdf",
)
OUTPUT_PDF_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究.pdf"
LEGACY_DUPLICATE_PDF_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究_送审版.pdf"
COMPOSED_TMP_PATH = TMP_DOC_DIR / "shu_thesis_composed.docx"
TMP_TEMPLATE_PDF_DIR = TMP_DOC_DIR / "template_pdf_pages"
TMP_DOCX_RENDER_DIR = TMP_DOC_DIR / "current_docx_render"
TMP_SUBMISSION_RENDER_DIR = TMP_DOC_DIR / "submission_pdf_pages"
TMP_FRONT_ASSET_DIR = TMP_DOC_DIR / "pdf_front_assets"
TEMPLATE_XML_PREFIX = TMP_TEMPLATE_PDF_DIR / "template_layout"

THESIS_TITLE = "基于大数据的热门行业识别与龙头股遴选研究"
STUDENT_INFO = {
    "college": "通信与信息工程学院",
    "major": "通信工程",
    "student_id": "22121527",
    "student_name": "唐梓涵",
    "advisor": "沈文辉",
    "duration": "2025.12-2026.6",
}

FIGURE_IMAGES = {
    "图 5.1 行业热度总览界面": {
        "path": SCREENSHOT_DIR / "industry-ranking-overview.png",
        "crop": (250, 140, 1360, 1160),
        "slug": "figure_5_1",
        "width_scale": 0.72,
    },
    "图 5.2 行业热力图界面": {
        "path": SCREENSHOT_DIR / "industry-heatmap-overview.png",
        "crop": (240, 120, 1360, 1180),
        "slug": "figure_5_2",
    },
    "图 5.3 龙头股详情界面": {
        "path": SCREENSHOT_DIR / "leader-stock-detail.png",
        "crop": (300, 90, 1160, 1080),
        "slug": "figure_5_3",
    },
}
FIGURE_INTRO_PARAGRAPHS = {
    "图 5.1 行业热度总览界面":
        "如图 5.1 所示，行业热度总览页面将热力图、行业排行榜、趋势面板与龙头股入口整合在同一研究界面中，便于研究者快速完成行业筛选、排序比较与详情联动。",
    "图 5.2 行业热力图界面":
        "如图 5.2 所示，热力图视图支持按颜色维度、尺寸维度和统计周期切换行业强弱，使热门行业的横截面分布特征能够被直观比较。",
    "图 5.3 龙头股详情界面":
        "如图 5.3 所示，龙头股详情页面展示综合得分、维度拆解、价格走势与关键财务字段，能够帮助用户理解候选股票被遴选出的具体依据。",
}
TMP_ASSET_DIR = TMP_DOC_DIR / "figure_assets"
ARCHITECTURE_FIGURE_PATH = TMP_ASSET_DIR / "figure_3_1_system_architecture.png"
CJK_FONT_SOURCES = (
    (Path("/Library/Fonts/SimSun.ttf"), 0),
    (Path("/System/Library/Fonts/SimSun.ttf"), 0),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 6),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 1),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 4),
)
BODY_PDF_ABSTRACT_PAGE_INDEX = 4
BODY_FIRST_LINE_INDENT = Emu(266700)
BODY_LINE_SPACING = Pt(23)
ABSTRACT_LINE_SPACING = Pt(20)
SOURCE_NOTE = "图片来源：系统运行截图（作者自制）"
ARCHITECTURE_SOURCE_NOTE = "图片来源：作者根据系统实现绘制"
HEADER_TEXT = "上海大学本科毕业论文（设计）"
SIGNED_DECLARATION_PAGE_CANDIDATES = (
    DOC_OUTPUT_DIR / "上海大学本科毕业论文_签字页扫描.pdf",
    DOC_OUTPUT_DIR / "上海大学本科毕业论文_签字页扫描.png",
    DOC_OUTPUT_DIR / "上海大学本科毕业论文_签字页扫描.jpg",
    DOC_OUTPUT_DIR / "上海大学本科毕业论文_签字页扫描.jpeg",
    DOC_OUTPUT_DIR / "签字页扫描.pdf",
    DOC_OUTPUT_DIR / "签字页扫描.png",
    DOC_OUTPUT_DIR / "签字页扫描.jpg",
    DOC_OUTPUT_DIR / "签字页扫描.jpeg",
)
LEADER_CASE_ANALYSIS_FALLBACK_TEXT = (
    "为补充说明龙头股遴选结果，本文固定引用项目中保存的一条龙头股双榜单样本。"
    "从该样本可以看到，系统不会把行业热度排序简单重复为单一股票列表，而是同时输出"
    "core（核心资产）与 hot（热点先锋）两类结果：前者更强调规模、估值与盈利质量，"
    "后者更强调短期涨幅和资金承接。该设计使行业识别结果能够进一步延伸为行业内个股的"
    "结构化筛选结果。"
)

COVER_FIELD_ORDER = [
    "title",
    "college",
    "major",
    "student_id",
    "student_name",
    "advisor",
    "duration",
]

COVER_FIELD_SPECS = {
    "title": {
        "text": THESIS_TITLE,
        # SHU cover spec: title in small No.2, rendered here on the
        # 150dpi template image as an approximately equivalent pixel size.
        "font_size": 38,
        "max_width": 520,
        "align": "center",
        "padding_x": 0,
        "gap_above_line": 4,
    },
    "college": {
        "text": STUDENT_INFO["college"],
        "font_size": 33,
        "max_width": 438,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "major": {
        "text": STUDENT_INFO["major"],
        "font_size": 33,
        "max_width": 438,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "student_id": {
        "text": STUDENT_INFO["student_id"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "student_name": {
        "text": STUDENT_INFO["student_name"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "advisor": {
        "text": STUDENT_INFO["advisor"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "duration": {
        "text": STUDENT_INFO["duration"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
}

COVER_FALLBACK_LINES = {
    "title": {"x_start": 458, "x_end": 964, "y_line": 848},
    "college": {"x_start": 444, "x_end": 949, "y_line": 976},
    "major": {"x_start": 444, "x_end": 949, "y_line": 1041},
    "student_id": {"x_start": 444, "x_end": 949, "y_line": 1106},
    "student_name": {"x_start": 444, "x_end": 949, "y_line": 1171},
    "advisor": {"x_start": 444, "x_end": 949, "y_line": 1236},
    "duration": {"x_start": 444, "x_end": 949, "y_line": 1301},
}

DECLARATION_FIELD_SPECS = {
    "student_name": {
        "text": STUDENT_INFO["student_name"],
        "anchor_label": "姓名：",
        "font_size": 20,
        "max_width": 160,
        "padding_x": 18,
        "vertical_align": "center",
    },
    "student_id": {
        "text": STUDENT_INFO["student_id"],
        "anchor_label": "学号：",
        "font_size": 19,
        "max_width": 170,
        "padding_x": 16,
        "vertical_align": "center",
    },
    "thesis_title": {
        "text": THESIS_TITLE,
        "anchor_label": "论文题目：",
        "font_size": 20,
        "max_width": 560,
        "padding_x": 22,
        "vertical_align": "center",
    },
}

DECLARATION_FALLBACK_LABEL_BOXES = {
    "姓名：": {"left": 214, "top": 160, "right": 339, "bottom": 185},
    "学号：": {"left": 747, "top": 160, "right": 822, "bottom": 185},
    "论文题目：": {"left": 214, "top": 208, "right": 339, "bottom": 233},
}


TOC_BLUEPRINT = [
    ("toc 1", "摘  要"),
    ("toc 1", "ABSTRACT"),
    ("toc 1", "1 绪论"),
    ("toc 2", "1.1 研究背景与意义"),
    ("toc 2", "1.2 国内外研究现状"),
    ("toc 2", "1.3 研究内容与技术路线"),
    ("toc 2", "1.4 论文结构安排"),
    ("toc 1", "2 相关技术与理论基础"),
    ("toc 2", "2.1 金融大数据的特征"),
    ("toc 2", "2.2 多源数据采集与清洗"),
    ("toc 2", "2.3 热门行业识别的理论基础"),
    ("toc 2", "2.4 龙头股评价的理论基础"),
    ("toc 2", "2.5 系统关键技术"),
    ("toc 1", "3 系统需求分析与总体设计"),
    ("toc 2", "3.1 设计目标"),
    ("toc 2", "3.2 功能需求分析"),
    ("toc 2", "3.3 非功能需求分析"),
    ("toc 2", "3.4 系统总体架构设计"),
    ("toc 2", "3.5 数据流程设计"),
    ("toc 2", "3.6 存储设计"),
    ("toc 1", "4 热门行业识别与龙头股遴选模型设计"),
    ("toc 2", "4.1 数据预处理与指标标准化"),
    ("toc 2", "4.2 热门行业识别模型设计"),
    ("toc 2", "4.3 行业波动率估计与聚类辅助分析"),
    ("toc 2", "4.4 龙头股综合评分模型设计"),
    ("toc 2", "4.5 快速评分机制设计"),
    ("toc 1", "5 系统实现"),
    ("toc 2", "5.1 数据提供器与回退机制实现"),
    ("toc 2", "5.2 行业分析模块实现"),
    ("toc 2", "5.3 龙头股评分模块实现"),
    ("toc 2", "5.4 后端接口实现"),
    ("toc 2", "5.5 前端可视化实现"),
    ("toc 2", "5.6 偏好配置与持续跟踪实现"),
    ("toc 1", "6 系统测试与结果分析"),
    ("toc 2", "6.1 测试环境"),
    ("toc 2", "6.2 功能测试"),
    ("toc 2", "6.3 热力图快照结果分析"),
    ("toc 2", "6.4 系统特征与不足分析"),
    ("toc 2", "6.5 对毕业设计目标的达成情况"),
    ("toc 1", "结 论"),
    ("toc 1", "参考文献"),
    ("toc 1", "致 谢"),
]


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_generated_padding_paragraphs(doc: Document) -> None:
    """Drop synthetic filler paragraphs left by a previous emergency expansion."""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if (
            text.startswith("围绕")
            and "中的“" in text
            and (
                "本文更关注它在真实项目链路" in text
                or "始终把数据来源、计算口径、缓存策略和页面展示" in text
            )
        ):
            delete_paragraph(paragraph)


def remove_section_break(paragraph) -> None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is not None:
        p_pr.remove(sect_pr)


def set_east_asia_font(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def replace_paragraph_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def replace_cover_line(paragraph, prefix: str, value: str, gap_spaces: int = 8) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(f"{prefix}{' ' * gap_spaces}{value}")


def ensure_page_break_before(paragraph) -> None:
    previous = paragraph._element.getprevious()
    if previous is not None and previous.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        previous_para = Paragraph(previous, paragraph._parent)
        if previous_para.text.strip() == "" and 'w:type="page"' in previous_para._element.xml:
            return
    break_para = paragraph.insert_paragraph_before()
    break_para.add_run().add_break(WD_BREAK.PAGE)


def remove_page_break_before_flag(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for node in p_pr.findall(qn("w:pageBreakBefore")):
        p_pr.remove(node)


def format_unnumbered_heading(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        run.font.bold = True
        run.font.size = Pt(18)
        run.font.name = "Times New Roman"
        set_east_asia_font(run, "黑体")


def format_toc_title(paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.page_break_before = True
    run = paragraph.add_run("目  录")
    run.font.bold = False
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "黑体")


def add_page_number_field(paragraph) -> None:
    clear_paragraph(paragraph)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def set_section_page_number_format(section, fmt: str | None = None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def configure_section_header(section, text: str | None) -> None:
    section.different_first_page_header_footer = False
    header_para = section.header.paragraphs[0]
    clear_paragraph(header_para)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.first_line_indent = Pt(0)
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    if text:
        run = header_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.bold = False
        set_east_asia_font(run, "宋体")


def configure_section_footer(section, with_page_number: bool) -> None:
    footer_para = section.footer.paragraphs[0]
    clear_paragraph(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.first_line_indent = Pt(0)
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)
    if with_page_number:
        add_page_number_field(footer_para)


def set_run_text_font(run, east_asia_font: str, size: Pt, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    set_east_asia_font(run, east_asia_font)


def format_body_run(run) -> None:
    if not run.text:
        return
    if run.font.size is None:
        run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "宋体")


def normalize_search_text(text: str) -> str:
    normalized = re.sub(r"\s+", "", text or "")
    return normalized.replace("（", "(").replace("）", ")")


def insert_paragraph_after(paragraph, text: str = ""):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before_element(element, parent):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    return Paragraph(new_p, parent)


def parse_iso_datetime(value: str | None) -> datetime | None:
    raw_value = str(value or "").strip()
    if not raw_value:
        return None
    try:
        return datetime.fromisoformat(raw_value)
    except ValueError:
        return None


def format_chinese_timestamp(value: str | None) -> str:
    parsed = parse_iso_datetime(value)
    if parsed is None:
        return str(value or "").strip()
    return f"{parsed.year} 年 {parsed.month} 月 {parsed.day} 日 {parsed.hour:02d}:{parsed.minute:02d}:{parsed.second:02d}"


def coerce_float(value, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _load_heatmap_snapshot_from_path(path: Path, days: int = 5) -> dict | None:
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(payload, dict):
        return None
    if int(payload.get("days", 0) or 0) != days:
        return None
    industries = payload.get("industries")
    if not isinstance(industries, list) or not industries:
        return None
    return payload


def load_latest_heatmap_snapshot(days: int = 5) -> dict | None:
    if not HEATMAP_HISTORY_PATH.exists():
        return None
    try:
        payload = json.loads(HEATMAP_HISTORY_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(payload, list):
        return None

    candidates = []
    for item in payload:
        if not isinstance(item, dict):
            continue
        if int(item.get("days", 0) or 0) != days:
            continue
        industries = item.get("industries")
        if not isinstance(industries, list) or not industries:
            continue
        candidates.append(item)
    if not candidates:
        return None

    return max(
        candidates,
        key=lambda item: (
            parse_iso_datetime(item.get("captured_at"))
            or parse_iso_datetime(item.get("update_time"))
            or datetime.min
        ),
    )


def load_thesis_heatmap_snapshot(days: int = 5) -> dict | None:
    frozen_snapshot = _load_heatmap_snapshot_from_path(THESIS_HEATMAP_SNAPSHOT_PATH, days=days)
    if frozen_snapshot is not None:
        frozen_payload = dict(frozen_snapshot)
        frozen_payload["_thesis_frozen"] = True
        return frozen_payload
    return load_latest_heatmap_snapshot(days=days)


def load_thesis_leader_case() -> dict | None:
    if not THESIS_LEADER_CASE_PATH.exists():
        return None
    try:
        payload = json.loads(THESIS_LEADER_CASE_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(payload, dict):
        return None
    focus_industry = str(payload.get("focus_industry") or "").strip()
    core = payload.get("core")
    hot = payload.get("hot")
    if not focus_industry or not isinstance(core, list) or not isinstance(hot, list):
        return None
    if not core or not hot:
        return None
    return payload


def build_leader_case_analysis_text(payload: dict | None) -> str:
    if payload is None:
        return LEADER_CASE_ANALYSIS_FALLBACK_TEXT

    focus_industry = str(payload.get("focus_industry") or "").strip()
    captured_display = format_chinese_timestamp(payload.get("captured_at"))
    related_heatmap_display = format_chinese_timestamp(payload.get("related_heatmap_captured_at"))
    core_names = "、".join(
        str(item.get("name") or "").strip()
        for item in payload.get("core", [])[:3]
        if str(item.get("name") or "").strip()
    )
    hot_names = "、".join(
        str(item.get("name") or "").strip()
        for item in payload.get("hot", [])[:3]
        if str(item.get("name") or "").strip()
    )
    if not focus_industry or not core_names or not hot_names:
        return LEADER_CASE_ANALYSIS_FALLBACK_TEXT

    if related_heatmap_display:
        prefix = (
            "为保证龙头股案例与热力图样例一样可复核，本文固定引用项目中保存的一条龙头股双榜单样本。"
            f"结合 {related_heatmap_display} 保留的热力图样例可以看到，{focus_industry}行业处于当期热点前列；"
        )
    else:
        prefix = (
            "为保证龙头股案例可复核，本文固定引用项目中保存的一条龙头股双榜单样本。"
        )

    if captured_display:
        prefix += f"根据 {captured_display} 记录的榜单结果，"

    return (
        prefix
        + f"{focus_industry}行业的 core 榜单前列主要为{core_names}，"
        f"hot 榜单前列则出现{hot_names}；前者更偏向规模、估值和盈利质量的综合筛选，"
        "后者更强调短期涨幅与资金承接。该结果说明龙头股遴选并非对行业热度排序的简单复制，"
        "而是在已识别热门行业基础上进一步提供“核心资产”和“热点先锋”两类互补视角。"
    )


def build_heatmap_result_section_text(snapshot: dict | None) -> tuple[str, str, str]:
    fallback_intro = (
        "为了增强论文结果分析的真实性，本文直接选取项目中已经保存的热力图历史快照作为样本。"
        "表 6.2 展示了项目当前保留的一条五日窗口快照中的前十个行业横截面字段，包括综合得分、5 日涨跌"
        "幅、资金流和换手率等信息，用于说明系统在固定统计窗口下的输出结果。需要说明的是，这里给出的结果"
        "主要用于样例分析，并不构成完整回测结论。"
    )
    fallback_analysis = (
        "从表 6.2 可以看出，样本时点前列行业并不完全由单一风格板块构成，而是呈现出多方向扩散特征。"
        "这说明系统并非只根据某一项价格指标排序，而是在统一横截面上综合考虑动量、资金流、活跃度和波动率"
        "代理后给出相对得分。"
    )
    fallback_conclusion = (
        "从快照字段本身也能看出，系统输出的不是单一涨跌幅排行。表 6.2 同时保留了 value、total_score、"
        "moneyFlow、turnoverRate、行业波动率和市值来源等信息，所以后文讨论某个行业为什么排在前面时，"
        "能够直接回到对应字段去解释。对毕业设计答辩来说，这样的结果更便于把评分逻辑讲清楚。"
    )
    if snapshot is None:
        return fallback_intro, fallback_analysis, fallback_conclusion

    industries = [
        item for item in snapshot.get("industries", [])
        if isinstance(item, dict) and str(item.get("name", "")).strip()
    ][:10]
    if not industries:
        return fallback_intro, fallback_analysis, fallback_conclusion

    captured_display = format_chinese_timestamp(snapshot.get("captured_at") or snapshot.get("update_time"))
    update_display = format_chinese_timestamp(snapshot.get("update_time") or snapshot.get("captured_at"))
    lead_names = "、".join(item["name"] for item in industries[:6])
    top_item = industries[0]
    contrast_item = next(
        (
            item for item in industries[1:]
            if coerce_float(item.get("moneyFlow")) < 0 and coerce_float(item.get("value")) > 0
        ),
        None,
    )

    is_frozen_snapshot = bool(snapshot.get("_thesis_frozen"))
    intro_prefix = (
        "为保证论文样例在多次生成过程中的一致性，本文固定引用项目中已经保存的一条五日窗口热力图快照作为样本。"
        if is_frozen_snapshot
        else "为了增强论文结果分析的真实性，本文直接选取项目当前仍保留的历史快照作为样本。"
    )

    if captured_display and captured_display == update_display:
        intro_text = (
            intro_prefix
            + f"根据 {captured_display} 记录的一条五日窗口热力图快照，表 6.2 展示了该样本时点前十个行业的综"
            "合得分、5 日涨跌幅、资金流和换手率等字段，用于说明系统在固定横截面上的输出结果。需要说明的是，"
            "这里给出的结果主要用于样例分析，并不构成完整回测结论。"
        )
    else:
        intro_text = (
            intro_prefix
            + f"根据 {captured_display} 记录的一条五日窗口热力图快照（对应快照更新时间为 {update_display}），"
            "表 6.2 展示了该样本时点前十个行业的综合得分、5 日涨跌幅、资金流和换手率等字段，用于说明系统"
            "在固定横截面上的输出结果。需要说明的是，这里给出的结果主要用于样例分析，并不构成完整回测结论。"
        )
    if contrast_item is None:
        analysis_text = (
            f"从表 6.2 可以看出，{lead_names}等行业位于前列，说明样本时点的市场热点呈现出多方向扩散特征，"
            "而非只集中在单一子行业。由此可见，系统给出的综合得分并不是对单一涨跌幅的简单重排，而是对多项横"
            "截面指标进行统一标准化之后的综合结果。"
        )
    else:
        contrast_flow_yi = coerce_float(contrast_item.get("moneyFlow")) / 1e8
        contrast_change = coerce_float(contrast_item.get("value"))
        analysis_text = (
            f"从表 6.2 可以看出，{lead_names}等行业位于前列，说明样本时点的市场热点呈现出多方向扩散特征，"
            f"而非只集中在单一子行业。值得注意的是，{contrast_item['name']}在 5 日涨跌幅达到 "
            f"{contrast_change:.2f}% 的同时资金流仍为 {contrast_flow_yi:.2f} 亿元，而{top_item['name']}"
            "的综合得分更高。这说明系统生成的 total_score 并不简单等同于单一涨跌幅排序，而是横截面标准"
            "化后多个因子共同作用的结果。"
        )
    conclusion_text = (
        "从这条快照保留的字段就能看出，系统输出的不是单一涨跌幅排行。表 6.2 同时记录了 value、"
        "total_score、moneyFlow、turnoverRate、行业波动率和市值来源等信息，所以后文讨论某个行业"
        "为什么排在前面时，能够直接回到对应字段去解释。对毕业设计答辩来说，这样的结果有一个现实好处："
        "老师看到的不只是“哪个行业涨得多”，而是还能继续追问资金、活跃度和波动代理分别起了什么作用。"
    )
    return intro_text, analysis_text, conclusion_text


def update_heatmap_sample_table(sample_table, snapshot: dict | None) -> None:
    if sample_table is None or len(sample_table.rows) < 2 or len(sample_table.columns) < 6:
        return

    sample_table.cell(0, 0).text = "序号"
    sample_table.cell(0, 4).text = "资金流(亿元)"

    industries = []
    if isinstance(snapshot, dict):
        industries = [
            item for item in snapshot.get("industries", [])
            if isinstance(item, dict)
        ][:10]

    for row_index in range(1, min(len(sample_table.rows), 11)):
        if row_index - 1 < len(industries):
            item = industries[row_index - 1]
            sample_table.cell(row_index, 0).text = str(row_index)
            sample_table.cell(row_index, 1).text = str(item.get("name", ""))
            sample_table.cell(row_index, 2).text = f"{coerce_float(item.get('total_score')):.2f}"
            sample_table.cell(row_index, 3).text = f"{coerce_float(item.get('value')):.2f}"
            sample_table.cell(row_index, 4).text = f"{coerce_float(item.get('moneyFlow')) / 1e8:.2f}"
            sample_table.cell(row_index, 5).text = f"{coerce_float(item.get('turnoverRate')):.2f}"
        else:
            for col_index in range(6):
                sample_table.cell(row_index, col_index).text = ""


def ensure_section_trailing_paragraph(doc: Document, next_heading_text: str, prefix: str, text: str):
    from docx.text.paragraph import Paragraph

    heading = find_paragraph_by_text(doc, next_heading_text)
    previous_element = heading._element.getprevious()
    while previous_element is not None:
        current_element = previous_element
        previous_element = current_element.getprevious()
        if current_element.tag != qn("w:p"):
            continue
        paragraph = Paragraph(current_element, heading._parent)
        stripped = paragraph.text.strip()
        if not stripped:
            continue
        if stripped.startswith(prefix) or stripped == text:
            current_element.getparent().remove(current_element)

    inserted = insert_paragraph_before_element(heading._element, heading._parent)
    inserted.add_run(text)
    return inserted


def insert_table_before_element(doc: Document, element, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    element.addprevious(tbl)
    return table


def fill_cover(doc: Document) -> None:
    replacements = {
        "题   目：": (THESIS_TITLE, 8),
        "学    院：": (STUDENT_INFO["college"], 8),
        "专    业：": (STUDENT_INFO["major"], 8),
        "学    号：": (STUDENT_INFO["student_id"], 8),
        "学生姓名：": (STUDENT_INFO["student_name"], 6),
        "指导教师：": (STUDENT_INFO["advisor"], 6),
        "起讫日期：": (STUDENT_INFO["duration"], 8),
    }
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        for prefix, (value, gap_spaces) in replacements.items():
            if paragraph.text.startswith(prefix):
                replace_cover_line(paragraph, prefix, value, gap_spaces=gap_spaces)


def fill_declaration_table(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    table.cell(0, 0).text = f"姓    名：{STUDENT_INFO['student_name']}"
    table.cell(0, 1).text = f"学号：{STUDENT_INFO['student_id']}"
    merged = table.cell(1, 0).merge(table.cell(1, 1))
    merged.text = f"论文题目：{THESIS_TITLE}"


def remove_elements_before_paragraph(doc: Document, text: str) -> None:
    body = doc._element.body
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            anchor = paragraph._element
            break
    if anchor is None:
        raise RuntimeError(f"Paragraph not found for strip-before: {text}")

    for child in list(body):
        if child is anchor:
            break
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def remove_elements_from_paragraph(doc: Document, text: str) -> None:
    body = doc._element.body
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            anchor = paragraph._element
            break
    if anchor is None:
        raise RuntimeError(f"Paragraph not found for strip-after: {text}")

    removing = False
    for child in list(body):
        if child is anchor:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def find_heading_break_paragraph(doc: Document, heading_text: str):
    heading_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading_text:
            heading_index = index
            break
    if heading_index is None:
        raise RuntimeError(f"Heading not found: {heading_text}")

    for index in range(heading_index - 1, max(-1, heading_index - 5), -1):
        if "w:sectPr" in doc.paragraphs[index]._element.xml:
            return doc.paragraphs[index]
    raise RuntimeError(f"Section break paragraph not found before heading: {heading_text}")


def replace_first_paragraph_starting_with(doc: Document, prefix: str, replacement: str) -> bool:
    prefix_hint = prefix[:24]
    replacement_prefix = replacement[:24]
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped.startswith(prefix) or stripped.startswith(prefix_hint) or stripped.startswith(replacement_prefix):
            replace_paragraph_text(paragraph, replacement)
            return True
    return False


def find_last_paragraph_index_by_text(doc: Document, text: str) -> int:
    matches = [idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == text]
    if not matches:
        raise RuntimeError(f"Paragraph not found: {text}")
    return matches[-1]


def replace_section_body(doc: Document, heading_text: str, next_heading_text: str, paragraphs: list[str]) -> None:
    start_index = find_last_paragraph_index_by_text(doc, heading_text)
    end_index = next(
        (
            idx
            for idx in range(start_index + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip() == next_heading_text
        ),
        None,
    )
    if end_index is None:
        raise RuntimeError(f"Section end heading not found after {heading_text}: {next_heading_text}")

    heading_para = doc.paragraphs[start_index]
    end_element = doc.paragraphs[end_index]._element
    current_element = heading_para._element.getnext()
    while current_element is not None and current_element is not end_element:
        next_element = current_element.getnext()
        current_element.getparent().remove(current_element)
        current_element = next_element

    current = heading_para
    for text in paragraphs:
        current = insert_paragraph_after(current, text)
        if re.match(r"^\d+\.\d+\.\d+\s+", text):
            current.style = "Heading 3"


def append_section_body(doc: Document, heading_text: str, next_heading_text: str, paragraphs: list[str]) -> None:
    start_index = find_last_paragraph_index_by_text(doc, heading_text)
    end_index = next(
        (
            idx
            for idx in range(start_index + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip() == next_heading_text
        ),
        None,
    )
    if end_index is None:
        raise RuntimeError(f"Section end heading not found after {heading_text}: {next_heading_text}")

    paragraph_set = {text.strip() for text in paragraphs}
    for index in range(end_index - 1, start_index, -1):
        paragraph = doc.paragraphs[index]
        if paragraph.text.strip() in paragraph_set:
            delete_paragraph(paragraph)

    start_index = find_last_paragraph_index_by_text(doc, heading_text)
    end_index = next(
        (
            idx
            for idx in range(start_index + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip() == next_heading_text
        ),
        None,
    )
    if end_index is None:
        raise RuntimeError(f"Section end heading not found after cleanup: {heading_text}: {next_heading_text}")

    current = doc.paragraphs[end_index - 1]
    for text in paragraphs:
        current = insert_paragraph_after(current, text)
        if re.match(r"^\d+\.\d+\.\d+\s+", text):
            current.style = "Heading 3"


def insert_table_after(doc: Document, paragraph, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    paragraph._element.addnext(tbl)
    return table


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str | None = None, size: str = "16") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        if color:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)
        else:
            node.set(qn("w:val"), "nil")


def format_architecture_cell(cell, title: str, body_lines: list[str], fill: str, border_color: str) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, fill)
    set_cell_borders(cell, border_color, size="18")
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(17)
    title_run = paragraph.add_run(title)
    set_run_text_font(title_run, "宋体", Pt(16), bold=True)
    title_run.add_break(WD_BREAK.LINE)
    body_run = paragraph.add_run("\n".join(body_lines))
    set_run_text_font(body_run, "宋体", Pt(12))


def format_architecture_arrow_cell(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, "F8FAFD")
    set_cell_borders(cell, None)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("↓")
    set_run_text_font(run, "宋体", Pt(16), bold=True)


def format_architecture_summary_cell(cell, lines: list[str]) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, "FFFFFF")
    set_cell_borders(cell, "B0B8C2", size="12")
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(16)
    run = paragraph.add_run("\n".join(lines))
    set_run_text_font(run, "宋体", Pt(11))


def build_architecture_diagram(output_path: Path) -> Path:
    TMP_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1280, 980
    image = Image.new("RGB", (width, height), "#F8FAFD")
    draw = ImageDraw.Draw(image)
    font_sources = get_available_cjk_font_sources()

    title_font = fit_cjk_font(draw, "系统总体架构图", 520, 34, min_size=28, font_sources=font_sources)
    layer_font = fit_cjk_font(draw, "表现层", 260, 28, min_size=22, font_sources=font_sources)
    text_font = fit_cjk_font(draw, "行业热力图 / 行业排行榜", 920, 22, min_size=18, font_sources=font_sources)
    summary_font = fit_cjk_font(draw, "多源适配、分层缓存与前后端协同", 900, 20, min_size=16, font_sources=font_sources)

    draw.text((width / 2, 44), "系统总体架构图", fill="#243447", font=title_font, anchor="ma")

    layer_specs = [
        ("表现层", "行业热力图 / 行业排行榜\n行业趋势面板 / 龙头股详情", "#D9EAFB", "#5B8FD9"),
        ("服务层", "FastAPI 行业接口 / 参数校验\n响应封装 / 端点缓存", "#E3F4E8", "#5BA36B"),
        ("分析层", "行业分析器 / 龙头股评分器\n快速评分链路 / 波动率估计", "#FDEBCF", "#C58B1C"),
        ("数据层", "THS 主数据 / AKShare 增强\nSina 与腾讯补充 / JSON 与 localStorage", "#EFE2FB", "#8B62C4"),
    ]

    box_left, box_right = 120, 1160
    box_height = 145
    gap = 28
    start_top = 88

    for index, (layer_name, layer_text, fill_color, line_color) in enumerate(layer_specs):
        top = start_top + index * (box_height + gap)
        bottom = top + box_height
        draw.rounded_rectangle((box_left, top, box_right, bottom), radius=28, fill=fill_color, outline=line_color, width=5)
        draw.text((width / 2, top + 42), layer_name, fill="#1F2D3D", font=layer_font, anchor="ma")
        draw.multiline_text(
            (width / 2, top + 88),
            layer_text,
            fill="#2E3A46",
            font=text_font,
            anchor="ma",
            align="center",
            spacing=10,
        )
        if index < len(layer_specs) - 1:
            arrow_center_x = width // 2
            arrow_top = bottom + 8
            arrow_bottom = bottom + gap - 4
            draw.line((arrow_center_x, arrow_top, arrow_center_x, arrow_bottom), fill="#7A8794", width=5)
            draw.polygon(
                [
                    (arrow_center_x, arrow_bottom + 10),
                    (arrow_center_x - 12, arrow_bottom - 6),
                    (arrow_center_x + 12, arrow_bottom - 6),
                ],
                fill="#7A8794",
            )

    draw.rounded_rectangle((185, 865, 1095, 935), radius=20, fill="#FFFFFF", outline="#B0B8C2", width=3)
    draw.multiline_text(
        (width / 2, 898),
        "多源适配、分层缓存与前后端协同\n共同支撑热门行业识别与龙头股遴选闭环",
        fill="#3A4652",
        font=summary_font,
        anchor="mm",
        align="center",
        spacing=10,
    )

    image.save(output_path)
    return output_path


def insert_architecture_figure(doc: Document) -> None:
    anchor = find_last_paragraph_by_text(doc, "3.5 数据流程设计")
    break_para = insert_paragraph_before_element(anchor._element, anchor._parent)
    break_para.paragraph_format.first_line_indent = Pt(0)
    break_para.paragraph_format.space_before = Pt(0)
    break_para.paragraph_format.space_after = Pt(0)
    break_para.paragraph_format.keep_with_next = True

    diagram_table = insert_table_after(doc, break_para, rows=8, cols=1)
    diagram_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    diagram_table.autofit = False
    table_width = int((doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin) * 0.88)
    for row in diagram_table.rows:
        row.cells[0].width = Emu(table_width)
        _set_row_no_split(row)

    row_specs = [
        ("layer", "表现层", ["行业热力图 / 行业排行榜", "行业趋势面板 / 龙头股详情"], "D9EAFB", "5B8FD9", Cm(2.2)),
        ("arrow",),
        ("layer", "服务层", ["FastAPI 行业接口 / 参数校验", "响应封装 / 端点缓存"], "E3F4E8", "5BA36B", Cm(2.1)),
        ("arrow",),
        ("layer", "分析层", ["行业分析器 / 龙头股评分器", "快速评分 / 波动率估计"], "FDEBCF", "C58B1C", Cm(2.1)),
        ("arrow",),
        ("layer", "数据层", ["THS 主数据 / AKShare 增强", "Sina 与腾讯补充 / JSON 与 localStorage"], "EFE2FB", "8B62C4", Cm(2.2)),
        ("summary", ["多源适配、分层缓存与前后端协同", "共同支撑热门行业识别与龙头股遴选闭环"], Cm(1.55)),
    ]

    for row, spec in zip(diagram_table.rows, row_specs):
        kind = spec[0]
        if kind == "layer":
            _, title, lines, fill, border, height = spec
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_cell(row.cells[0], title, lines, fill, border)
        elif kind == "summary":
            _, lines, height = spec
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_summary_cell(row.cells[0], lines)
        else:
            row.height = Cm(0.45)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_arrow_cell(row.cells[0])

    caption = insert_paragraph_before_element(anchor._element, anchor._parent)
    caption.add_run("图 3.1 系统总体架构图")
    caption.paragraph_format.keep_with_next = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(0)
    for run in caption.runs:
        set_run_text_font(run, "宋体", Pt(12))

    source = insert_paragraph_before_element(anchor._element, anchor._parent)
    source.add_run(ARCHITECTURE_SOURCE_NOTE)
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.first_line_indent = Pt(0)
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(0)
    source.paragraph_format.keep_with_next = True
    for run in source.runs:
        set_run_text_font(run, "宋体", Pt(12), bold=True)


def insert_task_completion_table(doc: Document) -> None:
    heading = find_last_paragraph_by_text(doc, "6.5 对毕业设计目标的达成情况")
    completion_intro = "对照任务书要求，如表 6.4 所示，本文在理论分析、模型设计和系统实现三个层面均完成了核心目标。"
    intro_para = None
    next_element = heading._element.getnext()
    if next_element is not None and next_element.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        intro_para = Paragraph(next_element, heading._parent)
    if intro_para is None:
        intro_para = insert_paragraph_after(heading, completion_intro)
    elif not intro_para.text.strip():
        replace_paragraph_text(intro_para, completion_intro)
    elif "表 6.3" in intro_para.text:
        replace_paragraph_text(
            intro_para,
            intro_para.text.replace("表 6.3", "表 6.4"),
        )
    elif "表 6.4" not in intro_para.text:
        intro_para = insert_paragraph_after(heading, completion_intro)

    caption = insert_paragraph_after(intro_para, "表 6.4 毕业设计任务书目标达成情况")
    caption.paragraph_format.keep_with_next = True

    table = insert_table_after(doc, caption, rows=4, cols=4)
    rows = [
        ("任务书要求", "论文对应内容", "系统对应实现", "达成情况"),
        ("理解金融大数据挖掘相关技术与研究现状", "第 1 章与第 2 章完成文献综述和理论基础梳理", "明确行业子系统的数据来源、分析流程与关键技术", "已达成"),
        ("掌握行业识别与龙头股筛选方法", "第 4 章构建行业热度模型与龙头股评分模型", "实现行业横截面评分、波动率代理与双路径个股评分", "已达成"),
        ("编程设计软件系统跟踪重点行业和关键企业", "第 5 章与第 6 章说明系统实现与验证过程", "完成 THS-first、FastAPI、React、热力图快照与偏好持久化闭环", "已达成"),
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value


def append_thesis_expansion_content(doc: Document) -> None:
    expansions = [
        (
            "1.1 研究背景与意义",
            "1.2 国内外研究现状",
            [
                "从资本市场研究的角度看，行业热点并不是孤立产生的。有效市场理论、资本资产定价模型和多因子模型都说明，资产收益既受到市场整体风险影响，也会受到风格、规模、价值和行业暴露等因素影响。[11-18] 因此，本文关注的行业热度并不是简单的短期价格榜单，而是希望在一个可解释的多维框架中观察行业相对强弱。",
                "行业动量研究进一步说明，行业层面的信息扩散和资金再配置具有一定持续性。Moskowitz 和 Grinblatt 对行业动量的研究表明，行业因素能够解释相当一部分动量收益；Hou 关于行业信息扩散的研究也指出，行业之间可能存在领先和滞后关系。[19-23] 这类研究给本文的启发是：如果系统能稳定保留行业横截面结果，就可以为后续复盘行业轮动提供材料。",
                "同时，行为金融文献提醒研究者不能把短期强势直接等同于基本面改善。投资者情绪、过度反应与反应不足都可能在短期内推高某些行业或个股表现。[24-28] 因此，本系统在行业排序之外继续保留资金流、换手率、波动率代理和市值来源等字段，目的就是避免把一个总分解释得过于绝对。",
                "在毕业设计场景下，选题还需要兼顾可实现性。机器学习和深度学习可以用于资产定价与行情预测，但它们对样本规模、特征稳定性和回测约束要求较高。[29-34] 本文选择可解释的加权评分和工程化回退策略，是为了让系统输出能被答辩老师、普通研究者和后续维护者直接追问与复核。",
            ],
        ),
        (
            "1.2 国内外研究现状",
            "1.3 研究内容与技术路线",
            [
                "1.2.4 量化资产定价与机器学习研究现状",
                "传统资产定价研究从均值方差组合、CAPM、套利定价理论逐步发展到三因子、五因子和四因子模型，这些研究为股票横截面比较提供了基本框架。[11-18] 对本文来说，它们最大的价值在于提醒系统设计必须区分市场整体风险、行业暴露和个股特征，不能把所有波动都归因于单一热点。",
                "近年机器学习方法进一步把非线性特征、交互项和高维变量引入资产定价。Gu、Kelly 和 Xiu 的研究表明，机器学习模型能够在大量特征中捕捉更复杂的收益结构；Kelly、Pruitt 和 Su 则从协方差角度统一解释特征和收益之间的关系。[29-30] 这些方法对后续系统迭代很有启发，但在本科毕业设计里，如果缺乏严格回测和样本外验证，直接上复杂模型反而容易弱化可解释性。",
                "金融时间序列预测文献也给出了类似提醒。LSTM、深度学习综述和机器学习预测综述都说明，模型效果往往受数据窗口、交易成本、样本划分和评价指标影响。[31-34] 因此，本文没有把系统定位为收益预测器，而是定位为行业研究辅助工具：先把热点横截面和候选龙头股解释清楚，再为后续更严格的回测留下接口。",
                "1.2.5 金融文本、舆情与外部信息研究现状",
                "除行情和财务字段外，金融文本与舆情也被大量用于市场研究。Tetlock 证明媒体内容能够反映投资者情绪，Loughran 和 McDonald 则针对金融文本构建了更适合年报和公告分析的词典，后续研究进一步把社交媒体情绪纳入市场预测。[35-39] 这些工作说明，政策、新闻和投资者表达确实可能影响行业热度。",
                "不过，本项目当前主链路仍然以结构化行情、资金流、估值和财务字段为主。这样做是一个有意取舍：结构化数据更容易形成稳定接口，评分结果也更容易在页面上解释。文本和舆情数据更适合作为后续扩展方向，尤其适合用来解释某些行业突然升温的外部原因，而不是在当前阶段直接并入核心得分。",
                "1.2.6 大数据系统与金融工程落地研究现状",
                "大数据系统研究强调的并不只是存储规模，还包括数据多样性、处理延迟、质量控制和系统可维护性。[40-45] 这些问题在金融工程原型里同样明显：外部接口会变化，字段可能缺失，缓存策略会影响页面体验，历史快照又会影响论文样例能否复核。",
                "因此，本文的研究现状梳理最终落回一个工程问题：已有文献能提供行业轮动、资产定价、金融文本和大数据系统的理论背景，但毕业设计必须把这些背景转化成一套能运行、能截图、能解释、能复现的系统。本文后续章节正是沿着这个方向，把理论依据和项目实现放在同一条叙述线上。",
            ],
        ),
        (
            "1.3 研究内容与技术路线",
            "1.4 论文结构安排",
            [
                "具体来说，本文的研究内容可以分成四个层次。第一层是数据层，重点解决多源行业数据、成分股数据和财务行情字段如何进入统一结构；第二层是模型层，重点解释行业热度得分和龙头股得分怎样计算；第三层是服务层，说明接口缓存、回退和响应封装如何保证页面可用；第四层是展示层，通过热力图、排行榜和详情弹窗把结果交给用户。",
                "技术路线并不是从零搭建一个孤立实验，而是从已有项目中抽取与论文题目最相关的行业子系统，再反向梳理其数据链路和模型逻辑。这样做的好处是论文里每个重要判断都能回到项目文件、接口返回或固定快照上，不需要依赖单纯描述性假设。",
                "在分析方法上，本文采用文献梳理、系统分析、模型归纳、工程实现说明和样例验证相结合的方式。文献梳理用于确认行业轮动、多因子筛选和金融大数据处理的研究基础；系统分析用于界定需求和架构；模型归纳用于解释评分逻辑；样例验证则使用项目保留的热力图快照和龙头股样本说明系统输出。",
                "在验证方式上，本文并不把单次样例解释为严格投资结论，而是把它作为系统可用性证据。真正的完整投资验证还需要更长历史窗口、交易成本假设、换仓规则和样本外检验，这些内容已经超出当前毕业设计主线，但会在结论和不足部分作为后续方向提出。[18][27][29-34]",
            ],
        ),
        (
            "2.1 金融大数据的特征",
            "2.2 多源数据采集与清洗",
            [
                "金融大数据还具有很强的时序性。行业热度不是一个静态标签，而是在不同统计窗口里不断变化的横截面状态。某个行业在 1 日窗口里可能明显升温，在 5 日或 10 日窗口里却已经回落；如果系统不保存窗口信息和快照时间，后续就很难解释同一行业在不同页面中为什么排序不同。",
                "另一个特征是噪声和缺失经常同时出现。市场数据的短时跳动、接口限流、字段更新延迟和行业名称不统一，都会让原始数据带有不稳定性。大数据系统研究中强调的数据质量、真实性和处理延迟，在这里不是抽象概念，而是直接影响热力图能否加载、排行榜是否合理、个股详情是否完整的问题。[40-45]",
                "金融数据还存在明显的尺度差异。市值可能以亿元甚至万亿元计量，涨跌幅通常以百分比计量，成交额和资金流又有不同单位。如果不做标准化，尺度较大的变量会在加权过程中天然占优。本文后续在行业评分和个股评分中都使用标准化、区间裁剪或中性值处理，就是为了尽量减少尺度差异造成的误导。",
                "从研究可解释性的角度看，金融大数据还有一个容易被忽视的特点：用户通常不会只接受一个最终分数。研究者会继续追问这个分数来自涨幅、资金、成交活跃度还是波动代理。因而系统必须在输出总分的同时保留关键字段，这也是热力图历史快照同时保存 value、total_score、moneyFlow、turnoverRate 和市值来源的原因。",
                "因此，本文所说的金融大数据并不只是“大量行情数据”的同义词，而是包含多源、异构、时变、噪声、缺失和解释需求的一组综合约束。只有把这些约束处理好，后面的行业识别与龙头股遴选才有稳定基础。",
            ],
        ),
        (
            "2.2 多源数据采集与清洗",
            "2.3 热门行业识别的理论基础",
            [
                "多源数据采集首先要解决来源分工问题。对本系统而言，同花顺接口更适合作为行业主数据入口，AKShare 适合补充行业列表、成分股和历史行情，新浪与腾讯则在部分字段缺失时承担补位角色。这样的组织方式接近数据工程里的主源加补源模式，可以减少任一单点失败对整体页面的影响。[9-10][52-55]",
                "其次要解决字段映射问题。行业名称、行业代码、股票代码、市值字段、资金流字段和涨跌幅字段在不同来源中的命名并不完全一致。项目中将这些字段统一到 industry_name、change_pct、flow_strength、avg_volume、turnover_rate 等内部字段，实际上就是把外部接口差异消化在适配层，而不是让分析器和前端分别处理。",
                "缺失值处理是第三个关键环节。对于行业层面，如果暂时拿不到历史波动率，系统会使用振幅、换手率或涨跌幅代理值；对于个股层面，如果快照场景下拿不到 ROE 或成长字段，快速评分会采用中性值，避免列表直接中断。这样的处理不代表数据完美，而是保证系统在真实环境中能够连续输出。",
                "异常值处理也不能忽略。金融数据里偶发的极端涨跌、单位异常或接口返回错误，都会影响综合得分。本文采用的评分逻辑会通过横截面标准化和分数压缩减少极端值对展示结果的冲击；在后续更严格的研究中，还可以继续引入分位数缩尾和异常源标记等方法。",
                "最后，多源清洗的结果需要面向复核。论文中使用的固定热力图快照和龙头股样本，实际上就是把采集清洗后的关键结果沉淀下来，便于在不同时间重新打开论文或代码时仍能看到同一组样例。这一点对毕业设计尤其重要，因为答辩材料必须稳定，而不能随着当天行情自动漂移。",
            ],
        ),
        (
            "2.3 热门行业识别的理论基础",
            "2.4 龙头股评价的理论基础",
            [
                "行业热点识别可以从三个角度理解。第一是资产定价角度，行业收益可能包含共同风险暴露和风格因子影响；第二是动量角度，强势行业可能在一段时间内继续受到资金关注；第三是行为金融角度，热点也可能包含情绪放大和过度反应。[11-28] 这三个角度共同说明，热门行业不能只靠单一指标判断。",
                "动量研究为本文提供了最直接的理论支撑。Jegadeesh 和 Titman 证明了股票层面的中期动量现象，Moskowitz 和 Grinblatt 则进一步讨论了行业动量。[19][21] 这些研究说明，行业层面的相对强弱具有研究意义，也支持本文按 1 日、5 日、10 日等窗口观察行业状态。",
                "行业信息扩散研究则解释了为什么行业内部和行业之间可能存在领先滞后关系。Hou 的研究表明，行业信息可能通过上下游或相关行业逐步反映到股票价格中。[23] 因此，系统只展示某个行业当天涨幅是不够的，还需要保留趋势、成分股和龙头股详情，让用户继续沿行业线索下钻。",
                "行为金融研究提醒本文在解释热点时保持克制。投资者情绪和市场过度反应可能造成短期强势，但这种强势未必转化为可持续收益。[24-28] 因此，本文在第六章分析样例结果时，只将其解释为样本时点的横截面状态，而不把它写成对未来收益的确定性预测。",
                "由此可见，热门行业识别的理论基础并不是单一模型，而是资产定价、行业动量、信息扩散和行为金融共同构成的解释框架。系统里的动量、资金流、活跃度和波动率代理，正是对这一框架的工程化简化。",
                "这种简化有一个明显优点：每个因子都能在页面或快照中找到对应字段，便于解释和复核。它也有局限：权重仍然主要依据工程经验，尚未通过大样本回测或机器学习自动学习。这个边界将在后文系统不足部分继续说明。",
            ],
        ),
        (
            "2.4 龙头股评价的理论基础",
            "2.5 系统关键技术",
            [
                "从资产定价角度看，个股是否能代表行业，首先要看它是否具有稳定的规模、盈利和风险暴露。市值、估值、ROE、收入增长和利润增长等指标并不是为了给企业贴标签，而是为了判断该公司是否具备行业代表性。[11-18] 如果一个股票只有短期涨幅而缺乏基本面支撑，它更适合进入热点观察，而不一定适合作为核心龙头。",
                "从机器学习资产定价研究看，大量公司特征确实可能与未来收益有关，但这些特征在不同市场和样本期里稳定性并不完全相同。[29-30] 因此，本文没有追求尽可能多的特征，而是选择规模、估值、盈利、成长、动量和活跃度六个容易解释、也较容易从项目中获得的维度。",
                "从市场情绪角度看，热点个股和核心资产往往并不重合。社交媒体、新闻情绪和短期资金关注可能迅速推高某些股票的热度，但这种热度未必意味着其长期代表性更强。[35-39] 项目中将榜单拆成 core 和 hot 两类，正是为了保留这两种不同语义。",
                "因此，龙头股评价应当同时考虑“稳定代表性”和“短期市场确认”。前者更依赖市值、估值、盈利和成长，后者更依赖涨跌幅、成交额、换手率和资金承接。本文的评分模型并不宣称能给出唯一正确答案，而是给用户一个透明的排序依据。",
                "在系统实现中，完整评分链路服务于解释，快速评分链路服务于响应速度。两者共用 raw_data 结构，能够减少字段含义不一致的问题；但当财务字段缺失时，快速链路采用中性处理，这也意味着 hot 榜单不能被简单理解为完整基本面评分。",
            ],
        ),
        (
            "2.5 系统关键技术",
            "系统需求分析与总体设计",
            [
                "本系统的关键技术可以概括为 Python 数据处理、机器学习辅助分析、Web 接口服务和前端可视化四类。Python 生态中的 pandas、NumPy、SciPy、Matplotlib 和 scikit-learn 已经形成较成熟的数据分析工具链，适合处理结构化行情、横截面标准化和聚类辅助分析。[9-10][52-55]",
                "在算法层面，系统主要使用加权评分和 K-Means 聚类。K-Means 本身是一种经典聚类方法，后续研究又提出了 k-means++ 初始化和轮廓系数评估等改进思路。[46-51] 在本文里，聚类不是为了替代行业得分，而是用于观察若干行业是否形成相近热点簇，属于辅助解释工具。",
                "在后端服务层，FastAPI 提供了较轻量的接口组织方式，适合把 Python 分析能力封装成前端可调用的 REST 接口。[58] REST 架构思想强调资源、统一接口和无状态交互，这与本项目中热门行业、热力图、趋势、龙头股详情等接口设计具有一致性。[56]",
                "在前端展示层，React 更适合组织状态复杂、交互频繁的单页应用。[59] 行业热度页面需要同时管理热力图、排行榜、详情弹窗、观察行业、时间窗口和偏好配置，如果只用静态报表，很难支撑这样的交互闭环。",
                "在数据源方面，AKShare 为国内金融数据研究提供了便利入口，适合补充行业、行情和财务字段。[60] 但任何开源或第三方数据源都可能遇到接口变动、限流或字段缺失，因此项目没有把单一来源当作绝对可靠前提，而是在适配层保留了多源回退。",
                "综合来看，本文采用的关键技术并不是彼此孤立的工具清单，而是服务于同一条研究链路：用 Python 处理数据，用 FastAPI 输出接口，用 React 呈现交互，用历史快照保证样例可复核。",
            ],
        ),
    ]

    expansions.extend(
        [
            (
                "3.1 设计目标",
                "3.2 功能需求分析",
                [
                    "系统设计目标首先是可用性。行业热度页必须在外部数据源短时波动时仍尽量给出结果，否则研究者每次打开页面都可能被接口状态打断。因此，系统需要缓存、回退和降级路径，而不是只在理想网络环境下运行。",
                    "第二个目标是可解释性。行业综合得分和龙头股综合得分都应当能拆回具体字段，用户至少要知道得分来自价格强弱、资金流、活跃度、波动率，还是来自市值、估值、盈利和成长。这个目标与行为金融和机器学习资产定价研究中的可解释性要求是一致的。[24-39]",
                    "第三个目标是可复核性。论文中的表格和截图不能依赖某一秒钟的临时行情，因此系统需要保存热力图历史快照和龙头股样本，使结果章节能够固定引用同一组数据。这个目标也是后续答辩和归档最实际的需求。",
                ],
            ),
            (
                "3.3 非功能需求分析",
                "3.4 系统总体架构设计",
                [
                    "性能需求主要体现在首屏加载和详情下钻两个场景。首屏需要尽快给出热门行业、热力图和排行榜，因此不适合在一开始就为每个行业拉取完整成分股和全部财务字段；详情下钻则更强调解释完整，可以接受更多补充请求。系统中的快速评分和完整评分正是围绕这两个场景拆开的。",
                    "稳定性需求主要来自外部接口的不确定性。金融数据接口经常会因为网络、限流、字段调整或节假日数据缺失而返回异常结果。系统不能把这些异常简单暴露给前端，而需要用缓存、代理字段和过期结果维持页面连续性。[40-45]",
                    "可维护性需求则体现在代码边界上。数据适配、评分计算、接口响应和前端展示如果互相缠在一起，后续任何字段变化都可能引起连锁修改。因此，项目把适配器、分析器、评分器、路由和前端组件分开，使每一层承担相对清晰的职责。",
                    "安全性和合规性在本课题中主要表现为边界控制。系统不处理真实交易指令，也不提供自动下单能力；论文中的结果只用于研究展示，不构成投资建议。这一边界需要在论文解释中保持明确，避免把毕业设计原型描述成生产级投资系统。",
                ],
            ),
            (
                "3.5 数据流程设计",
                "3.6 存储设计",
                [
                    "从一次页面请求看，数据流程可以分成请求、命中缓存、拉取数据、清洗映射、评分计算和响应封装几个阶段。用户打开行业页时，前端先请求热力图和热门行业列表；后端检查短期缓存；如果缓存不可用，再调用行业分析器；分析器继续向适配器请求行业摘要、资金流和必要的成分股字段。",
                    "当数据进入分析层后，系统会先对行业名称和数值字段进行整理，再计算动量、资金、活跃度和波动率代理。龙头股链路则会在行业成分股基础上继续补充市值、估值、财务和成交信息，并按 core 与 hot 两类语义组织结果。",
                    "响应封装阶段同样重要。前端不应直接面对不同数据源的原始返回，而应接收结构稳定的字段，例如 total_score、change_pct、moneyFlow、turnoverRate、leaders、trend 和 explanation。这样前端组件才能稳定复用，也便于测试脚本检查接口契约。",
                    "数据流程最后还会把部分结果沉淀为快照。热力图历史快照用于回看行业横截面，偏好配置用于保存观察行业和页面状态，财务缓存用于减少重复请求。这些文件共同构成本文样例分析的证据来源。",
                ],
            ),
            (
                "3.6 存储设计",
                "热门行业识别与龙头股遴选模型设计",
                [
                    "文件化存储还有一个好处，是非常适合毕业设计阶段的复核。老师或答辩委员如果追问第六章表格里的数据来源，论文可以直接说明其来自固定快照文件，而不是当天重新计算出的临时结果。对一个本地运行的研究原型来说，这种透明度比复杂数据库更重要。",
                    "当然，文件存储并不意味着可以忽略数据规模。热力图历史记录如果无限增长，会影响读取速度和版本管理；财务缓存如果长期不清理，也可能出现过期字段影响评分的问题。因此，项目对历史快照数量和缓存有效期都做了控制，并在论文中把它们作为工程取舍说明。",
                    "如果未来系统要走向多人协作或在线部署，存储层可以进一步迁移到数据库、对象存储或消息队列体系。但在当前本科毕业设计范围内，文件快照加浏览器状态已经能够满足样例复核、页面回放和偏好保存三类核心需求。",
                ],
            ),
            (
                "4.1 数据预处理与指标标准化",
                "4.2 热门行业识别模型设计",
                [
                    "数据预处理的第一步是数值化。来自接口的涨跌幅、资金流、成交额和市值字段可能带有百分号、单位字符或空字符串，必须统一转换成浮点数后才能参与计算。对于无法转换的字段，系统会结合字段含义选择中性值、缺失标记或回退值。",
                    "第二步是横截面标准化。行业热度比较关注的是同一时点不同行业之间的相对位置，因此标准化应当在同一横截面内完成，而不是把不同日期的数据混在一起处理。这样做可以减少市场整体涨跌对行业排序的干扰。",
                    "第三步是异常处理。极端资金流、极端换手率和短期异常涨跌都可能影响综合得分。本文采用的处理方式比较朴素，主要依靠代理字段、横截面标准化和分数压缩；如果后续做更严格研究，可以继续引入缩尾、稳健标准化和异常源标注。",
                    "个股评分部分还需要处理财务字段的低频更新问题。ROE、营收同比和利润同比不是每天变化的字段，如果要求它们和日行情完全同步，系统就会频繁遇到缺失。因此，完整评分使用缓存补充财务字段，快速评分则在必要时使用中性分，保证候选列表可以先输出。[29-34]",
                ],
            ),
            (
                "4.2 热门行业识别模型设计",
                "4.3 行业波动率估计与聚类辅助分析",
                [
                    "行业热度评分中的动量因子主要回答“近期是否走强”，资金因子回答“是否有资金承接”，活跃度因子回答“交易是否足够活跃”，波动率因子则用于约束过度剧烈的短期波动。四类信号组合起来，才能比较接近研究者实际看行业时的判断过程。[19-28]",
                    "权重设计采用动量和资金各 0.35、活跃度 0.15、波动率约束 0.15 的结构。这样的设计并不意味着该权重在所有市场环境下最优，而是为了在毕业设计阶段保证结果直观、口径稳定、容易解释。后续如果要进一步研究，可以把权重优化作为独立问题处理。",
                    "行业热度得分最后会被压缩到展示区间，这是前端体验上的必要处理。如果直接展示标准化后的原始分数，用户很容易看到过多负值或极端值；而压缩到 20 至 95 的区间后，热力图和排行榜更适合进行横截面比较。",
                    "需要强调的是，得分越高并不等于未来收益越高。它只说明在当前统计窗口里，该行业在模型关注的几个维度上处于相对靠前位置。这个解释边界与动量和数据窥探相关研究中的谨慎态度是一致的。[18][27]",
                ],
            ),
            (
                "4.3 行业波动率估计与聚类辅助分析",
                "4.4 龙头股综合评分模型设计",
                [
                    "波动率估计在系统中承担的是风险约束功能。一个行业如果短期涨幅很高但波动也很剧烈，综合热度不应被无条件抬高。真实行业指数历史收益率是更直接的波动来源，但在接口不稳定或首屏加载场景下，代理波动率可以先保证页面输出。",
                    "代理波动率的设计需要保持克制。振幅、换手率和涨跌幅都只能间接反映波动，并不等同于标准差意义上的历史波动率。因此，本文把它称为代理值，并在第六章说明样例结果是横截面状态而不是严格风险预测。",
                    "聚类分析使用 K-Means 和轮廓系数，主要是为了观察行业之间的相似结构。[46-51] 如果若干行业在涨跌幅、资金流和活跃度上接近，聚类结果可以帮助用户判断热点是否成片扩散；如果聚类不稳定，也可以提示当前市场结构比较分散。",
                    "在系统展示中，聚类不直接替代热度得分。原因是聚类结果更适合描述相似性，而排行榜需要给出相对顺序。本文将聚类定位为辅助分析，既能保留它的解释价值，也避免把无监督分类结果误用为投资排序。",
                ],
            ),
            (
                "4.4 龙头股综合评分模型设计",
                "4.5 快速评分机制设计",
                [
                    "龙头股综合评分中的市值规模维度反映公司在行业中的体量。体量较大的公司通常更容易成为行业研究和机构配置的代表，但市值过大也未必意味着成长性更好，因此系统只把它作为六个维度之一。",
                    "估值维度主要用于避免把过高估值的短期热门股无条件排在前面。市盈率处于合理区间时得分更高，过低或过高都需要谨慎解释。盈利能力和成长性则分别通过 ROE、营收同比和利润同比反映公司经营质量。[11-18]",
                    "动量和活跃度维度用于捕捉市场确认。一个公司基本面再好，如果短期几乎没有成交和关注，也未必适合作为当前热点行业中的候选；反过来，只有动量而缺乏基本面支撑，也更适合进入 hot 观察榜而不是 core 核心榜。",
                    "完整评分链路的意义在于解释，而不仅是排序。前端详情页展示维度拆解后，用户可以看到某只股票是因为规模、盈利、成长还是短期动量获得较高得分，这比单一总分更符合研究场景。",
                    "与机器学习选股模型相比，本文的六维评分明显更简单，但它的优势是透明。对于毕业设计而言，透明性和可复核性比追求复杂模型更重要；复杂模型可以作为后续研究方向，在更充分的数据和回测条件下再引入。[29-34]",
                ],
            ),
            (
                "4.5 快速评分机制设计",
                "系统实现",
                [
                    "快速评分机制的核心是承认页面有不同层级的信息需求。行业页列表只需要先筛出候选集合，详情页才需要展开完整解释。如果所有计算都按详情页标准执行，首屏体验会变差；如果所有计算都按快路径执行，结果解释又会不够充分。",
                    "因此，系统将快路径和完整路径并行保留。快路径优先使用已有快照字段，适合 core/hot 榜单的快速展示；完整路径补充估值、财务和价格序列，适合详情弹窗中的维度拆解。两条路径共享 raw_data，可以减少字段含义不一致的问题。",
                    "快路径中采用中性值并不是为了掩盖缺失，而是为了避免少数字段缺失导致整个候选列表中断。对于用户来说，先看到一组可解释的候选，再进入详情页补充信息，比等待所有字段完全齐备更符合页面使用习惯。",
                    "从模型边界看，快速评分结果不应被解释为完整基本面评价。论文中保留 core 与 hot 两类榜单说明，就是为了提醒用户区分长期代表性和短期市场冲击，避免把不同语义的榜单混为一谈。",
                ],
            ),
            (
                "5.1 数据提供器与回退机制实现",
                "5.2 行业分析模块实现",
                [
                    "在代码组织上，适配层承担了很多看起来琐碎但非常关键的工作。例如行业节点需要映射，股票名称需要转换成代码，部分字段需要从字符串转成数值，某些接口失败后还要尝试代理节点或缓存结果。这些工作如果散落到前端或评分器里，系统会很难维护。",
                    "适配器还承担了“尽量不让页面空掉”的职责。当主源短暂不可用时，系统会尝试补源或过期缓存；当某些行业成分股无法完整获取时，系统也会返回降级结果，并在后续详情分析中继续补充。这种策略不能保证每次都拿到最完整数据，但能提高研究页面的连续性。",
                    "从系统实现的角度看，数据提供器是论文中多源数据理论和实际代码之间的连接点。它把文献中常说的数据融合、质量控制和可用性问题，具体落到了名称映射、字段统一、缓存回退和异常处理这些操作上。[40-45][52-60]",
                ],
            ),
            (
                "5.2 行业分析模块实现",
                "5.3 龙头股评分模块实现",
                [
                    "行业分析器内部最重要的设计，是区分首屏分析和详情分析。首屏分析优先完成行业排序、热力图和关键指标展示；详情分析再补行业趋势、成分股覆盖、历史波动和解释字段。这个顺序和用户实际浏览页面的路径一致。",
                    "为了保证结果可解释，行业分析器在输出 total_score 的同时保留了原始或半加工字段。前端可以继续展示涨跌幅、资金流、换手率、市值来源和波动代理，测试脚本也可以检查这些字段是否存在。这样一来，综合分不是孤立数字，而是一个可拆解的研究结果。",
                    "缓存策略同样嵌入在行业分析模块周围。对于短期重复请求，缓存可以减少外部接口压力；对于数据源短时失败，过期缓存能维持页面可用。系统没有把缓存理解为性能小技巧，而是把它当作行业研究原型稳定运行的一部分。",
                ],
            ),
            (
                "5.3 龙头股评分模块实现",
                "5.4 后端接口实现",
                [
                    "龙头股评分模块的数据输入比行业评分更复杂。行业评分主要面对行业级字段，而个股评分还需要市值、PE、ROE、营收同比、利润同比、涨跌幅、成交额和换手率等信息。不同字段的更新频率不同，因此模块必须通过缓存和中性值处理保持评分流程连续。",
                    "在完整评分链路中，系统会尽量补齐估值和财务字段，再按六维权重生成综合得分。这个链路更适合详情页，因为用户在详情页中会关注一只股票为什么排在前面，以及它在哪些维度上占优。",
                    "在快速评分链路中，系统更关注候选集合的快速形成。hot 榜单强调短期涨幅和资金承接，core 榜单强调相对稳定的综合质量。两类榜单共同存在，使系统既能回应热点交易线索，也能保留行业代表性筛选逻辑。",
                ],
            ),
            (
                "5.4 后端接口实现",
                "5.5 前端可视化实现",
                [
                    "后端接口设计还需要考虑前端调用顺序。行业页通常会并行请求热力图、排行榜和龙头股摘要，如果接口返回结构不稳定，前端组件就需要写大量兼容逻辑。因此，后端尽量把字段整理成统一响应，让前端专注展示和交互。",
                    "接口缓存的粒度也需要区分。热门行业列表和热力图适合短时间缓存，避免用户切换视图时重复计算；行业详情和龙头股详情则需要根据行业名称、统计窗口和榜单类型区分缓存键。这样的设计能在性能和结果准确性之间取得平衡。",
                    "接口层还承担了错误翻译的职责。外部数据源的异常信息通常不适合直接返回前端，后端需要把它转成可理解的状态、降级结果或空数据提示。这样用户看到的是研究页面的局部不可用，而不是底层接口错误。",
                ],
            ),
            (
                "5.5 前端可视化实现",
                "5.6 偏好配置与持续跟踪实现",
                [
                    "从交互设计看，热力图负责快速建立市场全局印象。颜色维度让用户看到强弱，尺寸维度帮助用户理解市值或成交背景，统计窗口则支持比较短期和稍长周期的变化。对于行业研究而言，这比单纯表格更容易发现结构性热点。",
                    "排行榜负责把热力图中的视觉印象转成可排序、可筛选的列表。用户可以按综合得分、涨跌幅、资金流或波动条件缩小范围，再选择某个行业继续查看。这个过程对应研究中的“先粗筛、再细看”。",
                    "详情弹窗则负责解释。行业详情展示趋势、成分股和覆盖信息，龙头股详情展示得分拆解、价格走势和财务字段。这样，用户从热力图进入详情后，能沿着同一条数据链路继续追问，而不是被迫切换到另一套工具。",
                    "偏好配置和历史快照让前端页面具备持续研究能力。观察列表保存用户关心的行业，回放功能帮助用户比较不同时间点的横截面状态，本地状态则减少重复配置。这些能力虽然不是评分模型本身，却会明显影响系统是否像一个可持续使用的研究原型。",
                ],
            ),
            (
                "5.6 偏好配置与持续跟踪实现",
                "系统测试与结果分析",
                [
                    "持续跟踪功能的意义在于把一次性结果变成可回看的研究材料。行业热度今天排在前列，并不意味着明天仍然如此；只有保留历史快照，用户才能观察某个行业是短暂冲高、持续走强，还是在资金流和活跃度上出现分化。",
                    "偏好配置还降低了重复研究成本。用户每次打开页面时，如果都要重新选择观察行业、统计窗口和阈值，研究过程会被大量重复操作打断。按 profile 保存配置后，同一套研究视图可以反复使用，也方便后续扩展到更多研究场景。",
                    "从论文角度看，持续跟踪功能也提高了样例分析的可信度。第六章固定引用的快照和龙头股样本，正是通过这种持久化能力保存下来的。如果没有这类记录，论文结果会随着数据源更新不断变化，不利于提交和答辩。",
                ],
            ),
            (
                "6.1 测试环境",
                "6.2 功能测试",
                [
                    "测试环境选择本地运行，是因为毕业设计需要同时验证后端分析、前端展示和文档截图。后端负责接口和评分，前端负责交互与渲染，历史快照负责样例复核。三者必须放在同一环境里一起看，才能判断系统是否形成完整闭环。",
                    "测试过程中需要区分功能正确性和投资有效性。功能正确性关注接口能否返回、字段是否完整、页面能否交互、快照是否可回看；投资有效性则需要长期回测和交易成本假设。本文的测试重点是前者，后者作为未来研究方向保留。",
                    "为了避免样例随行情变化，论文使用固定热力图快照和固定龙头股案例。这样即使未来重新生成论文，表 6.2 和龙头股案例也能保持一致，便于老师检查论文、项目和数据文件之间的对应关系。",
                ],
            ),
            (
                "6.2 功能测试",
                "6.3 热力图快照结果分析",
                [
                    "功能测试首先检查分析层。行业评分需要在字段完整时输出合理得分，在字段缺失时走回退路径，在波动率拿不到时使用代理值。龙头股评分则需要分别验证完整评分和快速评分，保证 core 与 hot 两类榜单语义不混乱。",
                    "接口测试主要检查返回结构。热门行业、热力图、趋势、龙头股列表和详情接口，都需要返回前端能够直接消费的字段。如果后端某次改动导致字段名称变化，前端页面即使能加载，也可能出现空图、空榜单或详情缺失。",
                    "前端测试关注交互闭环。用户切换热力图周期、搜索行业、打开详情、查看龙头股和切换偏好配置时，页面状态应当保持一致。对于毕业设计展示来说，这类测试比单独跑一个函数更能说明系统已经具备可操作性。",
                    "测试结论中的“满足预期”并不是生产级质量保证，而是说明毕业设计主链路已经可运行。若要进一步走向真实投资研究，还需要增加长周期数据校验、异常行情回测、压力测试和更严格的接口监控。",
                ],
            ),
            (
                "6.3 热力图快照结果分析",
                "6.4 系统特征与不足分析",
                [
                    "从表 6.2 的样例可以进一步看到，行业综合得分和单一涨跌幅之间并不完全一致。某些行业涨幅靠前但资金流为负，某些行业涨幅稍低但资金承接更好，这正是多因子评分相对单指标排序的意义。",
                    "样例中电子、通信和能源金属等方向位于前列，说明系统能够在固定窗口内捕捉阶段性强势板块。但这里仍然需要强调，该结论只对应样本时点。市场环境、政策预期、行业景气度和资金风格变化之后，排序结果都可能发生明显变化。",
                    "龙头股双榜单样例也说明，core 与 hot 的结果并不完全相同。core 更接近行业代表性公司筛选，hot 更接近短期冲击信号筛选。两者同时展示，可以帮助用户区分“长期值得研究的公司”和“短期被市场推到前台的公司”。",
                    "如果后续要把样例分析推进到投资有效性验证，就需要把多个快照按时间顺序连接起来，构建行业信号形成、持有、换仓和收益评价流程。这部分工作已经超出当前论文篇幅，但系统保存快照的方式为后续研究留下了基础。",
                ],
            ),
            (
                "6.4 系统特征与不足分析",
                "6.5 对毕业设计目标的达成情况",
                [
                    "系统的第一个特征是主链路比较完整。从数据适配到行业评分，从热力图展示到龙头股详情，用户可以沿同一条路径完成“看行业、选行业、看个股、查原因”的过程。这个闭环是本文区别于单纯算法实验的地方。",
                    "第二个特征是结果解释性较强。行业得分可以拆回动量、资金、活跃度和波动代理，个股得分可以拆回规模、估值、盈利、成长、动量和活跃度。即使用户不同意当前权重，也能理解系统为什么给出这样的排序。",
                    "第三个特征是工程上保留了现实妥协。接口回退、过期缓存、中性值、快照保存和偏好持久化都不是理论模型里的核心概念，但它们决定了系统能否在真实数据环境中持续运行。",
                    "不足也同样明确。当前系统没有纳入政策文本和新闻舆情，对产业链上下游关系的刻画也比较弱；权重尚未通过大样本学习或优化；行业热度与未来收益之间还缺少完整回测。这些问题会限制系统作为投资决策工具的严谨性。",
                    "不过，作为本科毕业设计，当前系统已经完成了从问题提出到系统展示的主要目标。它更适合作为一个可解释的行业研究辅助原型，而不是直接用于自动交易或收益承诺。",
                ],
            ),
            (
                "结 论",
                "参考文献",
                [
                    "从论文扩展后的整体结构看，本文已经把金融大数据、行业轮动、资产定价、机器学习辅助分析和系统工程实现放在同一研究框架下。前两章回答为什么要这样做，第三章回答系统要满足什么需求，第四章解释模型怎样计算，第五章说明代码如何落地，第六章用固定快照和测试结果验证系统主链路。",
                    "本文的主要贡献可以概括为三点。第一，围绕行业热度和龙头股遴选建立了可解释的多维评分框架；第二，把 THS-first 多源适配、缓存回退、历史快照和前端可视化连接成可运行系统；第三，通过固定样例避免论文结果随行情变化而漂移，提高了毕业设计材料的可复核性。",
                    "同时，本文也保持了研究边界。系统当前主要支持行业研究和候选股观察，不能直接等同于投资建议；第六章结果分析使用的是固定样例，不是完整回测；权重和代理字段仍有经验性。把这些边界写清楚，有助于让论文结论更稳，也有助于后续迭代继续推进。",
                ],
            ),
        ]
    )

    expansions.extend(
        [
            (
                "3.4 系统总体架构设计",
                "3.5 数据流程设计",
                [
                    "从软件工程角度看，四层结构还对应了不同的变更频率。前端交互会随着展示需求变化较快，接口层主要跟随页面数据结构调整，分析层受模型和字段口径影响，数据层则最容易受到外部数据源变化影响。把这些变化点拆开，可以降低后续维护成本。[56-57]",
                    "系统没有把行业分析做成完全离线的批处理流程，是因为用户打开页面时需要看到较新的横截面状态；也没有把它做成高频流式系统，是因为毕业设计关注的是行业研究辅助，而不是毫秒级交易。当前 REST 加缓存的方式，正好处在这两个极端之间。",
                    "另外，架构设计还要考虑论文材料的生成。截图、固定快照、附录材料和主论文表格都依赖系统能够稳定重复运行。如果架构过于依赖临时状态，论文成稿就很容易和项目当前状态脱节；因此，本系统在设计时特别保留了历史快照和本地文件证据链。",
                ],
            ),
            (
                "4.2 热门行业识别模型设计",
                "4.3 行业波动率估计与聚类辅助分析",
                [
                    "在实际解释行业得分时，本文更强调因子方向而不是绝对数值。动量和资金流为正向信号，活跃度为辅助确认信号，波动率为约束信号。这样设置的直觉是：一个行业既要有价格表现，也要有交易和资金层面的配合，同时不能完全由高波动推高。",
                    "分数压缩函数还承担了降低展示噪声的作用。横截面标准化后的原始分数可能在样本很少或极端值较多时波动很大，压缩后虽然牺牲了一部分精确度，但换来了更稳定的可视化效果。对于以研究展示为目标的系统来说，这是一种合理折中。",
                    "如果后续要把该模型推进到严谨策略研究，可以增加敏感性分析，分别测试权重变化、窗口长度变化和波动率代理变化对行业排序的影响。这样可以判断当前得分是否对某个参数过于敏感，也能为权重优化提供依据。[18][27]",
                ],
            ),
            (
                "5.4 后端接口实现",
                "5.5 前端可视化实现",
                [
                    "接口实现还需要处理参数边界。行业名称、统计窗口、榜单类型和回放时间都可能由前端传入，如果后端不做校验，就容易出现空结果、错误缓存命中或异常查询。项目中的接口层因此会统一处理参数默认值和响应结构。",
                    "对于热力图历史接口，返回数据既要满足页面回放，也要控制体积。系统不会无限制返回全部历史，而是通过记录数和文件大小限制保持快照可控。这样的设计虽然简单，但对于本地研究原型已经足够。",
                    "对于龙头股详情接口，重点不只是返回一个股票列表，而是把综合得分、维度拆解、价格走势和原始字段一起组织出来。这样前端弹窗才能解释候选股票被选中的原因，也能让论文中的图 5.3 与后端能力对应起来。",
                ],
            ),
            (
                "6.5 对毕业设计目标的达成情况",
                "结 论",
                [
                    "从文档和系统完成度看，扩展后的论文已经补充了理论依据、模型解释和工程实现说明。参考文献覆盖大数据管理、资产定价、行业动量、金融文本、机器学习、聚类算法和 Web 系统实现等方向，能够更好支撑题目中的“基于大数据”和“龙头股遴选”两个关键词。",
                    "从后续扩展看，系统仍可继续加入舆情、政策文本、产业链关系和回测验证。当前论文没有把这些扩展写成已完成功能，而是放在不足和未来工作中，结论更稳，也更符合项目真实状态。",
                ],
            ),
        ]
    )

    expansions.extend(
        [
            (
                "1.3 研究内容与技术路线",
                "1.4 论文结构安排",
                [
                    "结合旧稿中保留下来的技术路线描述，本文还需要强调研究闭环本身。系统不是先给出一个抽象指标体系，再另外寻找页面展示，而是从行业热度页面的真实使用路径反推写作结构：用户先看到市场横截面，再选择行业，再进入个股候选，最后回到评分字段和接口链路解释原因。",
                    "因此，论文中的章节安排也尽量避免脱离工程实际。数据获取、字段统一、指标计算、综合评分、结果展示和结果验证这六个步骤，对应项目中的适配器、分析器、评分器、接口、前端组件和历史快照文件。这样写可以让论文、代码和答辩演示形成同一套证据链。",
                ],
            ),
            (
                "3.2 功能需求分析",
                "3.3 非功能需求分析",
                [
                    "功能需求首先体现在行业全局观察。系统需要在一个页面中给出热门行业列表、热力图和基础市场概况，使用户不必在多个行情终端之间反复切换。对于毕业设计而言，这一需求对应的是“识别热门行业”的入口能力，也是后续龙头股筛选的前置条件。",
                    "第二类需求是行业下钻分析。用户选中某个行业后，需要继续看到行业趋势、成分股覆盖、涨跌分布和龙头股候选，而不是只停留在一个行业名称和综合得分上。只有支持下钻，行业热度结果才真正能转化为后续研究线索。",
                    "第三类需求是龙头股解释。系统不仅要返回候选股票，还要展示为什么这些股票排在前面，包括市值、估值、盈利、成长、动量和活跃度等维度。这样既方便用户判断结果是否可信，也方便论文把模型设计和页面展示对应起来。",
                    "第四类需求是研究状态保存。观察行业、热力图回放、统计窗口和部分筛选条件都需要在本地保留下来，否则用户每次重新打开页面都要从头配置。偏好持久化与历史快照并不直接改变模型结果，但会明显影响系统是否具备持续研究价值。",
                    "第五类需求是结果可导向复核。论文中的表格、截图和样例结论需要能回到固定快照和固定接口结构上；如果系统每次运行都完全依赖当天行情，论文结果就很难被老师和后续读者复查。因此，快照保存和字段稳定也是功能需求的一部分。",
                ],
            ),
            (
                "3.6 存储设计",
                "热门行业识别与龙头股遴选模型设计",
                [
                    "旧稿中曾把 JSON 快照和研究工作台放在同一条复核线索下说明，这一点仍然有价值。行业子系统当前主要依赖 JSON 文件和浏览器本地状态，优点是轻量、透明、容易随论文材料一起检查；研究工作台等其他模块若使用 SQLite，则更适合沉淀结构化任务和研究结论。",
                    "这种存储选择反映了毕业设计阶段的边界：行业热度主链路追求的是可运行、可回看和可解释，而不是一开始就搭建复杂数据平台。只要热力图历史、偏好配置和财务缓存能够稳定保存，系统就已经具备支持论文样例复核的基础。",
                    "从后续扩展看，若要把行业信号进一步接入回测或研究报告生成流程，可以把固定快照转成更规范的数据表，再与研究工作台的任务记录关联。这样能够形成从行业发现、样例保存到策略验证的更完整闭环，但这已经属于当前课题之后的工作。",
                ],
            ),
            (
                "4.5 快速评分机制设计",
                "系统实现",
                [
                    "快速评分机制还承担了一个容易被忽略的功能：它让系统在字段不完整时仍能给出可解释的候选集合。真实数据环境里，财务字段、估值字段和资金字段不一定同时可用，如果必须等全部字段齐备，行业页就会经常停在加载状态。",
                    "因此，快路径的设计不是为了降低模型要求，而是为了把首屏筛选和深度解释拆成两个阶段。首屏先根据已有快照和交易字段形成候选，详情页再尽量补齐估值、财务和价格序列。这个分层思路与前端交互路径一致，也能减少外部接口波动对用户体验的影响。",
                    "在论文表述上，需要把快路径边界写清楚。快速评分结果适合用于候选预筛和页面响应，不应被解释为完整基本面评价；完整评分则更适合在详情页中展开维度拆解。两条链路共同存在，反而能让系统更贴近真实研究流程。",
                ],
            ),
            (
                "5.6 偏好配置与持续跟踪实现",
                "系统测试与结果分析",
                [
                    "偏好配置的实现使行业研究从一次性浏览变成可持续跟踪。用户可以保存观察行业、阈值和研究视图，后续重新进入页面时继续沿用同一套观察口径，这比每次从默认页面重新筛选更符合研究工作习惯。",
                    "热力图历史快照则解决了另一个问题：行业热度结果会随行情变化，如果没有快照，论文中的表 6.2 很快就无法复现。系统将关键字段写入历史文件后，样例分析可以固定在同一时点，答辩时也能说明数据不是临时手工编写。",
                    "持续跟踪能力还为后续研究留下接口。若将多个快照按时间顺序整理，就可以观察行业热度的形成、延续和衰减过程；再结合龙头股列表，就能进一步分析行业内部代表性公司的变化。这些工作超出当前论文主线，但与本系统的存储设计是衔接的。",
                    "从工程角度看，本地 JSON、浏览器 localStorage 和前端状态并不是彼此孤立的。后端负责保存可复核的历史数据，前端负责保留用户当前研究状态，二者配合后，页面既能支持答辩展示，也能支持日常复盘。",
                ],
            ),
            (
                "6.1 测试环境",
                "6.2 功能测试",
                [
                    "测试环境除了说明软硬件条件，还需要说明测试边界。本文测试的是行业研究原型能否完成数据接入、评分计算、接口返回、页面展示和样例复核，而不是验证一个可直接投入交易的生产系统。",
                    "本地测试采用固定快照与当前代码结合的方式。固定快照保证第六章样例不会随行情漂移，当前代码验证则保证系统仍然能够从适配器、分析器和接口层完成主链路运行。两者结合，比只展示静态截图更可靠。",
                    "前端测试重点放在用户路径：打开行业页、切换热力图窗口、搜索或选择行业、查看排行榜、进入龙头股详情、再检查评分拆解是否存在。这个路径正好对应论文题目中的“热门行业识别”和“龙头股遴选”两个动作。",
                    "后端测试重点放在接口契约和降级行为。即使某些外部字段暂时缺失，接口也应返回结构稳定的结果或清晰的降级说明，而不是让前端出现空白页面。对毕业设计来说，这种稳定性比单次样例得分高低更重要。",
                ],
            ),
            (
                "6.4 系统特征与不足分析",
                "6.5 对毕业设计目标的达成情况",
                [
                    "综合来看，本系统最适合作为行业研究辅助原型，而不是交易执行系统。它能够把多源数据整理成可视化线索，把行业热度和龙头股候选放在同一工作区中解释，但并不负责生成自动交易指令。",
                    "这一定位也决定了本文对不足的评价方式。当前最需要补强的不是页面数量，而是更长周期的信号验证、权重敏感性分析、政策和舆情数据融合，以及行业链上下游关系建模。只有这些工作进一步完成，系统才可能从展示型原型走向更严谨的投资研究工具。",
                    "不过，就任务书要求而言，系统已经覆盖了金融大数据理解、热门行业识别方法、龙头股筛选方法和软件系统实现四个方面。论文通过理论基础、模型说明、系统实现和固定样例验证，把这些目标逐项落到了可检查的章节中。",
                ],
            ),
            (
                "结 论",
                "参考文献",
                [
                    "进一步从应用边界看，本文完成的系统更接近研究辅助工具，而不是自动化投资决策工具。它能够把多源数据、行业热度、个股候选和页面交互组织到一个连续流程中，但最终判断仍需要研究者结合宏观环境、行业基本面和风险偏好进行复核。",
                    "这种定位使系统输出具有较强解释性。用户不仅能看到行业或股票的综合得分，还能继续查看涨跌幅、资金流、换手率、波动代理、估值、盈利和成长等字段。对于毕业设计而言，这种“能追问原因”的结果比单纯给出一个排序更重要。",
                    "本文的另一个收获是把工程可用性纳入了论文主线。外部数据源不稳定、字段不完整、页面响应速度和历史样例复现，都是金融数据系统真正落地时必须面对的问题。通过缓存、回退、快照和偏好保存，系统在这些问题上形成了初步解决方案。",
                    "当然，本文仍然存在明显局限。首先，行业热度和龙头股评分权重主要依据经验设定，尚未经过系统参数优化；其次，样例分析使用的是固定快照，不能替代长周期回测；再次，政策文本、新闻舆情和产业链上下游关系尚未进入核心得分。",
                    "后续研究可以沿三个方向继续推进：一是构建更长时间序列的行业热度信号并加入交易成本假设，检验其对后续收益的解释能力；二是引入文本和产业链数据，增强热点形成原因的解释；三是把行业识别结果与研究工作台、回测模块和报告生成模块联动起来。",
                    "总体而言，本文在本科毕业设计范围内完成了从问题分析、理论梳理、模型设计、系统实现到结果验证的闭环。虽然它还不是生产级投资系统，但已经能够支撑热门行业识别与龙头股遴选这一主题的展示、复核和后续迭代。",
                ],
            ),
        ]
    )

    for heading_text, next_heading_text, paragraphs in expansions:
        append_section_body(doc, heading_text, next_heading_text, paragraphs)


def apply_project_focused_thesis_revision(doc: Document) -> None:
    """Keep the recovered thesis focused on the real project instead of page count."""
    section_revisions = [
        (
            "1.1 研究背景与意义",
            "1.2 国内外研究现状",
            [
                "在证券市场里，真正难做的往往不是盯住一两只股票，而是同时判断热点行业怎么切换、行业内部哪些公司更值得继续跟踪。一个行业持续走强，通常意味着资金、预期和景气度正在往同一方向聚集；而行业里的代表性公司又会因为规模、盈利和市场关注度，被更频繁地拿来比较。于是，如何从海量数据中尽快识别当期热点，并在行业内部继续缩小到更有代表性的股票，就成了一个既有研究价值也有实际用途的问题。",
                "过去这类工作很大程度依赖研究员人工切行情终端、翻行业报告和积累经验。行情节奏慢的时候，这种方式还勉强可行；但市场数据量和更新频率上来之后，单靠人工已经很难同时完成横截面比较、历史跟踪和多维交叉验证。也正因为如此，把大数据处理、量化评分和可视化页面真正接成一条可持续运行的研究链路，比单独写一套纸面分析方法更有工程意义。[1-2]",
                "本文依托的项目并不是单一的课程设计程序，而是一个较为完整的量化研究平台。当前公开仓库对外聚焦策略回测、实时行情与行业热度三个主工作区，同时仍保留部分相关研究能力代码与接口。在这些能力中，行业热度页面与毕业设计任务书的目标最为一致，因此本文不再泛化讨论平台全部模块，而是从现有工程实现中抽取行业热度子系统作为核心研究对象，在真实代码、真实接口与真实历史快照的基础上完成毕业论文撰写。",
                "从研究角度看，行业热点并不是孤立产生的。有效市场、资产定价、动量和行为金融相关研究都说明，资产价格既会受到共同风险和风格暴露影响，也会受到信息扩散、资金再配置和投资者情绪影响。[8-20] 因此，本文关注的行业热度不是简单的短期涨幅榜，而是希望在一个可解释的多维框架中观察行业相对强弱。",
                "从工程角度看，这项工作的价值不只在于“做出一个页面”。一方面，系统把多源金融数据整理成可比较的行业研究结果，能够减少人工筛选成本；另一方面，评分过程保持较强可解释性，便于把工程实现、模型口径和样例结果放在一起说明。更现实的一点是，这个原型已经可以支撑毕业设计展示、运行截图获取和后续迭代，而不是停留在方案层面。",
            ],
        ),
        (
            "1.2 国内外研究现状",
            "1.3 研究内容与技术路线",
            [
                "1.2.1 金融大数据与多源数据处理研究现状",
                "把现有文献和实际系统实现放在一起看，最先暴露出来的问题往往不是算法本身，而是数据怎么接、怎么对齐。行业快照、资金流、估值和财务表来自不同接口时，更新时间、单位和字段名经常对不上；同一个行业在不同平台里也可能对应不同叫法。大数据管理、商业智能和数据密集型系统相关研究都强调，容量、速度、多样性、真实性和可维护性会共同影响金融数据系统能否稳定运行。[1][24-29]",
                "因此，本文借鉴这些研究时，更看重它们对数据组织方式的启发，而不是再额外设计一套抽象工具链。具体到本课题，系统采用 pandas、NumPy、SciPy 和 scikit-learn 等成熟工具，是因为它们已经足够支撑项目里真实存在的字段清洗、横截面比较、聚类辅助和评分计算流程。[6-7][34-37]",
                "1.2.2 行业轮动与热点识别研究现状",
                "在行业轮动与热点识别方面，国内外研究普遍不主张只看单一指标。无论是从资产定价、行业暴露出发，还是从价格动量、资金流向和交易活跃度出发，最终都要回到多维信号综合判断上来。[3][8-16] 这也解释了为什么本文没有把“热门行业”简单理解成短期涨幅排名，而是希望在横截面上同时观察价格、资金、活跃度和风险约束。",
                "行业动量研究进一步说明，行业层面的信息扩散和资金再配置具有一定持续性。Moskowitz 和 Grinblatt 对行业动量的研究表明，行业因素能够解释相当一部分动量收益；Hou 关于行业信息扩散的研究也指出，行业之间可能存在领先和滞后关系。[13-16] 这类研究给本文的启发是：如果系统能稳定保留行业横截面结果，就可以为后续复盘行业轮动提供材料。",
                "1.2.3 龙头股筛选与可解释评价研究现状",
                "把文献里的龙头股识别方法和当前项目放在一起看，有一个很直接的共识：真正能长期代表行业状态的股票，通常不会只靠某一天的涨幅来判断。规模、盈利、成长和估值提供的是相对稳定的基本面线索，成交额、换手率和近期动量则更像市场有没有继续确认这一判断。[4-5][8-12]",
                "同时，行为金融文献提醒研究者不能把短期强势直接等同于基本面改善。投资者情绪、过度反应与反应不足都可能在短期内推高某些行业或个股表现。[17-20] 因此，本系统在行业排序之外继续保留资金流、换手率、波动率代理和市值来源等字段，目的就是避免把一个总分解释得过于绝对。",
                "1.2.4 机器学习方法与工程落地边界",
                "近年机器学习方法已经被广泛用于资产定价和金融时间序列研究，相关文献说明复杂模型能够在大量特征中捕捉非线性结构，但模型效果往往受样本划分、交易成本、数据窥探和样本外检验影响。[19][21-23] 因此，本文没有把系统定位为收益预测器，而是定位为行业研究辅助工具：先把热点横截面和候选龙头股解释清楚，再为后续更严格的回测留下接口。",
                "除行情和财务字段外，政策、公告和新闻文本也可能影响行业热度，金融文本研究已经证明文本信息对市场分析具有参考价值。[42] 不过，本项目当前主链路仍然以结构化行情、资金流、估值和财务字段为主。这样做是一个有意取舍：结构化数据更容易形成稳定接口，评分结果也更容易在页面上解释；文本和舆情数据更适合作为后续扩展方向。",
                "总体来看，现有研究已经为热门行业识别和龙头股筛选提供了较充分的理论基础，但一落到工程实现，就还会遇到两个明显空档：一是多源适配、字段回退、缓存和页面展示这些真正影响可用性的细节经常被简化；二是行业识别和龙头股遴选常被拆开讨论，缺少统一的研究闭环。本文的重点不是再堆一个复杂模型，而是在真实项目基础上，把多源数据获取、行业评分、龙头股筛选和前端展示真正接起来。",
            ],
        ),
        (
            "1.3 研究内容与技术路线",
            "1.4 论文结构安排",
            [
                "本文没有把研究内容拆成几块互不相干的模块，而是尽量按系统真实使用的顺序来梳理。更接近日常研究场景的路径，是先打开行业页，看热力图和排行榜怎样把横截面结果摆出来；接着沿着某个行业继续查看趋势、成分股和龙头股详情；最后再回头解释这些结果是怎样由适配器、缓存和评分器一步步算出来的。",
                "对应到实现链路，前端页面先请求热门行业、热力图和详情相关数据，路由层再把请求分发给 IndustryAnalyzer 和 LeaderStockScorer，分析层继续向 THS-first 主路径和补充数据源取回字段并整理结果。本文之所以按这条顺序展开，是因为这样最容易让论文里的截图、表格和项目代码彼此对上，也更方便答辩时解释页面里每一步是怎么来的。",
                "具体来说，本文的研究内容可以分成四个层次。第一层是数据层，重点解决多源行业数据、成分股数据和财务行情字段如何进入统一结构；第二层是模型层，重点解释行业热度得分和龙头股得分怎样计算；第三层是服务层，说明接口缓存、回退和响应封装如何保证页面可用；第四层是展示层，通过热力图、排行榜和详情弹窗把结果交给用户。",
                "在分析方法上，本文采用文献梳理、系统分析、模型归纳、工程实现说明和样例验证相结合的方式。文献梳理用于确认行业轮动、多因子筛选和金融大数据处理的研究基础；系统分析用于界定需求和架构；模型归纳用于解释评分逻辑；样例验证则使用项目保留的热力图快照和龙头股样本说明系统输出。",
                "在验证方式上，本文并不把单次样例解释为严格投资结论，而是把它作为系统可用性证据。真正的完整投资验证还需要更长历史窗口、交易成本假设、换仓规则和样本外检验，这些内容已经超出当前毕业设计主线，但会在结论和不足部分作为后续方向提出。[19-23]",
            ],
        ),
        (
            "2.1 金融大数据的特征",
            "2.2 多源数据采集与清洗",
            [
                "落到这个课题里，金融大数据最突出的麻烦并不只是“量大”这一个词，而是不同数据一起进来时口径很不整齐。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率不同，结构也不同，可信程度还会跟着来源变化。比如行业快照可以按日更新，财务指标却按季度披露；同一个字段在不同接口里又可能出现名称不一致、单位不同或缺失值写法不同的情况。真要把这些数据送进同一条分析链路，先解决标准化和可用性问题，比直接套模型更重要。[24-29]",
                "对于热门行业识别尤其如此。只盯涨跌幅，很容易把短期情绪波动当成热点；只看资金流，又可能忽略趋势已经转弱的行业。行业研究真正需要的是把价格、资金、活跃度乃至风险约束放在同一横截面里比较，这也是本文后续采用多维评分而不是单指标排序的原因。",
                "金融数据还具有很强的时序性。行业热度不是一个静态标签，而是在不同统计窗口里不断变化的横截面状态。某个行业在 1 日窗口里可能明显升温，在 5 日或 10 日窗口里却已经回落；如果系统不保存窗口信息和快照时间，后续就很难解释同一行业在不同页面中为什么排序不同。",
                "从研究可解释性的角度看，金融大数据还有一个容易被忽视的特点：用户通常不会只接受一个最终分数。研究者会继续追问这个分数来自涨幅、资金、成交活跃度还是波动代理。因而系统必须在输出总分的同时保留关键字段，这也是热力图历史快照同时保存 value、total_score、moneyFlow、turnoverRate 和市值来源的原因。",
            ],
        ),
        (
            "2.2 多源数据采集与清洗",
            "2.3 热门行业识别的理论基础",
            [
                "多源采集在本项目里不是可有可无的补充，而是整条行业分析主链路的一部分。行业子系统按 THS-first 组织数据接口：同花顺负责提供行业目录、行业摘要、资金流和领涨股这类主数据，AKShare 补行业元数据、成分股、估值与财务字段，新浪和腾讯主要承担回退与补缺。这样做并不是为了“多接几个接口”，而是因为单一来源很难同时覆盖行业研究真正需要的全部字段。[24-29][41]",
                "数据拿到之后，马上要处理的不是模型，而是口径统一。项目里会先统一行业名称、字段命名和数值类型，再根据字段特征做缺失值处理、异常值裁剪和重复剔除。例如行业名称需要映射到统一的 industry_name，涨跌幅和资金流要变成可运算的浮点数，个股快照里暂时缺失的财务字段则用中性值参与快速评分。只有经过这一步，不同来源的数据才能进入后面的行业评分和龙头股筛选链路。[6][34-37]",
                "缺失值处理是第三个关键环节。对于行业层面，如果暂时拿不到历史波动率，系统会使用振幅、换手率或涨跌幅代理值；对于个股层面，如果快照场景下拿不到 ROE 或成长字段，快速评分会采用中性值，避免列表直接中断。这样的处理不代表数据完美，而是保证系统在真实环境中能够连续输出。",
                "最后，多源清洗的结果需要面向复核。论文中使用的固定热力图快照和龙头股样本，实际上就是把采集清洗后的关键结果沉淀下来，便于在不同时间重新打开论文或代码时仍能看到同一组样例。这一点对毕业设计尤其重要，因为答辩材料必须稳定，而不能随着当天行情自动漂移。",
            ],
        ),
        (
            "2.3 热门行业识别的理论基础",
            "2.4 龙头股评价的理论基础",
            [
                "行业热点识别可以从三个角度理解。第一是资产定价角度，行业收益可能包含共同风险暴露和风格因子影响；第二是动量角度，强势行业可能在一段时间内继续受到资金关注；第三是行为金融角度，热点也可能包含情绪放大和过度反应。[8-20] 这三个角度共同说明，热门行业不能只靠单一指标判断。",
                "动量研究为本文提供了最直接的理论支撑。Jegadeesh 和 Titman 证明了股票层面的中期动量现象，Moskowitz 和 Grinblatt 则进一步讨论了行业动量。[13][15] 这些研究说明，行业层面的相对强弱具有研究意义，也支持本文按 1 日、5 日、10 日等窗口观察行业状态。",
                "行业信息扩散研究则解释了为什么行业内部和行业之间可能存在领先滞后关系。Hou 的研究表明，行业信息可能通过上下游或相关行业逐步反映到股票价格中。[16] 因此，系统只展示某个行业当天涨幅是不够的，还需要保留趋势、成分股和龙头股详情，让用户继续沿行业线索下钻。",
                "行为金融研究提醒本文在解释热点时保持克制。投资者情绪和市场过度反应可能造成短期强势，但这种强势未必转化为可持续收益。[17-20] 因此，本文在第六章分析样例结果时，只将其解释为样本时点的横截面状态，而不把它写成对未来收益的确定性预测。",
            ],
        ),
        (
            "2.4 龙头股评价的理论基础",
            "2.5 系统关键技术",
            [
                "在这个课题里，“龙头股”并不等于某一天涨得最快的股票。更常见的情况是，短期涨幅靠前的个股未必真的能代表行业，有时只是情绪推动；真正更能代表行业状态的，往往是那些规模、盈利、成长和市场关注度都更稳的公司。传统龙头企业识别也大多沿着这个思路，从市值规模、产业地位、盈利质量、成长能力、估值合理性和市场表现等多个维度综合判断。[4-5][8-12]",
                "从资产定价角度看，个股是否能代表行业，首先要看它是否具有稳定的规模、盈利和风险暴露。市值、估值、ROE、收入增长和利润增长等指标并不是为了给企业贴标签，而是为了判断该公司是否具备行业代表性。若一个股票只有短期涨幅而缺乏基本面支撑，它更适合进入热点观察，而不一定适合作为核心龙头。",
                "从市场确认角度看，热点个股和核心资产往往并不重合。短期资金关注可能迅速推高某些股票的热度，但这种热度未必意味着其长期代表性更强。项目中将榜单拆成 core 和 hot 两类，正是为了保留这两种不同语义。",
                "因此，龙头股评价应当同时考虑“稳定代表性”和“短期市场确认”。前者更依赖市值、估值、盈利和成长，后者更依赖涨跌幅、成交额、换手率和资金承接。本文的评分模型并不宣称能给出唯一正确答案，而是给用户一个透明的排序依据。",
            ],
        ),
        (
            "2.5 系统关键技术",
            "系统需求分析与总体设计",
            [
                "本系统的关键技术可以概括为 Python 数据处理、机器学习辅助分析、Web 接口服务和前端可视化四类。Python 生态中的 pandas、NumPy、SciPy、Matplotlib 和 scikit-learn 已经形成较成熟的数据分析工具链，适合处理结构化行情、横截面标准化、聚类辅助分析和图表生成。[6-7][34-37]",
                "在算法层面，系统主要使用加权评分和 K-Means 聚类。K-Means 本身是一种经典聚类方法，后续研究又提出了 k-means++ 初始化和轮廓系数评估等改进思路。[30-33] 在本文里，聚类不是为了替代行业得分，而是用于观察若干行业是否形成相近热点簇，属于辅助解释工具。",
                "在后端服务层，FastAPI 提供了较轻量的接口组织方式，适合把 Python 分析能力封装成前端可调用的 REST 接口。[38-39] REST 架构思想强调资源、统一接口和无状态交互，这与本项目中热门行业、热力图、趋势、龙头股详情等接口设计具有一致性。",
                "在前端展示层，React 更适合组织状态复杂、交互频繁的单页应用。[40] 行业热度页面需要同时管理热力图、排行榜、详情弹窗、观察行业、时间窗口和偏好配置，如果只用静态报表，很难支撑这样的交互闭环。",
                "在数据源方面，AKShare 为国内金融数据研究提供了便利入口，适合补充行业、行情和财务字段。[41] 但任何开源或第三方数据源都可能遇到接口变动、限流或字段缺失，因此项目没有把单一来源当作绝对可靠前提，而是在适配层保留了多源回退。",
            ],
        ),
        (
            "4.2 热门行业识别模型设计",
            "4.3 行业波动率估计与聚类辅助分析",
            [
                "根据项目中的实际实现，行业热度判断主要围绕四类信号展开：价格动量、资金流强度、交易活跃度和行业波动率。实际代码并没有停留在抽象的“四因子加总”上，而是把每个信号都落到了具体字段和回退口径上。动量主要来自 change_pct 或 weighted_change，资金因子对应 flow_strength，活跃度优先看 avg_volume，拿不到时再退化到 turnover_rate。波动率则优先尝试行业指数历史收益率，拿不到时退回振幅、换手率或涨跌幅代理值。这种写法更符合行业轮动研究里强调多维信号综合判断的思路。[3][13-20]",
                "Sindustry=0.35×Zm+0.35×Zf+0.15×Zv-0.15×Zr",
                "（4.2.1）",
                "该公式中的权重与项目代码中的默认权重保持一致。公式（4.2.1）对应的是横截面原始评分，并不直接等同于前端最终展示值。系统在完成加权之后，会进一步调用分数压缩函数，把结果统一映射到约 20 至 95 的展示区间，以减少样本集中时 0 分和 100 分贴边的问题。因此，排行榜和热力图中的 total_score 更适合用于相对比较，而不应被理解为行业的绝对评价值。",
                "权重设计采用动量和资金各 0.35、活跃度 0.15、波动率约束 0.15 的结构。这样的设计并不意味着该权重在所有市场环境下最优，而是为了在毕业设计阶段保证结果直观、口径稳定、容易解释。若后续要推进到严谨策略研究，可以再增加权重敏感性分析、样本外检验和交易成本假设。[19-23]",
            ],
        ),
        (
            "4.3 行业波动率估计与聚类辅助分析",
            "4.4 龙头股综合评分模型设计",
            [
                "波动率估计在系统中承担的是风险约束功能。一个行业如果短期涨幅很高但波动也很剧烈，综合热度不应被无条件抬高。真实行业指数历史收益率是更直接的波动来源，但在接口不稳定或首屏加载场景下，代理波动率可以先保证页面输出。",
                "代理波动率的设计需要保持克制。振幅、换手率和涨跌幅都只能间接反映波动，并不等同于标准差意义上的历史波动率。因此，本文把它称为代理值，并在第六章说明样例结果是横截面状态而不是严格风险预测。",
                "聚类分析使用 K-Means 和轮廓系数，主要是为了观察行业之间的相似结构。[30-33] 如果若干行业在涨跌幅、资金流和活跃度上接近，聚类结果可以帮助用户判断热点是否成片扩散；如果聚类不稳定，也可以提示当前市场结构比较分散。",
                "在系统展示中，聚类不直接替代热度得分。原因是聚类结果更适合描述相似性，而排行榜需要给出相对顺序。本文将聚类定位为辅助分析，既能保留它的解释价值，也避免把无监督分类结果误用为投资排序。",
            ],
        ),
        (
            "4.4 龙头股综合评分模型设计",
            "4.5 快速评分机制设计",
            [
                "龙头股评分模型的目标，是在已经识别出的热门行业或指定行业内部，从多只成分股里找出更能代表行业状态的股票。结合项目中的实际实现，系统没有只盯短期涨幅，而是同时看市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度六个维度。这样的口径与既有企业价值分析和多因子评价思路基本一致。[4-5][8-12]",
                "其中，市值规模通过对数标准化反映企业体量；估值水平以市盈率（Price-to-Earnings Ratio, PE）处于合理区间得分更高为原则；盈利能力通过净资产收益率（Return on Equity, ROE）体现；成长性通过营业收入同比和利润同比共同体现；价格动量反映短期市场强弱；交易活跃度通过成交额或换手率衡量。设六个维度得分分别为 s1 至 s6，则龙头股综合得分可以表示为：",
                "Sleader=100×(0.20×s1+0.15×s2+0.25×s3+0.20×s4+0.10×s5+0.10×s6)",
                "（4.4.1）",
                "从工程实现看，公式（4.4.1）对应的是完整评分链路下的六维加权总分，其中 s1 至 s6 分别表示市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度的归一化得分。完整评分链路会在统一的 raw_data 结构上聚合估值、财务和行情字段，并按该权重体系生成总分。因此，前端 core 榜单主要反映这一综合评分语义；hot 榜单则更强调短期涨幅与资金承接。",
                "完整评分链路的意义在于解释，而不仅是排序。前端详情页展示维度拆解后，用户可以看到某只股票是因为规模、盈利、成长还是短期动量获得较高得分，这比单一总分更符合研究场景。与复杂机器学习选股模型相比，本文的六维评分明显更简单，但它的优势是透明，适合本科毕业设计阶段的展示和复核。[21-23]",
            ],
        ),
        (
            "4.5 快速评分机制设计",
            "系统实现",
            [
                "快速评分机制的核心是承认页面有不同层级的信息需求。行业页列表只需要先筛出候选集合，详情页才需要展开完整解释。如果所有计算都按详情页标准执行，首屏体验会变差；如果所有计算都按快路径执行，结果解释又会不够充分。",
                "因此，系统将快路径和完整路径并行保留。快路径优先使用已有快照字段，适合 core/hot 榜单的快速展示；完整路径补充估值、财务和价格序列，适合详情弹窗中的维度拆解。两条路径共享 raw_data，可以减少字段含义不一致的问题。",
                "快路径中采用中性值并不是为了掩盖缺失，而是为了避免少数字段缺失导致整个候选列表中断。对于用户来说，先看到一组可解释的候选，再进入详情页补充信息，比等待所有字段完全齐备更符合页面使用习惯。",
                "从模型边界看，快速评分结果不应被解释为完整基本面评价。论文中保留 core 与 hot 两类榜单说明，就是为了提醒用户区分长期代表性和短期市场冲击，避免把不同语义的榜单混为一谈。",
            ],
        ),
        (
            "5.6 偏好配置与持续跟踪实现",
            "系统测试与结果分析",
            [
                "偏好配置和历史快照让前端页面具备持续研究能力。观察列表保存用户关心的行业，回放功能帮助用户比较不同时间点的横截面状态，本地状态则减少重复配置。这些能力虽然不是评分模型本身，却会明显影响系统是否像一个可持续使用的研究原型。",
                "持续跟踪功能的意义在于把一次性结果变成可回看的研究材料。行业热度今天排在前列，并不意味着明天仍然如此；只有保留历史快照，用户才能观察某个行业是短暂冲高、持续走强，还是在资金流和活跃度上出现分化。",
                "从论文角度看，持续跟踪功能也提高了样例分析的可信度。第六章固定引用的快照和龙头股样本，正是通过这种持久化能力保存下来的。如果没有这类记录，论文结果会随着数据源更新不断变化，不利于提交和答辩。",
                "从工程角度看，本地 JSON、浏览器 localStorage 和前端状态并不是彼此孤立的。后端负责保存可复核的历史数据，前端负责保留用户当前研究状态，二者配合后，页面既能支持答辩展示，也能支持日常复盘。",
            ],
        ),
        (
            "6.1 测试环境",
            "6.2 功能测试",
            [
                "测试环境选择本地运行，是因为毕业设计需要同时验证后端分析、前端展示和文档截图。后端负责接口和评分，前端负责交互与渲染，历史快照负责样例复核。三者必须放在同一环境里一起看，才能判断系统是否形成完整闭环。",
                "测试过程中需要区分功能正确性和投资有效性。功能正确性关注接口能否返回、字段是否完整、页面能否交互、快照是否可回看；投资有效性则需要长期回测和交易成本假设。本文的测试重点是前者，后者作为未来研究方向保留。",
                "为了避免样例随行情变化，论文使用固定热力图快照和固定龙头股案例。这样即使未来重新生成论文，表 6.2 和龙头股案例也能保持一致，便于老师检查论文、项目和数据文件之间的对应关系。",
            ],
        ),
        (
            "6.4 系统特征与不足分析",
            "6.5 对毕业设计目标的达成情况",
            [
                "如果只看当前仓库真正保留下来的行业模块，它最有用的地方在于同一套数据入口已经能支撑从热力图到个股详情的连续分析。用户先在页面里看到行业排序，再点进趋势、龙头股和评分拆解时，背后用的仍是同一条适配、评分和缓存链路，因此结果解释不会前后脱节。",
                "系统的第二个特征是结果解释性较强。行业得分可以拆回动量、资金、活跃度和波动代理，个股得分可以拆回规模、估值、盈利、成长、动量和活跃度。即使用户不同意当前权重，也能理解系统为什么给出这样的排序。",
                "第三个特征是工程上保留了现实妥协。接口回退、过期缓存、中性值、快照保存和偏好持久化都不是理论模型里的核心概念，但它们决定了系统能否在真实数据环境中持续运行。",
                "但它的边界也很明确。权重目前还是工程经验优先，部分行业在成分股映射不完整时仍会走降级路径；模块刷新主要依赖 REST、缓存和历史快照，不是持续流式的实时系统；行业识别结果和后续收益之间的量化验证也还可以继续补强。这些问题不会影响毕业设计展示，但确实是后续迭代最该继续打磨的地方。",
                "综合来看，本系统更适合作为行业研究辅助原型，而不是交易执行系统。它能够把多源数据整理成可视化线索，把行业热度和龙头股候选放在同一工作区中解释，但并不负责生成自动交易指令，也不构成投资建议。",
            ],
        ),
        (
            "6.5 对毕业设计目标的达成情况",
            "结 论",
            [
                "因此，本文已把任务书中的理论梳理、模型设计、系统实现和结果验证落到论文与项目证据中，主线完整。",
            ],
        ),
        (
            "结 论",
            "参考文献",
            [
                "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
                "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率代理为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型。行业热度部分采用横截面标准化与加权合成的方法，并进一步压缩到便于展示的相对得分区间；龙头股部分则同时支持完整评分与快照快速评分两条链路，在保证口径一致的前提下兼顾分析完整性与页面响应效率。",
                "在工程实现方面，本文说明了 THS-first 多源数据回退、分析层与路由层缓存、热力图历史快照、财务缓存以及按 profile 偏好持久化的协同工作方式。第六章结合项目当前保留的五日热力图快照、龙头股双榜单样例和测试结果说明，该系统能够识别样本时点的阶段性强势行业，并进一步给出行业内部具有代表性的龙头股候选结果，具有较好的结果解释性与展示性。",
                "本文的不足也比较明确。首先，行业热度和龙头股评分权重主要依据工程经验设定，尚未经过系统参数优化；其次，样例分析使用的是固定快照，不能替代长周期回测；再次，政策文本、新闻舆情和产业链上下游关系尚未进入核心得分。把这些边界写清楚，有助于让论文结论更稳，也有助于后续迭代继续推进。",
                "未来工作可以从以下方向展开：第一，构建更长时间序列的行业热度信号并加入交易成本假设，检验其对后续收益的解释能力；第二，引入政策文本、新闻舆情和产业链关系，增强热点形成原因的解释；第三，继续提高成分股映射完整度并优化核心行业分析链路。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。",
            ],
        ),
    ]

    for heading_text, next_heading_text, paragraphs in section_revisions:
        replace_section_body(doc, heading_text, next_heading_text, paragraphs)

    citation_replacements = {
        "[11-18]": "[8-12]",
        "[19-23]": "[13-16]",
        "[24-28]": "[17-20]",
        "[29-34]": "[21-23]",
        "[31-34]": "[23]",
        "[35-39]": "[42]",
        "[40-45]": "[24-29]",
        "[46-51]": "[30-33]",
        "[52-55]": "[34-37]",
        "[56-57]": "[29][38]",
        "[52-60]": "[24-29][34-41]",
        "[58]": "[39]",
        "[56]": "[38]",
        "[59]": "[40]",
        "[60]": "[41]",
        "[9-10]": "[6-7]",
        "[9]": "[6]",
        "[10]": "[7]",
        "[4][7-8]": "[3][13-16]",
        "[3][5-6]": "[4-5]",
        "[5-6]": "[4-5]",
        "[18][27][29-34]": "[19-23]",
        "[18][27]": "[19]",
        "[24-39]": "[17-23][42]",
        "[19-28]": "[13-20]",
        "[11-28]": "[8-20]",
        "[29-30]": "[21-22]",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in citation_replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            replace_paragraph_text(paragraph, updated)

    seen_long_paragraphs: set[str] = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if "旧稿" in text:
            delete_paragraph(paragraph)
            continue
        if len(text) > 30 and text in seen_long_paragraphs:
            delete_paragraph(paragraph)
            continue
        seen_long_paragraphs.add(text)


def append_project_depth_content(doc: Document) -> None:
    """Add project-grounded depth so the thesis is fuller without artificial padding."""
    additions = [
        (
            "3.4 系统总体架构设计",
            "3.5 数据流程设计",
            [
                "从毕业设计展示角度看，四层架构还有一个好处：它能把论文中的模型说明和真实页面运行过程对应起来。答辩时如果从前端热力图开始演示，页面请求会先进入服务层，再调用分析层和数据层；如果从代码角度说明，则可以反过来从数据适配器、行业分析器、评分器一路讲到前端展示。两种讲法对应的是同一条链路，能够减少论文和项目脱节的问题。",
                "表现层并不是简单负责“好看”。行业热度页面需要同时承载总览、排序、筛选、下钻和回放等动作，因此前端组件必须把热力图、排行榜、详情弹窗和偏好配置组织在同一个状态流里。如果前端只展示静态表格，用户就很难从行业热度继续推进到龙头股遴选。",
                "服务层的价值主要体现在接口契约。前端并不关心底层数据来自同花顺、AKShare、新浪还是腾讯，而是需要结构稳定的字段，例如行业名称、综合得分、涨跌幅、资金流、换手率、龙头股列表和评分拆解。服务层把这些字段封装起来，既降低了前端复杂度，也让后端替换数据源时不至于破坏页面。",
                "分析层是论文模型和代码实现最直接的交汇点。IndustryAnalyzer 负责把行业摘要、资金流、交易活跃度和波动代理转换成行业横截面得分；LeaderStockScorer 负责把个股估值、盈利、成长、动量和活跃度转换成候选龙头股排序。论文第四章的两个公式，正是对这两类分析逻辑的概括。",
                "数据层则承担了系统可靠性的底座作用。真实金融接口不可能永远稳定，字段也不可能永远完整，因此数据层必须处理名称映射、主源失败、补充来源、过期缓存和字段缺失等问题。若没有这一层，模型公式即使写得很清楚，系统也很难在真实环境中连续输出。",
            ],
        ),
        (
            "3.6 存储设计",
            "热门行业识别与龙头股遴选模型设计",
            [
                "从生命周期看，行业热度数据至少包含三种状态：当次接口返回的临时结果、用于页面回放的历史快照、以及支持后续详情补充的缓存字段。临时结果强调新鲜度，历史快照强调可复核性，缓存字段强调响应速度。三者目标不同，因此不能简单用同一种保存策略处理。",
                "热力图历史快照更适合长期保留。它记录的是某个时间窗口下的行业横截面状态，后续即使行情已经变化，论文仍然可以引用同一组快照进行结果分析。这样的设计让第六章的表格和图示具有稳定来源，也让答辩时可以解释“样例不是临时手工编出来的”。",
                "偏好配置则更接近用户状态。观察行业、阈值、回放窗口和当前选择会随着用户习惯变化而变化，适合按 profile 保存。这样处理可以让同一台机器上的不同研究场景互不干扰，也能为后续扩展到研究工作台、报告生成或回测模块留下入口。",
                "财务缓存的设计主要服务于龙头股评分。估值、ROE、营收增长和利润增长不是高频字段，如果每次打开行业页都重新请求，既会拖慢页面，也会增加外部接口失败概率。缓存这些字段可以让完整评分链路更稳定，而快速评分链路又能在缓存缺口较大时先给出候选集合。",
                "需要说明的是，本课题没有把存储层升级成数据库，并不是因为数据库不重要，而是因为当前毕业设计更关注可运行、可解释和可复核。文件化存储在本地原型阶段足够透明，老师或后续读者也更容易顺着文件路径检查样例来源。",
            ],
        ),
        (
            "4.2 热门行业识别模型设计",
            "4.3 行业波动率估计与聚类辅助分析",
            [
                "在具体实现中，行业热度得分并不是一次性由单个字段决定，而是经历字段清洗、缺失处理、横截面归一化、加权合成和分数压缩几个步骤。这样做的原因是行业之间的原始字段尺度差异很大，成交额、资金流和涨跌幅不能直接相加，必须先转成可比较的相对位置。",
                "动量因子的解释相对直观，它反映行业在当前统计窗口内是否走强。资金因子强调市场是否有持续承接，避免只有价格上涨却缺少成交和资金支持的行业被过度高估。活跃度因子用于判断该行业是否有足够交易参与，波动率代理则用于约束短期剧烈波动带来的噪声。",
                "分数压缩是页面展示中的一个重要细节。若直接展示标准化后的原始得分，用户可能看到大量负值或极端值，不利于热力图颜色映射；压缩到相对稳定区间后，虽然牺牲了一部分数学精度，但更适合毕业设计中的可视化展示和人工解释。",
                "模型权重目前仍属于经验型设置。本文选择动量与资金流权重较高，是因为行业热点通常同时需要价格表现和资金确认；活跃度与波动率权重较低，是因为它们更像辅助和约束信号。后续若要增强严谨性，可以基于历史样本进行权重敏感性分析和滚动回测。[19-23]",
                "因此，行业热度模型更适合被理解为研究筛选器，而不是收益预测器。它帮助用户把大量行业先缩小到一批值得继续观察的对象，再结合行业基本面、政策环境和个股结构做进一步判断。",
            ],
        ),
        (
            "4.4 龙头股综合评分模型设计",
            "4.5 快速评分机制设计",
            [
                "龙头股评分中，市值维度体现行业代表性，但它不能单独决定排序。市值大的公司更容易成为行业配置和研究报告中的代表，但过大的市值也可能意味着成长弹性有限。因此，系统只把市值作为六个维度之一，而不是直接按市值排序。",
                "估值维度的作用是约束短期炒作。若一只股票短期涨幅很高但估值已经明显偏离合理区间，系统不应无条件把它放到 core 榜单最前面。相反，hot 榜单可以保留这类短期关注度较高的标的，让用户区分“市场正在追逐”和“长期更具代表性”。",
                "盈利和成长维度更接近基本面判断。ROE、营收同比和利润同比能够反映公司经营质量，但这些字段更新频率低、接口覆盖并不总是完整。因此完整评分链路依赖缓存补充，快速评分链路则在必要时采用中性值，保证页面候选列表不会因个别字段缺失而中断。",
                "动量和活跃度维度用于补充市场确认。即使一家企业基本面较好，如果短期几乎没有成交和关注，也未必适合作为当前热点行业中的候选；反过来，只有动量而缺少基本面支撑的股票，则更适合放在 hot 观察榜中继续跟踪。",
                "这种双榜单设计是本文比较重要的工程取舍。core 和 hot 不是重复列表，而是分别回答两个问题：前者回答“谁更像行业核心代表”，后者回答“谁正在被市场短期推到前台”。把二者同时展示，可以减少用户把短期涨幅等同于长期龙头的误解。",
            ],
        ),
        (
            "5.4 后端接口实现",
            "5.5 前端可视化实现",
            [
                "从接口组织看，行业模块并不是只提供一个综合结果接口，而是围绕前端研究动作拆成多个端点。热门行业接口服务于首屏列表，热力图接口服务于行业横截面展示，趋势接口服务于行业下钻，龙头股接口服务于个股候选，偏好接口则服务于持续跟踪。",
                "这种拆分可以降低单次请求压力。若把所有字段都塞进一个接口，首屏加载会被详情数据拖慢；若每个小字段都拆成单独接口，前端又会出现过多请求。当前设计介于两者之间，把常用汇总结果和详情结果分开，既能保证首屏速度，也能支持用户继续下钻。",
                "接口缓存的设置也与页面使用方式有关。热门行业和热力图结果在短时间内可以复用，因此适合设置分钟级缓存；完整成分股和财务字段成本更高，适合异步构建或长期缓存；偏好配置则更强调写入后的稳定读取。不同接口采用不同缓存策略，比统一缓存时间更合理。",
                "异常返回需要尽量结构化。外部数据源失败时，后端不能只抛出原始异常，否则前端很难判断是完全不可用、部分字段缺失还是可以使用过期缓存。项目中保留降级说明和状态字段，就是为了让页面在异常场景下仍能给用户一个可理解的反馈。",
                "从论文角度看，接口层是模型和页面之间的证据桥。第四章说明模型如何计算，第五章说明这些计算如何通过接口交给前端，第六章再用页面截图和固定快照验证结果。三者连在一起，论文才不会变成只有公式没有系统，或者只有页面没有模型。",
            ],
        ),
        (
            "6.2 功能测试",
            "6.3 热力图快照结果分析",
            [
                "功能测试还需要关注降级场景。金融数据接口在实际运行中可能出现空字段、超时或字段名变化，如果测试只覆盖理想返回，就很难证明系统具备真实可用性。因此，行业评分与接口测试都应检查字段缺失、缓存命中和代理值生效等情况。",
                "前端测试的重点不是逐个按钮截图，而是验证用户路径是否连贯。一个完整路径应当包括进入行业页、查看热力图、切换统计窗口、选择行业、打开龙头股详情、查看评分拆解，并在必要时回到历史快照。只要这条路径稳定，毕业设计展示的可信度就会明显提高。",
                "测试结果也需要保持边界意识。本文的测试能够说明系统主链路可运行、结果可解释、页面可展示，但不能说明该模型已经具备可交易收益。若要把系统推进到投资策略层面，还需要加入更长时间窗口、交易成本、换仓规则和样本外检验。",
                "因此，第六章的测试更像工程验收，而不是策略绩效报告。它回答的是“系统是否按论文描述运行”，而不是“系统是否一定具备稳定收益能力”。这一区分对于金融类毕业设计尤其重要。",
            ],
        ),
        (
            "6.3 热力图快照结果分析",
            "6.4 系统特征与不足分析",
            [
                "固定快照的意义在于把动态市场状态固定下来。若论文每次生成都重新抓取最新行情，表 6.2 的行业排序会不断变化，老师和读者很难复核。使用固定快照后，样例分析就可以围绕同一时点展开，结论也不会随着行情刷新而漂移。",
                "样例中电子、通信和能源金属等方向位于前列，说明系统能够在固定窗口内捕捉阶段性强势板块。但这里仍然需要强调，该结论只对应样本时点。市场环境、政策预期、行业景气度和资金风格变化之后，排序结果都可能发生明显变化。",
                "龙头股双榜单样例也说明，core 与 hot 的结果并不完全相同。core 更接近行业代表性公司筛选，hot 更接近短期冲击信号筛选。两者同时展示，可以帮助用户区分“长期值得研究的公司”和“短期被市场推到前台的公司”。",
                "如果后续要把样例分析推进到投资有效性验证，就需要把多个快照按时间顺序连接起来，构建行业信号形成、持有、换仓和收益评价流程。这部分工作已经超出当前论文篇幅，但系统保存快照的方式为后续研究留下了基础。",
            ],
        ),
    ]

    for heading_text, next_heading_text, paragraphs in additions:
        append_section_body(doc, heading_text, next_heading_text, paragraphs)


def apply_final_project_consistency_cleanup(doc: Document) -> None:
    """Polish project-truth wording after all generated expansions are applied."""
    project_boundary_text = (
        "本文依托的项目并不是单一的课程设计程序，而是一个较为完整的量化研究平台。"
        "当前公开仓库对外保留今日研究、策略回测、实时行情和行业热度四个公开入口，"
        "其中今日研究用于汇总当天研究档案，策略回测和实时行情用于验证与监控，"
        "行业热度页面则与毕业设计任务书的目标最为一致。"
        "因此，本文不再泛化讨论平台全部模块，而是从现有工程实现中抽取行业热度子系统作为核心研究对象，"
        "在真实代码、真实接口与真实历史快照的基础上完成毕业论文撰写。"
    )

    paragraph_replacements = {
        "本文依托的项目并不是单一的课程设计程序": project_boundary_text,
        "从系统实现的角度看，数据提供器是论文中多源数据理论和实际代码之间的连接点。它把文献中常说的数据融合、质量控制和可用性问题，具体落到了名称映射、字段统一、缓存回退和异常处理这些操作上。[24-29][24-29][34-41]":
            "从系统实现的角度看，数据提供器是论文中多源数据理论和实际代码之间的连接点。它把文献中常说的数据融合、质量控制和可用性问题，具体落到了名称映射、字段统一、缓存回退和异常处理这些操作上。[24-29][34-41]",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for start_text, replacement in paragraph_replacements.items():
            if text.startswith(start_text):
                replace_paragraph_text(paragraph, replacement)
                break

    section_text_replacements = {
        "研究工作台的任务记录": "今日研究或研究档案的任务记录",
        "研究工作台、报告生成或回测模块": "今日研究、报告生成或回测模块",
        "研究工作台、回测模块和报告生成模块": "今日研究/研究档案、回测模块和报告生成模块",
        "样例中电子、通信和能源金属等方向位于前列": "样例中能源金属、元件、军工电子和小金属等方向位于前列",
        "系统是否一定能赚钱": "系统是否一定具备稳定收益能力",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        for old_text, new_text in section_text_replacements.items():
            updated = updated.replace(old_text, new_text)
        updated = re.sub(r"(\[\d+(?:-\d+)?\])\1+", r"\1", updated)
        if updated != original:
            replace_paragraph_text(paragraph, updated)

    redundant_starts = (
        "从表 6.2 可以看出，行业综合得分和单一涨跌幅之间并不完全一致。",
        "从表 6.2 的样例可以进一步看到，行业综合得分和单一涨跌幅之间并不完全一致。",
        "表 6.1 中的“满足预期”并不是生产级质量保证，",
    )
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith(redundant_starts):
            delete_paragraph(paragraph)

    seen_long_paragraphs: set[str] = set()
    for paragraph in list(doc.paragraphs):
        stripped = paragraph.text.strip()
        if len(stripped) <= 80:
            continue
        if stripped in seen_long_paragraphs:
            delete_paragraph(paragraph)
            continue
        seen_long_paragraphs.add(stripped)

    replace_section_body(
        doc,
        "3.1 设计目标",
        "3.2 功能需求分析",
        [
            "系统设计目标首先是可用性。行业热度页必须在外部数据源短时波动时仍尽量给出结果，否则研究者每次打开页面都可能被接口状态打断。因此，系统需要缓存、回退和降级路径，而不是只在理想网络环境下运行。",
            "第二个目标是可解释性。行业综合得分和龙头股综合得分都应当能拆回具体字段，用户至少要知道得分来自价格强弱、资金流、活跃度、波动率，还是来自市值、估值、盈利和成长。这个目标与行为金融和机器学习资产定价研究中的可解释性要求是一致的。[17-23][42]",
            "第三个目标是可复核性。论文中的表格和截图不能依赖某一秒钟的临时行情，因此系统需要保存热力图历史快照和龙头股样本，使结果章节能够固定引用同一组数据。这个目标也是后续答辩和归档最实际的需求。",
        ],
    )


def insert_project_depth_tables(doc: Document) -> None:
    architecture_anchor = find_paragraph_by_text(doc, "3.5 数据流程设计")
    arch_intro = insert_paragraph_before_element(architecture_anchor._element, architecture_anchor._parent)
    arch_intro.add_run("结合系统总体架构，行业子系统各层职责和证据来源可归纳为表 3.2。该表用于说明论文中的模型、接口和页面证据如何对应到真实项目实现。")
    arch_caption = insert_paragraph_before_element(architecture_anchor._element, architecture_anchor._parent)
    arch_caption.add_run("表 3.2 行业子系统核心模块职责与证据来源")
    arch_caption.paragraph_format.keep_with_next = True
    arch_table = insert_table_after(doc, arch_caption, rows=5, cols=4)
    arch_rows = [
        ("层次", "核心职责", "项目对应内容", "论文证据"),
        ("表现层", "组织行业研究交互", "React 行业仪表盘、热力图、排行榜、详情弹窗", "图 5.1-图 5.3"),
        ("服务层", "封装稳定接口契约", "FastAPI 行业接口、缓存、降级响应", "第 5.4 节与表 5.1"),
        ("分析层", "完成评分与解释", "IndustryAnalyzer、LeaderStockScorer", "公式（4.2.1）和（4.4.1）"),
        ("数据层", "多源适配与快照保存", "THS-first、AKShare/Sina/腾讯补齐、JSON 快照", "第 3.6 节与表 6.2"),
    ]
    for r_idx, row in enumerate(arch_rows):
        for c_idx, value in enumerate(row):
            arch_table.cell(r_idx, c_idx).text = value

    score_anchor = find_paragraph_by_text(doc, "4.3 行业波动率估计与聚类辅助分析")
    score_intro = insert_paragraph_before_element(score_anchor._element, score_anchor._parent)
    score_intro.add_run("在具体计算中，行业热度模型需要同时处理字段缺失和数据源波动。表 4.1 汇总了核心指标的主要字段、回退口径和解释作用。")
    score_caption = insert_paragraph_before_element(score_anchor._element, score_anchor._parent)
    score_caption.add_run("表 4.1 行业热度指标口径与回退策略")
    score_caption.paragraph_format.keep_with_next = True
    score_table = insert_table_after(doc, score_caption, rows=5, cols=4)
    score_rows = [
        ("指标", "主要字段", "回退口径", "解释作用"),
        ("价格动量", "change_pct / weighted_change", "统计窗口内可用涨跌幅", "判断行业是否走强"),
        ("资金强度", "flow_strength / moneyFlow", "缺失时保留中性或降级说明", "判断资金是否承接"),
        ("交易活跃度", "avg_volume", "turnover_rate / 成交额代理", "判断关注度和交易参与"),
        ("波动率约束", "industry_volatility", "振幅、换手率或涨跌幅代理", "约束高波动噪声"),
    ]
    for r_idx, row in enumerate(score_rows):
        for c_idx, value in enumerate(row):
            score_table.cell(r_idx, c_idx).text = value

    leader_anchor = find_paragraph_by_text(doc, "4.5 快速评分机制设计")
    leader_intro = insert_paragraph_before_element(leader_anchor._element, leader_anchor._parent)
    leader_intro.add_run("龙头股综合评分并不把单一维度绝对化，而是将规模、估值、盈利、成长、动量和活跃度放在同一框架内比较。各维度含义和解释边界见表 4.2。")
    leader_caption = insert_paragraph_before_element(leader_anchor._element, leader_anchor._parent)
    leader_caption.add_run("表 4.2 龙头股评分维度与解释边界")
    leader_caption.paragraph_format.keep_with_next = True
    leader_table = insert_table_after(doc, leader_caption, rows=7, cols=4)
    leader_rows = [
        ("维度", "主要字段", "筛选含义", "解释边界"),
        ("市值规模", "market_cap", "衡量行业代表性和市场关注度", "市值大不必然代表成长性更好"),
        ("估值水平", "PE / PB 等估值字段", "避免高估值短期热门股无条件靠前", "估值区间需结合行业差异解释"),
        ("盈利能力", "ROE / 利润率", "反映经营质量与核心资产属性", "低频字段依赖缓存补充"),
        ("成长性", "营收同比 / 利润同比", "衡量公司未来扩张能力", "短期缺失时快速路径采用中性值"),
        ("价格动量", "涨跌幅 / 价格趋势", "捕捉近期市场确认", "不能单独解释为长期龙头"),
        ("交易活跃度", "成交额 / 换手率", "判断是否有足够市场参与", "高活跃也可能来自短期情绪"),
    ]
    for r_idx, row in enumerate(leader_rows):
        for c_idx, value in enumerate(row):
            leader_table.cell(r_idx, c_idx).text = value

    api_anchor = find_paragraph_by_text(doc, "5.5 前端可视化实现")
    api_intro = insert_paragraph_before_element(api_anchor._element, api_anchor._parent)
    api_intro.add_run("后端接口与前端页面之间的对应关系如表 5.1 所示。该表体现了模型结果不是孤立输出，而是被组织为排行榜、热力图、趋势面板、详情弹窗和偏好配置等具体研究动作。")
    api_caption = insert_paragraph_before_element(api_anchor._element, api_anchor._parent)
    api_caption.add_run("表 5.1 行业接口与前端用途对应关系")
    api_caption.paragraph_format.keep_with_next = True
    api_table = insert_table_after(doc, api_caption, rows=6, cols=4)
    api_rows = [
        ("接口类型", "核心返回字段", "前端用途", "可靠性设计"),
        ("热门行业", "industry_name、total_score、change_pct", "首屏排行榜", "短期缓存与降级说明"),
        ("热力图", "value、moneyFlow、turnoverRate", "行业横截面可视化", "固定快照与历史回放"),
        ("行业趋势", "窗口收益、波动代理、样本点", "行业下钻分析", "缺失字段代理处理"),
        ("龙头股列表/详情", "core/hot、score、raw_data", "候选股筛选与解释", "快路径与完整路径并行"),
        ("偏好配置", "profile、watchlist、thresholds", "持续跟踪", "JSON 持久化与本地状态"),
    ]
    for r_idx, row in enumerate(api_rows):
        for c_idx, value in enumerate(row):
            api_table.cell(r_idx, c_idx).text = value

    risk_anchor = find_paragraph_by_text(doc, "6.5 对毕业设计目标的达成情况")
    risk_intro = insert_paragraph_before_element(risk_anchor._element, risk_anchor._parent)
    risk_intro.add_run("在结果解释中，系统不足需要和测试结论一起呈现。表 6.3 从权重、数据、文本信息、投资验证和部署形态五个方面列出了当前边界及后续改进方向。")
    risk_intro.paragraph_format.page_break_before = True
    risk_intro.paragraph_format.keep_with_next = True
    risk_caption = insert_paragraph_before_element(risk_anchor._element, risk_anchor._parent)
    risk_caption.add_run("表 6.3 系统不足与后续改进方向")
    risk_caption.paragraph_format.keep_with_next = True
    risk_table = insert_table_after(doc, risk_caption, rows=6, cols=4)
    risk_rows = [
        ("不足类型", "当前表现", "影响范围", "后续改进方向"),
        ("权重经验性", "行业和个股权重主要依据工程经验", "影响排序稳健性", "引入敏感性分析和滚动回测"),
        ("数据覆盖", "部分行业成分股和财务字段可能降级", "影响详情完整度", "补充数据源并记录缺失来源"),
        ("文本信息不足", "政策、新闻和公告暂未进入核心得分", "影响热点原因解释", "加入文本情绪和事件标签"),
        ("投资验证不足", "固定快照不能替代长期收益检验", "限制策略有效性判断", "构建换仓规则和样本外检验"),
        ("部署形态有限", "当前以本地研究原型为主", "限制多人协作和长期服务", "后续迁移数据库和任务调度"),
    ]
    for r_idx, row in enumerate(risk_rows):
        for c_idx, value in enumerate(row):
            risk_table.cell(r_idx, c_idx).text = value


def ensure_blank_page_before_declaration(doc: Document) -> None:
    if not doc.tables:
        return

    declaration_table = doc.tables[0]

    next_sibling = declaration_table._element.getnext()
    if next_sibling is not None and next_sibling.tag == qn("w:p") and 'w:type="page"' in next_sibling.xml:
        next_sibling.getparent().remove(next_sibling)

    previous_sibling = declaration_table._element.getprevious()
    if previous_sibling is not None and previous_sibling.tag == qn("w:p") and 'w:type="page"' in previous_sibling.xml:
        return

    break_para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    break_para.append(run)
    declaration_table._element.addprevious(break_para)


def set_core_properties(doc: Document) -> None:
    doc.core_properties.author = STUDENT_INFO["student_name"]
    doc.core_properties.title = THESIS_TITLE
    doc.core_properties.subject = THESIS_TITLE
    doc.core_properties.keywords = "金融大数据, 热门行业识别, 龙头股遴选, 多源数据, 行业热度分析系统"


def run_checked(command: list[str]) -> None:
    subprocess.run(command, check=True)


def get_available_cjk_font_sources() -> list[tuple[Path, int]]:
    sources = [(path, index) for path, index in CJK_FONT_SOURCES if path.exists()]
    if not sources:
        raise RuntimeError("No available CJK font found for thesis front-page rendering.")
    return sources


def load_cjk_font(size: int, font_sources: list[tuple[Path, int]] | None = None):
    sources = font_sources or get_available_cjk_font_sources()
    for path, index in sources:
        try:
            return ImageFont.truetype(str(path), size=size, index=index)
        except OSError:
            continue
    raise RuntimeError("Failed to load a usable CJK font source for thesis rendering.")


def fit_cjk_font(
    draw: ImageDraw.ImageDraw,
    text: str,
    max_width: int,
    start_size: int,
    min_size: int = 14,
    font_sources: list[tuple[Path, int]] | None = None,
):
    for size in range(start_size, min_size - 1, -1):
        font = load_cjk_font(size, font_sources=font_sources)
        bbox = draw.textbbox((0, 0), text, font=font)
        if bbox[2] - bbox[0] <= max_width:
            return font
    return load_cjk_font(min_size, font_sources=font_sources)


def make_box(left: float, top: float, right: float, bottom: float) -> dict[str, float]:
    return {
        "left": float(left),
        "top": float(top),
        "right": float(right),
        "bottom": float(bottom),
    }


def merge_boxes(boxes: list[dict[str, float]]) -> dict[str, float]:
    return make_box(
        min(box["left"] for box in boxes),
        min(box["top"] for box in boxes),
        max(box["right"] for box in boxes),
        max(box["bottom"] for box in boxes),
    )


def scale_box(box: dict[str, float], scale_x: float, scale_y: float) -> dict[str, float]:
    return make_box(
        box["left"] * scale_x,
        box["top"] * scale_y,
        box["right"] * scale_x,
        box["bottom"] * scale_y,
    )


def normalize_pdf_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def ensure_template_pdf_xml() -> Path:
    TMP_TEMPLATE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    existing = sorted(TMP_TEMPLATE_PDF_DIR.glob("template_layout*.xml"))
    if existing:
        return existing[0]

    run_checked([
        "pdftohtml",
        "-xml",
        "-f",
        "1",
        "-l",
        "2",
        "-nomerge",
        "-noroundcoord",
        str(TEMPLATE_PDF_PATH),
        str(TEMPLATE_XML_PREFIX),
    ])
    generated = sorted(TMP_TEMPLATE_PDF_DIR.glob("template_layout*.xml"))
    if not generated:
        raise RuntimeError("Failed to extract template XML layout for thesis front pages.")
    return generated[0]


def ensure_template_pdf_pages() -> tuple[Path, Path]:
    TMP_TEMPLATE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    page_1 = TMP_TEMPLATE_PDF_DIR / "template_page-01.png"
    page_2 = TMP_TEMPLATE_PDF_DIR / "template_page-02.png"
    if page_1.exists() and page_2.exists():
        return page_1, page_2

    prefix = TMP_TEMPLATE_PDF_DIR / "template_page"
    run_checked([
        "pdftoppm",
        "-f",
        "1",
        "-l",
        "2",
        "-png",
        str(TEMPLATE_PDF_PATH),
        str(prefix),
    ])
    if not page_1.exists() or not page_2.exists():
        raise RuntimeError("Failed to render official template PDF pages.")
    return page_1, page_2


def detect_cover_lines_in_image(image: Image.Image) -> dict[str, dict[str, int]]:
    grayscale = image.convert("L")
    width, height = grayscale.size
    search_left = int(width * 0.34)
    search_right = int(width * 0.80)
    search_top = int(height * 0.43)
    search_bottom = int(height * 0.75)
    min_length = int(width * 0.20)
    candidates: list[tuple[int, int, int]] = []

    for y in range(search_top, search_bottom):
        best_start = None
        best_end = None
        run_start = None
        for x in range(search_left, search_right + 1):
            is_dark = grayscale.getpixel((x, y)) < 120
            if is_dark:
                if run_start is None:
                    run_start = x
            elif run_start is not None:
                if best_start is None or x - run_start > best_end - best_start:
                    best_start, best_end = run_start, x - 1
                run_start = None
        if run_start is not None:
            if best_start is None or search_right - run_start > best_end - best_start:
                best_start, best_end = run_start, search_right
        if best_start is not None and best_end is not None and best_end - best_start >= min_length:
            candidates.append((y, best_start, best_end))

    groups: list[list[tuple[int, int, int]]] = []
    for row in candidates:
        if not groups or row[0] - groups[-1][-1][0] > 3:
            groups.append([row])
        else:
            groups[-1].append(row)

    normalized_groups = []
    for group in groups:
        line = {
            "y_line": int(round(sum(item[0] for item in group) / len(group))),
            "x_start": int(round(sum(item[1] for item in group) / len(group))),
            "x_end": int(round(sum(item[2] for item in group) / len(group))),
        }
        if line["x_end"] - line["x_start"] >= min_length:
            normalized_groups.append(line)

    if len(normalized_groups) < len(COVER_FIELD_ORDER):
        raise RuntimeError("Failed to detect all cover underline anchors from the official template.")
    if len(normalized_groups) > len(COVER_FIELD_ORDER):
        normalized_groups = sorted(
            normalized_groups,
            key=lambda line: line["x_end"] - line["x_start"],
            reverse=True,
        )[:len(COVER_FIELD_ORDER)]
        normalized_groups.sort(key=lambda line: line["y_line"])

    return {
        key: line
        for key, line in zip(COVER_FIELD_ORDER, normalized_groups[: len(COVER_FIELD_ORDER)])
    }


def detect_cover_lines(template_page: Path) -> dict[str, dict[str, int]]:
    try:
        with Image.open(template_page) as image:
            return detect_cover_lines_in_image(image)
    except Exception:
        return deepcopy(COVER_FALLBACK_LINES)


def extract_template_text_boxes(
    template_xml_path: Path,
    template_page: Path,
    page_number: int,
    labels: list[str],
) -> dict[str, dict[str, float]]:
    tree = ET.parse(template_xml_path)
    root = tree.getroot()
    page = next(
        (node for node in root.findall("page") if node.get("number") == str(page_number)),
        None,
    )
    if page is None:
        raise RuntimeError(f"Template XML page {page_number} not found.")

    with Image.open(template_page) as page_image:
        width, height = page_image.size
    scale_x = width / float(page.get("width"))
    scale_y = height / float(page.get("height"))

    nodes = []
    for element in page.findall("text"):
        raw_text = "".join(element.itertext())
        normalized = normalize_pdf_text(raw_text)
        if not normalized:
            continue
        left = float(element.get("left"))
        top = float(element.get("top"))
        box = make_box(
            left * scale_x,
            top * scale_y,
            (left + float(element.get("width"))) * scale_x,
            (top + float(element.get("height"))) * scale_y,
        )
        nodes.append({
            "text": raw_text,
            "normalized": normalized,
            "left_raw": left,
            "top_raw": top,
            "box": box,
        })

    lines: list[list[dict[str, object]]] = []
    for node in sorted(nodes, key=lambda item: (item["top_raw"], item["left_raw"])):
        if not lines or abs(node["top_raw"] - lines[-1][0]["top_raw"]) > 4:
            lines.append([node])
        else:
            lines[-1].append(node)

    boxes: dict[str, dict[str, float]] = {}
    for label in labels:
        normalized_label = normalize_pdf_text(label)
        for line in lines:
            sorted_line = sorted(line, key=lambda item: item["left_raw"])
            for start in range(len(sorted_line)):
                matched_nodes = []
                candidate = ""
                for end in range(start, len(sorted_line)):
                    piece = sorted_line[end]["normalized"]
                    next_candidate = candidate + piece
                    if not normalized_label.startswith(next_candidate):
                        break
                    candidate = next_candidate
                    matched_nodes.append(sorted_line[end]["box"])
                    if candidate == normalized_label:
                        boxes[label] = merge_boxes(matched_nodes)
                        break
                if label in boxes:
                    break
            if label in boxes:
                break

    for label in labels:
        if label not in boxes and label in DECLARATION_FALLBACK_LABEL_BOXES:
            boxes[label] = deepcopy(DECLARATION_FALLBACK_LABEL_BOXES[label])
    missing = [label for label in labels if label not in boxes]
    if missing:
        raise RuntimeError(f"Failed to extract template label boxes: {', '.join(missing)}")
    return boxes


def draw_text_above_line(
    draw: ImageDraw.ImageDraw,
    text: str,
    line_box: dict[str, int],
    spec: dict[str, object],
    font_sources: list[tuple[Path, int]],
) -> dict[str, float]:
    line_width = line_box["x_end"] - line_box["x_start"]
    padding = int(spec["padding_x"])
    available_width = min(int(spec["max_width"]), line_width - 2 * padding)
    font = fit_cjk_font(
        draw,
        text,
        available_width,
        int(spec["font_size"]),
        font_sources=font_sources,
    )
    bbox = draw.textbbox((0, 0), text, font=font)
    text_width = bbox[2] - bbox[0]
    if spec["align"] == "center":
        desired_left = line_box["x_start"] + (line_width - text_width) / 2
    else:
        desired_left = line_box["x_start"] + padding
    desired_bottom = line_box["y_line"] - int(spec["gap_above_line"])
    draw_x = desired_left - bbox[0]
    draw_y = desired_bottom - bbox[3]
    draw.text((draw_x, draw_y), text, font=font, fill=(18, 18, 18, 255))
    return make_box(
        draw_x + bbox[0],
        draw_y + bbox[1],
        draw_x + bbox[2],
        draw_y + bbox[3],
    )


def draw_text_after_label(
    draw: ImageDraw.ImageDraw,
    text: str,
    label_box: dict[str, float],
    spec: dict[str, object],
    font_sources: list[tuple[Path, int]],
) -> dict[str, float]:
    font = fit_cjk_font(
        draw,
        text,
        int(spec["max_width"]),
        int(spec["font_size"]),
        font_sources=font_sources,
    )
    bbox = draw.textbbox((0, 0), text, font=font)
    desired_left = label_box["right"] + int(spec["padding_x"])
    label_center_y = (label_box["top"] + label_box["bottom"]) / 2
    desired_center_y = label_center_y
    draw_x = desired_left - bbox[0]
    draw_y = desired_center_y - (bbox[1] + bbox[3]) / 2
    draw.text((draw_x, draw_y), text, font=font, fill=(18, 18, 18, 255))
    return make_box(
        draw_x + bbox[0],
        draw_y + bbox[1],
        draw_x + bbox[2],
        draw_y + bbox[3],
    )


def render_cover_page(template_page: Path, output_path: Path) -> dict[str, object]:
    line_boxes = detect_cover_lines(template_page)
    font_sources = get_available_cjk_font_sources()
    image = Image.open(template_page).convert("RGBA")
    draw = ImageDraw.Draw(image)
    placements = {}
    for key in COVER_FIELD_ORDER:
        placements[key] = draw_text_above_line(
            draw,
            str(COVER_FIELD_SPECS[key]["text"]),
            line_boxes[key],
            COVER_FIELD_SPECS[key],
            font_sources,
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)
    return {
        "placements": placements,
        "anchors": line_boxes,
        "image_size": image.size,
    }


def render_declaration_page(template_page: Path, output_path: Path) -> dict[str, object]:
    font_sources = get_available_cjk_font_sources()
    template_xml = ensure_template_pdf_xml()
    label_boxes = extract_template_text_boxes(
        template_xml,
        template_page,
        2,
        [spec["anchor_label"] for spec in DECLARATION_FIELD_SPECS.values()],
    )
    image = Image.open(template_page).convert("RGBA")
    draw = ImageDraw.Draw(image)
    placements = {}
    for key, spec in DECLARATION_FIELD_SPECS.items():
        placements[key] = draw_text_after_label(
            draw,
            str(spec["text"]),
            label_boxes[str(spec["anchor_label"])],
            spec,
            font_sources,
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)
    return {
        "placements": placements,
        "label_boxes": label_boxes,
        "image_size": image.size,
    }


def prepare_declaration_page_asset(template_page: Path, output_path: Path) -> tuple[Path, dict[str, object] | None]:
    signed_page = find_first_existing_path(*SIGNED_DECLARATION_PAGE_CANDIDATES)
    if signed_page is None:
        return output_path, render_declaration_page(template_page, output_path)

    if signed_page.suffix.lower() == ".pdf":
        output_path.parent.mkdir(parents=True, exist_ok=True)
        override_prefix = output_path.with_name(f"{output_path.stem}_signed_override")
        run_checked([
            "pdftoppm",
            "-f",
            "1",
            "-l",
            "1",
            "-singlefile",
            "-png",
            str(signed_page),
            str(override_prefix),
        ])
        return override_prefix.with_suffix(".png"), None

    return signed_page, None


def build_front_pdf(front_cover_png: Path, declaration_png: Path, output_pdf: Path) -> None:
    first = Image.open(front_cover_png).convert("RGB")
    second = Image.open(declaration_png).convert("RGB")
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    first.save(output_pdf, "PDF", save_all=True, append_images=[second], resolution=150.0)


def export_docx_pdf(doc_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    run_checked([
        "soffice",
        "-env:UserInstallation=file:///tmp/lo_profile_shu_thesis",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(doc_path),
    ])
    pdf_path = output_dir / f"{doc_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError(f"Failed to export PDF from DOCX: {pdf_path}")
    return pdf_path


def merge_submission_pdf(front_pdf: Path, body_pdf: Path, output_pdf: Path) -> None:
    front_reader = PdfReader(str(front_pdf))
    body_reader = PdfReader(str(body_pdf))
    writer = PdfWriter()

    for page in front_reader.pages:
        writer.add_page(page)
    for page in body_reader.pages[BODY_PDF_ABSTRACT_PAGE_INDEX:]:
        writer.add_page(page)

    with output_pdf.open("wb") as fp:
        writer.write(fp)


def render_pdf_preview(input_pdf: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    run_checked([
        "pdftoppm",
        "-f",
        "1",
        "-l",
        "40",
        "-png",
        str(input_pdf),
        str(prefix),
    ])


def get_rendered_preview_page(output_dir: Path, page_number: int) -> Path:
    candidates = [
        output_dir / f"page-{page_number:02d}.png",
        output_dir / f"page-{page_number}.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    matches = sorted(output_dir.glob(f"page-*{page_number:02d}.png"))
    if matches:
        return matches[0]
    raise RuntimeError(f"Rendered preview page not found: {page_number}")


def validate_front_page_alignment(
    rendered_page: Path,
    placements: dict[str, dict[str, float]],
    source_size: tuple[int, int],
) -> None:
    with Image.open(rendered_page) as image:
        rendered_lines = detect_cover_lines_in_image(image)
        scale_x = image.size[0] / source_size[0]
        scale_y = image.size[1] / source_size[1]

    for key in COVER_FIELD_ORDER:
        placement = scale_box(placements[key], scale_x, scale_y)
        line_box = rendered_lines[key]
        gap = line_box["y_line"] - placement["bottom"]
        if placement["bottom"] >= line_box["y_line"]:
            raise RuntimeError(f"Cover field '{key}' overlaps the underline in the final PDF.")
        if gap < 1.5 or gap > 10:
            raise RuntimeError(
                f"Cover field '{key}' is not aligned above the underline: gap={gap:.2f}px."
            )


def validate_declaration_page_alignment(
    rendered_page: Path,
    placements: dict[str, dict[str, float]],
    label_boxes: dict[str, dict[str, float]],
    source_size: tuple[int, int],
) -> None:
    with Image.open(rendered_page) as image:
        scale_x = image.size[0] / source_size[0]
        scale_y = image.size[1] / source_size[1]

    for key, spec in DECLARATION_FIELD_SPECS.items():
        placement = scale_box(placements[key], scale_x, scale_y)
        label_box = scale_box(label_boxes[str(spec["anchor_label"])], scale_x, scale_y)
        center_delta = abs(
            ((placement["top"] + placement["bottom"]) / 2)
            - ((label_box["top"] + label_box["bottom"]) / 2)
        )
        if center_delta > 2:
            raise RuntimeError(
                f"Declaration field '{key}' is vertically misaligned: delta={center_delta:.2f}px."
            )
        if placement["left"] <= label_box["right"]:
            raise RuntimeError(f"Declaration field '{key}' overlaps its label region.")


def write_toc_entry(paragraph, style_name: str, title: str, page: str) -> None:
    clear_paragraph(paragraph)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
    run = paragraph.add_run(f"{title}\t{page}")
    set_run_text_font(run, "宋体", Pt(12), bold=False)


def normalize_chapter_headings(doc: Document) -> None:
    chapter_pattern = re.compile(r"^\d+\s+")
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Heading 1":
            continue
        text = paragraph.text.strip()
        if not text:
            paragraph.style = "Normal"
            continue
        if text in {"结 论", "参考文献", "致 谢"}:
            format_unnumbered_heading(paragraph)
            continue
        if chapter_pattern.match(text):
            clear_paragraph(paragraph)
            run = paragraph.add_run(chapter_pattern.sub("", text))
            run.font.name = "Times New Roman"
            set_east_asia_font(run, "黑体")


def rebuild_toc(doc: Document, toc_pages: dict[str, str] | None = None) -> None:
    title_para = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "目  录":
            title_para = paragraph
            break
    if title_para is None:
        raise RuntimeError("Failed to locate TOC title paragraph.")

    toc_end = None
    search_started = False
    for paragraph in doc.paragraphs:
        if paragraph._element is title_para._element:
            search_started = True
            continue
        if search_started and "w:sectPr" in paragraph._element.xml:
            toc_end = paragraph
            break
    if toc_end is None:
        raise RuntimeError("Failed to locate TOC section end.")

    remove_section_break(title_para)
    current = title_para._element.getnext()
    toc_end_element = toc_end._element
    while current is not None and current is not toc_end_element:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    format_toc_title(title_para)
    resolved_pages = toc_pages or {title: "" for _, title in TOC_BLUEPRINT}
    for style_name, title in TOC_BLUEPRINT:
        paragraph = toc_end.insert_paragraph_before()
        write_toc_entry(paragraph, style_name, title, resolved_pages.get(title, ""))


def find_table_containing_text(doc: Document, needle: str):
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if needle in text:
            return table
    return None


def ensure_table_caption_immediately_before(
    doc: Document,
    table_needle: str,
    caption_text: str,
    intro_text: str | None = None,
    page_break_before: bool = False,
) -> None:
    """Keep table captions adjacent to their tables, as required by the SHU template."""
    table = find_table_containing_text(doc, table_needle)
    if table is None:
        return

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == caption_text:
            delete_paragraph(paragraph)

    caption = insert_paragraph_before_element(table._tbl, table._parent)
    caption.add_run(caption_text)
    caption.paragraph_format.page_break_before = page_break_before
    caption.paragraph_format.keep_with_next = True

    if intro_text and not any(paragraph.text.strip() == intro_text for paragraph in doc.paragraphs):
        intro = insert_paragraph_before_element(caption._element, caption._parent)
        intro.add_run(intro_text)


def normalize_numbered_table_captions(doc: Document) -> None:
    """Fix inherited template captions that can drift after section body replacement."""
    ensure_table_caption_immediately_before(
        doc,
        "数据采集模块",
        "表 3.1 系统功能需求表",
        "根据上述需求分析，行业热度子系统的核心功能可整理为表 3.1。该表把用户操作、系统处理和输出结果对应起来，便于后续设计和测试逐项核对。",
    )
    ensure_table_caption_immediately_before(doc, "行业评分计算与回退", "表 6.1 主要功能测试结果")
    ensure_table_caption_immediately_before(
        doc,
        "能源金属",
        "表 6.2 五日窗口行业热力图快照样例",
        page_break_before=True,
    )


def update_body_content(doc: Document) -> None:
    remove_generated_padding_paragraphs(doc)

    replacements = {
        "本文依托的项目并不是单一的课程设计程序":
            "本文依托的项目并不是单一的课程设计程序，而是一个较为完整的量化研究平台。当前公开仓库对外保留今日研究、策略回测、实时行情和行业热度四个公开入口，其中今日研究用于汇总当天研究档案，策略回测和实时行情用于验证与监控，行业热度页面则与毕业设计任务书的目标最为一致。因此，本文不再泛化讨论平台全部模块，而是从现有工程实现中抽取行业热度子系统作为核心研究对象，在真实代码、真实接口与真实历史快照的基础上完成毕业论文撰写。",
        "随着金融市场数据规模持续增长":
            "A 股市场的行业轮动很快，真正做研究时，光靠人工盯盘、翻资讯和手工比对往往跟不上节奏。毕业设计任务书希望解决的正是这一类问题，所以本文没有另外搭一套脱离项目背景的新系统，而是把现有量化研究平台里已经在运行的行业热度子系统拿出来，按真实代码、真实接口和真实历史快照重新梳理一遍，把数据接入、行业评分、个股筛选和可视化展示这几条主链路讲清楚。",
        "A 股行情更新很快":
            "A 股市场的行业轮动很快，真正做研究时，光靠人工盯盘、翻资讯和手工比对往往跟不上节奏。毕业设计任务书希望解决的正是这一类问题，所以本文没有另外搭一套脱离项目背景的新系统，而是把现有量化研究平台里已经在运行的行业热度子系统拿出来，按真实代码、真实接口和真实历史快照重新梳理一遍，把数据接入、行业评分、个股筛选和可视化展示这几条主链路讲清楚。",
        "系统采用前后端分离架构。前端基于 React 与 Ant Design 构建行业热力图":
            "从工程实现看，这个子系统没有追求大而全，而是围绕日常行业研究最常用的界面来组织。前端主要保留热力图、排行榜、趋势面板和龙头股详情这些入口；后端用 FastAPI 把热门行业、成分股、热力图、趋势和龙头股相关接口串起来；评分和整理逻辑则放在 Python 分析层中完成。数据侧也不是单一来源，而是先由同花顺优先（Tonghuashun-first，以下简称 THS-first）适配器拿行业目录和摘要，再按字段缺口补接 AKShare、新浪财经和腾讯接口。",
        "具体实现上，前端保留了行业研究最常用的几个界面":
            "从工程实现看，这个子系统没有追求大而全，而是围绕日常行业研究最常用的界面来组织。前端主要保留热力图、排行榜、趋势面板和龙头股详情这些入口；后端用 FastAPI 把热门行业、成分股、热力图、趋势和龙头股相关接口串起来；评分和整理逻辑则放在 Python 分析层中完成。数据侧也不是单一来源，而是先由同花顺优先（Tonghuashun-first，以下简称 THS-first）适配器拿行业目录和摘要，再按字段缺口补接 AKShare、新浪财经和腾讯接口。",
        "系统采用前后端分离架构。前端围绕行业热力图、行业排行榜、行业趋势与龙头股详情构建研究界面":
            "从工程实现看，这个子系统没有追求大而全，而是围绕日常行业研究最常用的界面来组织。前端主要保留热力图、排行榜、趋势面板和龙头股详情这些入口；后端用 FastAPI 把热门行业、成分股、热力图、趋势和龙头股相关接口串起来；评分和整理逻辑则放在 Python 分析层中完成。数据侧也不是单一来源，而是先由同花顺优先（Tonghuashun-first，以下简称 THS-first）适配器拿行业目录和摘要，再按字段缺口补接 AKShare、新浪财经和腾讯接口。",
        "在模型设计方面，本文依据项目中已经实现的 IndustryAnalyzer 与 LeaderStockScorer 两个核心模块":
            "模型部分没有额外引入难以解释的黑箱方法，而是直接沿用项目里已经落地的两个分析类。IndustryAnalyzer 负责判断行业热度，看的主要是价格变化、资金承接、交易活跃度和波动率代理；LeaderStockScorer 负责在行业内部比较个股，把规模、估值、盈利、成长、动量和活跃度六个维度合成综合分，并保留快照场景下的快速评分链路。这样的写法虽然不炫技，但更容易和现有工程实现一一对应。",
        "模型部分没有采用难以解释的黑箱方法":
            "模型部分没有额外引入难以解释的黑箱方法，而是直接沿用项目里已经落地的两个分析类。IndustryAnalyzer 负责判断行业热度，看的主要是价格变化、资金承接、交易活跃度和波动率代理；LeaderStockScorer 负责在行业内部比较个股，把规模、估值、盈利、成长、动量和活跃度六个维度合成综合分，并保留快照场景下的快速评分链路。这样的写法虽然不炫技，但更容易和现有工程实现一一对应。",
        "结合项目保存的行业热力图历史快照数据":
            "论文写作时，本文直接使用项目保留下来的热力图历史快照和本地运行结果做样例验证。按五日窗口观察，电子化学品、通信设备、半导体、消费电子和能源金属等阶段性强势行业能够被系统稳定排到前列，行业内部也能继续给出代表性龙头候选。换句话说，这套原型系统至少已经能够支撑“先识别热点，再往行业内部找代表股票”这一最核心的研究流程。",
        "结合项目保存的行业热力图历史快照和系统运行结果进行验证":
            "论文写作时，本文直接使用项目保留下来的热力图历史快照和本地运行结果做样例验证。按五日窗口观察，电子化学品、通信设备、半导体、消费电子和能源金属等阶段性强势行业能够被系统稳定排到前列，行业内部也能继续给出代表性龙头候选。换句话说，这套原型系统至少已经能够支撑“先识别热点，再往行业内部找代表股票”这一最核心的研究流程。",
        "关键词：金融大数据；行业热度识别；龙头股遴选；多源数据；量化研究平台":
            "关键词：金融大数据；热门行业识别；龙头股遴选；多源数据；量化研究平台",
        "With the rapid growth of financial market data":
            "Against the backdrop of rapidly expanding financial market datasets, approaches that rely primarily on manual experience are no longer sufficient for identifying hot industries and representative leader stocks in a timely and interpretable manner. Focusing on the graduation-project topic of hot-industry identification and leader-stock selection, this thesis takes the industry-heat subsystem of an existing quantitative research platform as its research object and develops an A-share-oriented prototype for financial big-data analysis that integrates data acquisition, industry scoring, stock screening, and visual presentation.",
        "The system adopts a front end and back end separated architecture. The front end is implemented with React and Ant Design":
            "The system follows a front-end/back-end separated architecture. The front end, implemented with React and Ant Design, provides core research views including industry heatmaps, ranking boards, trend panels, and leader-stock detail pages. The back end is built on FastAPI, while the analytical layer is implemented in Python. At the data layer, a Tonghuashun-first (THS-first) adapter organizes industry catalog data, summary snapshots, and money-flow information, with AKShare, Sina Finance, and Tencent valuation endpoints serving as multi-source complements and fallback sources.",
        "According to the actual implementation of the project, the thesis summarizes two core models.":
            "Based on the actual implementation of the project, the thesis summarizes two core models. The industry-heat model evaluates momentum, money-flow strength, trading activity, and a volatility proxy, while the leader-stock model evaluates market capitalization, valuation, profitability, growth, price momentum, and trading activity. A lightweight fast-scoring path is additionally retained for scenarios in which complete financial data are temporarily unavailable.",
        "Historical heatmap snapshots stored in the project show that the system can iden":
            "Historical heatmap snapshots stored in the project show that the system can identify stage-specific strong industries and further screen representative leader-stock candidates within those industries. The results indicate that the proposed prototype maintains reasonable engineering usability, interpretability, and visual readability, and can provide a practical reference for continuous industry tracking.",
        "Keywords: financial big data; industry heat; leader stock selection; quantitative platform":
            "Keywords: financial big data; hot-industry identification; leading stock selection; multi-source data; quantitative research platform",
        "在热门行业识别与龙头股遴选过程中，不同数据字段的量纲差异很大。例如，行业资金流通常以亿元计量，涨跌幅以百分比计量，市值则以十亿或万亿规模计量。如果直接对原始数据进行加权，容易导致尺度较大的变量主导结果。因此，系统在分析阶段首先对指标做数值清洗和标准化处理。":
            "在热门行业识别与龙头股遴选过程中，不同数据字段的量纲差异很大。例如，行业资金流通常以亿元计量，涨跌幅以百分比计量，市值则以十亿或万亿规模计量。如果直接对原始数据进行加权，容易导致尺度较大的变量主导结果。因此，系统在分析阶段首先对指标做数值清洗和标准化处理。[9-10]",
        "行业评分部分主要使用标准化方法将动量、资金流、活跃度和波动率等因子转换到统一尺度后再进行加权；个股评分部分则结合线性归一化、对数标准化和区间裁剪等方式处理市值、成交额、ROE、营收同比和利润同比等指标。对于缺失数据，系统采用中性值或回退值，保证评分流程不因少量字段缺失而中断。":
            "行业评分部分主要使用标准化方法将动量、资金流、活跃度和波动率等因子转换到统一尺度后再进行加权；个股评分部分则结合线性归一化、对数标准化和区间裁剪等方式处理市值、成交额、净资产收益率（Return on Equity, ROE）、营业收入同比和利润同比等指标。对于缺失数据，系统采用中性值或回退值，保证评分流程不因少量字段缺失而中断。[9-10]",
        "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心，但在工程实现上并非简单依赖抽象的四因子加总。动量因子主要来自 change_pct 或 weighted_change，用于刻画行业近期相对强弱；资金因子主要来自 flow_strength，用于衡量行业吸引资金的能力；活跃度因子优先采用 avg_volume，缺失时退化为 turnover_rate；波动率因子则优先使用行业指数历史收益率计算的真实波动率，不足时再回退到振幅、换手率或涨跌幅代理值。":
            "根据项目中的实际实现，行业热度判断主要围绕四类信号展开：价格动量、资金流强度、交易活跃度和行业波动率。不过代码里并不是把“四因子模型”当成一句口号来写，而是把每个字段都落到了可追踪的数据口径上。动量主要来自 change_pct 或 weighted_change，资金因子对应 flow_strength，活跃度优先看 avg_volume，拿不到时再退化到 turnover_rate。波动率这一项最能体现工程取舍：热力图和排行榜的冷启动阶段先复用振幅、换手率或涨跌幅代理值，等用户缩小到指定行业集合或进入详情分析后，再尽量补充行业指数历史收益率波动率。即使历史序列暂时拿不到，系统也会沿用代理值保持输出连续性。这种做法与行业轮动研究里强调多维信号综合判断的思路是一致的。[4][7-8]",
        "在实际运行中，并不是所有行业都能稳定获得完整的历史波动率数据。为提高系统鲁棒性，项目在行业分析引擎中设计了多级回退机制：优先使用真实行业指数历史收益率计算波动率；若历史数据不可用，则依次回退到振幅代理、换手率代理和涨跌幅代理。该设计保证了波动率因子在多数场景下都能被估算出来。":
            "在实际运行中，并不是所有行业都能稳定获得完整的历史波动率数据。为兼顾首屏响应速度与结果稳定性，项目在行业分析引擎中设计了分层波动率策略：页面冷启动时先用振幅代理、换手率代理或涨跌幅代理把热力图和排行榜跑起来；当用户缩小到指定行业集合，或者继续打开详情分析时，系统再补拉真实行业指数历史收益率去修正波动率。如果历史数据仍然拿不到，就继续沿用代理值。这样做的重点不是追求某一个字段绝对完美，而是让结果在大多数场景下都能持续输出，同时保持口径大体稳定。[9-10]",
        "此外，系统还使用 K-Means 聚类对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。":
            "此外，系统还使用 K-Means 聚类方法对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。[10]",
        "龙头股评分模型的目标是在已识别出的热门行业或指定行业内部，从多只成分股中筛选更具代表性的股票。结合项目中的实际实现，系统从六个维度构建综合评分模型：市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度。":
            "龙头股评分模型的目标，是在已经识别出的热门行业或指定行业内部，从多只成分股里找出更能代表行业状态的股票。结合项目中的实际实现，系统没有只盯短期涨幅，而是同时看市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度六个维度。这样的口径与既有龙头企业识别和财务可视化研究中常见的分析框架基本一致。[3][5-6]",
        "其中，市值规模通过对数标准化反映企业体量；估值水平以 PE 处于合理区间得分更高为原则；盈利能力通过 ROE 体现；成长性通过营收同比和利润同比共同体现；价格动量反映短期市场强弱；交易活跃度通过成交额或换手率衡量。设六个维度得分分别为 s1 至 s6，则龙头股综合得分可以表示为：":
            "其中，市值规模通过对数标准化反映企业体量；估值水平以市盈率（Price-to-Earnings Ratio, PE）处于合理区间得分更高为原则；盈利能力通过净资产收益率（Return on Equity, ROE）体现；成长性通过营业收入同比和利润同比共同体现；价格动量反映短期市场强弱；交易活跃度通过成交额或换手率衡量。设六个维度得分分别为 s1 至 s6，则龙头股综合得分可以表示为：[3][5-6]",
        "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。":
            "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。该接口组织方式与前文围绕多源金融数据处理和 Python 分析工具链构建研究系统的技术基础相衔接。[1-2][9-10]",
        "技术路线方面，系统采用前后端分离架构。前端使用 React 实现可视化交互":
            "技术路线并不是先写一套抽象模型再去找实现载体，而是从现有行业子系统的真实链路反推出来。前端用 React 承载热力图、榜单和详情弹窗，后端以 FastAPI 暴露行业分析接口，分析侧通过 pandas、NumPy 与 scikit-learn 完成字段清洗、标准化、评分和辅助聚类，数据侧则由 THS-first 适配器先接同花顺，再用 AKShare、新浪和腾讯补齐缺失字段。围绕这条实际链路，本文再展开数据获取、预处理、指标计算、结果展示和样例验证。",
        "在投资研究实践中，市值较大的企业更容易体现行业代表性":
            "落到行业内个股比较时，系统并没有把“龙头股”简单理解成短期涨幅最大的股票，而是把规模、盈利、成长、估值和市场交易信号放在一起看。原因也比较直接：市值更大的公司通常更能代表行业，净资产收益率和营收、利润增速能够反映经营质量，而成交额、换手率和近期涨跌幅则补足市场确认信息。把这些维度合在一起，得到的筛选结果通常比只看某一个指标更稳定。",
        "多源数据采集是本系统实现的基础。结合项目代码，系统通过数据提供器工厂统一管理不同来源的数据接口。":
            "多源数据采集是整个行业子系统能跑起来的前提。结合项目代码来看，这部分并不是把几个接口简单拼在一起，而是按 THS-first 的思路来组织。对于 A 股行业分析任务，同花顺相关接口优先提供行业目录、行业摘要、资金流与领涨股等主数据；AKShare 负责补齐行业元数据、成分股、估值与财务信息；新浪财经和腾讯接口则更多承担回退和字段补缺职责。这样处理的目的很直接，就是尽量减少单一数据源波动对整条分析链路的影响。",
        "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心。":
            "根据项目中的实际实现，行业热度判断主要围绕四类信号展开：价格动量、资金流强度、交易活跃度和行业波动率。实际代码并没有停留在抽象的“四因子加总”上，而是把每个信号都落到了具体字段和回退口径上。动量主要来自 change_pct 或 weighted_change，资金因子对应 flow_strength，活跃度优先看 avg_volume，拿不到时再退化到 turnover_rate。波动率则优先尝试行业指数历史收益率，拿不到时退回振幅、换手率或涨跌幅代理值。这种写法更符合行业轮动研究里强调多维信号综合判断的思路。[4][7-8]",
        "设行业动量、资金流、活跃度和波动率标准化后的结果分别为 Zm、Zf、Zv 和 Zr":
            "设行业动量、资金流、活跃度和波动率标准化后的结果分别为 Zm、Zf、Zv 和 Zr，则行业横截面原始评分可表示为：",
        "系统采用前后端分离架构，前端使用 React 与 Ant Design 构建行业热力图、行业排行榜、龙头股面板、行业趋势图和轮动对比图等交互页面，后端基于 FastAPI 提供数据接口与分析服务，核心分析逻辑由 Python 完成。数据层以 AKShare 获取 A 股行业分类、行业资金流、行业指数、个股估值和财务数据，同时引入新浪行业适配层作为回退机制，以提升系统运行中的稳定性和可用性。":
            "系统采用前后端分离架构，前端使用 React 与 Ant Design 构建行业热力图、行业排行榜、行业趋势图和龙头股详情等核心页面，后端基于 FastAPI 提供行业分析接口与偏好配置服务，核心分析逻辑由 Python 完成。行业数据层并非简单依赖 AKShare 单一数据源，而是由 THS-first 适配器主导：同花顺行业目录与行业摘要提供主要的行业热度底座，AKShare 负责行业元数据、成分股、财务与历史行情补齐，新浪财经和腾讯估值接口承担回退与字段补充，从而提升行业模块的连通性与抗故障能力。",
        "结合项目真实实现，本文的研究内容主要包括四个方面：第一，对现有量化研究平台中的行业热度子系统进行模块化梳理，明确其在整个平台中的功能定位；第二，围绕行业热度识别建立以动量、资金流、活跃度和波动率为核心的行业评分机制；第三，围绕行业内核心公司筛选建立以规模、估值、盈利、成长、动量和活跃度为核心的龙头股评分机制；第四，对系统的前后端实现、缓存机制、历史快照保存和运行结果进行分析。":
            "结合项目真实实现，本文的研究内容是沿着行业子系统已经跑通的链路来展开的。首先需要把 THS-first 数据适配器、FastAPI 路由层和前端行业页面之间的关系交代清楚；随后再说明 IndustryAnalyzer 如何完成行业热度横截面评分、热力图生成和趋势统计，以及 LeaderStockScorer 如何组织完整评分与快照快速评分两条龙头股筛选链路；最后再回到缓存、历史快照和偏好持久化这些工程细节，对系统的整体闭环做总结。",
        "本系统的关键技术包括以下几个方面。第一，后端采用 FastAPI 构建 REST 接口，具备较好的开发效率和接口组织能力。第二，前端采用 React 与 Ant Design 构建仪表盘式交互页面，便于完成热力图、表格、趋势图和弹窗详情的联动。第三，数据分析部分主要使用 pandas、NumPy 和 scikit-learn，实现数据处理、标准化与聚类。第四，数据持久化部分采用 JSON 与 SQLite 相结合的方式，以较低成本实现历史快照和用户偏好的保存。[9-10]":
            "本系统的关键技术包括以下几个方面。第一，后端采用 FastAPI 构建 REST 接口，行业模块通过参数化路由输出热门行业、热力图、成分股、龙头股、趋势和偏好配置等核心数据。第二，前端采用 React 与 Ant Design 实现仪表盘式页面，并结合 Recharts 等组件完成热力图、榜单、趋势和详情弹窗联动。第三，分析部分使用 pandas、NumPy 和 scikit-learn 完成字段清洗、横截面标准化、综合评分和辅助聚类分析。第四，行业子系统的持久化主要依赖 JSON 文件与浏览器 localStorage；其他研究模块即便采用结构化存储方式，也不属于本文行业分析主链路的核心内容。[9-10]",
        "系统总体上采用四层结构：表现层、服务层、分析层和数据层。表现层由 IndustryDashboard 及相关前端组件构成，负责用户交互和图表展示；服务层由 FastAPI 行业接口组成，负责请求接收、参数校验、缓存与响应封装；分析层由 IndustryAnalyzer 和 LeaderStockScorer 组成，负责评分计算和结果整理；数据层则通过数据提供器工厂统一管理 AKShare、新浪等数据源。":
            "系统总体上采用四层结构：表现层、服务层、分析层和数据层。表现层以前端行业主页面、排行榜组件以及行业详情和龙头股详情弹窗为核心，回放和偏好配置组件作为辅助交互存在；服务层由 FastAPI 行业接口组成，负责请求接收、参数校验、缓存与响应封装；分析层由 IndustryAnalyzer 行业分析器和 LeaderStockScorer 龙头股评分器组成，负责评分计算和结果整理；数据层则以 THS-first 适配器为核心，统一协调同花顺、AKShare、新浪和腾讯等数据源。",
        "系统数据流程可概括为：前端页面发起请求后，接口层根据请求类型调用行业分析器或龙头股评分器；分析器从主数据源获取行业与个股原始数据，必要时启用回退数据源；随后对数据进行清洗、字段统一和异常兜底；在标准化后计算行业或个股综合得分；最后将结果封装为统一响应结构返回前端，并在本地缓存和快照文件中记录必要信息。":
            "系统数据流程可概括为：前端页面通过 REST 请求调用行业接口；路由层优先检查端点缓存，随后根据请求类型调用 IndustryAnalyzer 或 LeaderStockScorer；分析层优先从 THS-first 适配器获取行业摘要、行业资金流和成分股快照，必要时再回退到 AKShare、新浪或腾讯接口，并在必要时返回过期缓存以维持结果连续性；在字段统一、异常处理和标准化之后，系统计算行业或个股综合得分，将结果封装为统一响应结构返回前端，并把必要的历史快照与偏好配置持久化到本地文件或浏览器缓存。",
        "项目并未采用复杂的大型数据库方案，而是根据实际需求选择轻量化存储方式。行业偏好配置和热力图历史快照保存为 JSON 文件，便于快速读写和迁移；研究工作台任务则使用 SQLite 持久化，以支持更稳定的结构化存储。":
            "项目并未为行业子系统引入独立数据库，而是根据真实实现选择轻量化存储方式。后端将热力图历史快照写入 data/industry/heatmap_history.json，将行业偏好按 profile 写入 data/industry_preferences/<profile>.json，并把龙头股财务缓存保存在 cache/financial_cache.json；前端还将热力图回放快照和当前选中状态保存到浏览器 localStorage 中。其他研究模块若采用结构化存储方式，也不影响本文行业子系统以文件缓存和浏览器状态为主的实现路径。",
        "系统的数据层以数据提供器工厂为统一入口，对不同数据源进行注册、管理与切换。对于行业分析任务，AKShare 是主数据源，承担行业分类、行业资金流、行业指数、成分股、估值与财务信息的获取工作；新浪行业适配层则负责在部分接口失败或字段不完整时进行补充。":
            "系统的数据层以 THS-first 适配器作为行业分析主入口。该适配器优先利用同花顺行业目录与行业一览表获取行业名称、涨跌幅、资金流与领涨股等主数据，再由 AKShare 补齐行业元数据、成分股、财务与历史 K 线，新浪财经接口承担行业列表、成分股与行情补充，腾讯财经接口则补单股估值核心字段。",
        "这种设计使系统形成了主数据源优先、备用数据源兜底的运行模式。对于面向真实市场数据的应用场景而言，这种容错机制比单纯依赖单一接口更具工程可用性。":
            "这种设计使系统形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行模式。对于面向真实市场数据的行业研究场景而言，多源适配比依赖单一接口更具工程可用性，也更能覆盖行业目录、资金流、成分股、估值和历史走势等不同类型的数据需求。",
        "行业分析模块的核心类为 IndustryAnalyzer。该模块负责完成行业资金流分析、行业动量计算、行业波动率估计、热门行业排序、热力图数据生成、行业趋势分析和聚类分析等任务。为提高效率，模块内部设置了 30 分钟缓存，避免在短时间内重复计算相同结果。":
            "行业分析模块的核心类为 IndustryAnalyzer。该模块负责行业资金流分析、行业动量计算、波动率估计、热门行业排序、热力图数据生成和行业趋势统计等任务，聚类与轮动分析则作为辅助能力存在。其实现中优先采用快速路径：直接复用行业资金流与横截面摘要结果，避免逐行业拉取全部成分股；在此基础上，再对 change_pct 或 weighted_change、flow_strength、avg_volume 或 turnover_rate 以及 industry_volatility 做标准化，并按默认权重计算原始分数。得到原始分数后，系统还会进一步压缩到约 20 至 95 的展示区间，以减少样本集中时 0 分和 100 分贴边的问题。为提高效率，分析器内部设置了 30 分钟缓存，避免在短时间内重复计算相同结果。",
        "在接口调用链路中，热门行业接口首先获取行业分类列表，再根据行业资金流和动量数据生成横截面数据表，随后计算综合得分并输出行业排行榜。热力图接口则在此基础上增加尺寸映射和颜色映射所需字段。趋势接口会进一步拉取行业指数序列并汇总行业内涨跌分布。":
            "在接口调用链路中，热门行业接口首先获取行业分类列表，再根据行业资金流、涨跌幅、换手率和波动代理生成横截面数据表，随后计算综合得分并输出行业排行榜。热力图接口则在此基础上增加尺寸映射、颜色映射和市值来源标签所需字段；当真实市值缺失时，系统还会退化到成交额、资金流或成分股数量代理。趋势接口会进一步拉取行业指数序列、汇总成分股涨跌分布，并返回覆盖率与降级说明字段。",
        "龙头股评分模块的核心类为 LeaderStockScorer。该模块负责对单只股票评分、在行业内部进行排名、生成龙头股列表以及输出评分拆解与原始数据。模块同时支持基于完整财务数据的完整评分和基于快照数据的轻量快速评分。":
            "龙头股评分模块的核心类为 LeaderStockScorer，即系统中的龙头股评分器。该模块负责对单只股票评分、在行业内部进行排名、生成龙头股列表以及输出评分拆解与原始数据。模块同时支持两条链路：其一是基于估值与财务数据的完整评分，其二是直接复用行业成分股快照的轻量快速评分。前者用于龙头股详情和深度研究，后者用于行业列表页中的快速筛选；在结果组织上，系统进一步把输出划分为 core（核心资产）与 hot（热点先锋）两个榜单，用于区分综合质量筛选和短期热点确认。",
        "在实现层面，模块通过缓存机制对财务数据进行暂存，避免在批量评分过程中重复请求；同时使用统一的原始数据结构将估值、财务和行情字段映射为评分输入，使不同评分路径共享同一套维度计算逻辑。":
            "在实现层面，模块会将财务数据以 24 小时缓存方式保存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、市盈率（Price-to-Earnings Ratio, PE）、净资产收益率（Return on Equity, ROE）、营收同比、利润同比、涨跌幅和成交额等数据。完整评分链路按六维权重体系生成综合得分；快速评分链路则复用相同字段结构，在无法获取 ROE 与增长指标时采用中性分处理，以兼顾结果稳定性与页面响应效率。",
        "在实现层面，模块会将财务数据以 24 小时 TTL 缓存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、PE、ROE、营收同比、利润同比、涨跌幅和成交额等数据，使完整评分与快速评分共享同一套六维评分逻辑。对于快速评分暂时无法获得的 ROE 与增长指标，系统采用中性分处理，以换取页面响应效率。":
            "在实现层面，模块会将财务数据以 24 小时缓存方式保存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、市盈率（Price-to-Earnings Ratio, PE）、净资产收益率（Return on Equity, ROE）、营收同比、利润同比、涨跌幅和成交额等数据。完整评分链路按六维权重体系生成综合得分；快速评分链路则复用相同字段结构，在无法获取 ROE 与增长指标时采用中性分处理，以兼顾结果稳定性与页面响应效率。",
        "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表接口、行业成分股接口、热力图接口、热力图历史接口、行业趋势接口、行业聚类接口、行业轮动接口、龙头股列表接口、龙头股详情接口以及偏好配置接口。":
            "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。该接口组织方式与前文围绕多源金融数据处理和 Python 分析工具链构建研究系统的技术基础相衔接。[1-2][9-10]",
        "为了提升接口稳定性，接口层额外设置了结果缓存，并在必要时使用过期缓存作为兜底返回。这一设计能够在外部数据源短暂波动时尽量保持页面可用。":
            "为了提升接口稳定性，路由层额外维护了 3 分钟端点缓存、30 分钟 parity 缓存以及完整版成分股异步构建机制，并在必要时返回过期缓存作为补充保障。热力图历史数据还设置了最多 48 条记录和 2MB 文件大小上限，从而在外部数据源短暂波动时尽量保持页面可用。",
        "行业热度页面由 IndustryDashboard 及其子组件实现。页面集成热力图、行业排行榜、行业趋势面板、龙头股面板、轮动分析图、预警面板、观察列表和评分雷达图弹窗等功能。用户可以通过点击行业名称进入详细面板，进一步查看该行业的趋势、成分股和龙头股。":
            "行业热度页面由 IndustryDashboard 及其子组件实现。页面以热力图、行业排行榜、行业趋势面板和龙头股面板为核心，同时保留回放与偏好配置等辅助功能。用户可以通过点击行业名称进入详细面板，进一步查看该行业的趋势、成分股和龙头股。",
        "与传统静态报表不同，本系统更强调交互式研究体验。热力图支持不同颜色维度和尺寸维度切换，排行榜支持多字段排序，详情页支持评分拆解和价格走势查看。这种设计更加符合行业分析工作的实际使用场景。":
            "与传统静态报表不同，本系统更强调交互式研究体验。热力图支持不同颜色维度、尺寸维度和市值来源筛选，排行榜支持回看窗口、波动过滤和排序口径切换，详情页支持评分拆解和价格走势查看；历史快照回放与偏好配置则作为研究辅助能力存在。这种设计更加符合行业分析工作的实际使用场景。",
        "项目中已经实现行业偏好持久化功能，支持保存观察行业、已保存视图和告警阈值。默认告警条件覆盖行业得分、涨跌幅、资金净流入和波动率等指标。该功能说明系统并非一次性分析工具，而是支持持续跟踪的研究原型。":
            "项目中已经实现行业偏好持久化功能，用于保存观察行业和研究配置等辅助信息。后端以 profile 维度将观察行业、研究视图和阈值配置写入 JSON 文件，前端则把热力图回放快照和当前选中状态保存在浏览器本地存储中。上述能力主要服务于持续跟踪与复盘分析，而不是改变行业热度识别与龙头股遴选这一核心研究流程。",
        "此外，系统还会记录热力图历史快照，为后续对比分析提供基础。研究工作台模块则为进一步沉淀分析结果和任务状态提供了可能。":
            "此外，系统后端还会持续记录热力图历史快照，为后续对比分析和前端热力图回放提供基础；这些快照不仅保存涨跌幅，还保留 total_score、资金流、换手率、市值来源和估值字段，便于后续复盘不同时间点的行业状态。",
        "本文的系统测试基于项目当前本地环境展开。系统采用 Python 作为后端分析语言，FastAPI 作为接口框架，React 作为前端框架，数据源主要为 AKShare，并结合项目中保存的历史快照进行结果验证。由于本课题属于毕业设计性质，测试重点放在功能闭环、数据连通性、结果可解释性和页面展示效果，而非面向生产环境的极限压测。":
            "本文的系统测试基于项目当前本地环境展开。系统采用 Python 作为后端分析语言，FastAPI 作为接口框架，React 作为前端框架，行业主数据经由 THS-first 适配器提供，并在需要时调用 AKShare、新浪与腾讯接口补齐字段，同时结合项目中保存的历史快照进行结果验证。除样例快照分析外，项目还保留了行业评分逻辑、适配器映射、偏好服务与前端行业页面的测试脚本，用于验证评分逻辑、回退逻辑和页面交互闭环。由于本课题属于毕业设计性质，测试重点放在功能闭环、数据连通性、结果可解释性和页面展示效果，而非面向生产环境的极限压测。",
        "为了增强论文结果分析的真实性，本文直接选取项目中已经保存的热力图历史快照作为样本。根据 2026 年 4 月 11 日 22:59:20 保存的五日窗口行业热力图数据，系统输出的前十个行业结果如表 6.2 所示。":
            "为了增强论文结果分析的真实性，本文直接选取项目中已经保存的热力图历史快照作为样本。选择 2026 年 4 月 11 日 22:59:20 保存的五日窗口快照，是因为该样本由系统自动记录、字段较为完整，并同时覆盖综合得分、涨跌幅、资金流和换手率等核心指标，便于展示系统在固定统计窗口下的横截面输出。表 6.2 展示了该样本时点按热力图默认顺序截取的前十个行业横截面字段，包括 total_score、5 日涨跌幅、资金流和换手率等信息。需要说明的是，这里给出的结果用于样例分析，并不构成完整回测结论。",
        "从表 6.2 可以看出，通信设备、半导体、电子化学品、消费电子、光学光电子和自动化设备等科技成长风格行业在该时间窗口内表现较强，说明系统能够较好识别当期市场热点方向。值得注意的是，综合得分并不完全等同于涨跌幅。例如，能源金属行业在五日涨跌幅较高的同时，还表现出正向资金流，因此综合得分达到较高水平；而部分行业虽然涨幅较大，但资金流为负，说明系统在排序时并非只依赖价格因子。":
            "从表 6.2 可以看出，电子化学品、通信设备、其他电子、元件、半导体和消费电子等电子科技链行业在该时间窗口内整体表现较强，说明样本时点的市场热点明显集中在成长风格板块。值得注意的是，表中多数科技行业虽然 5 日涨幅较高，但资金流仍为负；只有能源金属表现出较明显的正向资金流。这说明系统生成的 total_score 并不简单等同于单一资金因子或单一涨跌幅，而是横截面标准化后多个因子共同作用的结果。",
        "这一结果与系统的模型设计是一致的。行业热度评分同时考虑动量、资金流、活跃度和波动率，因此能够较好平衡短期趋势与资金确认。对于毕业设计来说，这种结果不仅能够支持论文中的实验分析，也便于在答辩中解释评分逻辑。":
            "这一结果与系统代码中的模型设计是一致的。热力图快照不仅保留 value 和 total_score，还同步记录资金流、换手率、行业波动率和市值来源等字段，因此能够较好平衡短期趋势、资金确认和数据来源差异。需要强调的是，total_score 反映的是样本时点的多因子横截面综合状态，并不直接构成对未来收益的预测。对于毕业设计来说，这种结果既支持论文中的样例分析，也便于在答辩中解释为什么某些行业会在资金流暂时偏弱时仍保持较高的综合得分。",
        "同时，系统仍存在一些不足。首先，行业评分和龙头股评分的权重主要依据工程经验设定，尚未通过更系统的学习或优化方法进行自适应调整。其次，系统当前对政策文本、新闻舆情和产业链关系等数据的融合仍不充分。再次，系统的“实时性”主要体现为准实时刷新与历史快照持续积累，并非高频交易系统意义上的毫秒级实时。最后，行业识别结果与未来收益之间的定量回测仍可进一步强化。":
            "同时，系统仍存在一些不足。首先，行业评分和龙头股评分的权重主要依据工程经验设定，尚未通过更系统的学习或优化方法进行自适应调整。其次，虽然 THS-first 适配器提高了整体可用性，但部分行业成分股在实际运行中仍可能触发降级路径，只能返回不完整明细或代理字段。再次，行业模块当前主要依赖 REST 拉取、缓存和历史快照实现准实时更新，而非基于持续流式订阅的实时分析系统。最后，行业识别结果与后续收益表现之间的定量回测，以及前瞻性验证工作，仍有进一步强化空间。",
        "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的数据提供器、行业分析引擎、龙头股评分器、FastAPI 接口层和 React 仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。":
            "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
        "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型；在实现方面，本文说明了多源数据回退、缓存机制、历史快照、偏好持久化和前端可视化协同工作方式；在结果方面，本文结合项目保存的历史热力图快照展示了系统对阶段性强势行业的识别能力。":
            "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率代理为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型；在实现方面，本文说明了 THS-first 多源数据回退、分析层与路由层缓存、热力图历史快照和按 profile 偏好持久化的协同工作方式；在结果方面，本文结合项目保存的历史热力图快照展示了系统对阶段性强势行业的识别能力。",
        "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果的回测验证；第四，强化与研究工作台的联动，形成更完整的研究闭环。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。":
            "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果与后续收益表现之间的回测验证；第四，继续提高成分股映射完整度并优化核心行业分析链路。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。",
        "该公式中的权重与项目代码中的默认权重保持一致":
            "该公式中的权重与项目代码中的默认权重保持一致。需要强调的是，公式（4.2.1）对应的是横截面原始评分，并不直接等同于前端最终展示值。系统在完成加权之后，会进一步调用分数压缩函数，把结果统一映射到约 20 至 95 的展示区间，以减少样本集中时 0 分和 100 分贴边的问题。因此，排行榜和热力图中的 total_score 更适合用于相对比较，而不应被理解为行业的绝对评价值。",
        "此外，系统还使用 K-Means 聚类对行业进行辅助划分":
            "此外，系统还使用 K-Means 聚类对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。",
        "该模型突出盈利和成长的权重":
            "从工程实现看，公式（4.4.1）对应的是完整评分链路下的六维加权总分，其中 s1 至 s6 分别表示市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度的归一化得分。系统会在统一的 raw_data 结构上聚合估值、财务和行情字段，并按该权重体系生成总分。由于盈利能力和成长性权重相对较高，该模型更适合作为行业内核心标的筛选依据；与单纯按涨幅排序相比，它能够减少短期情绪对结果的扰动。同时需要说明的是，公式（4.4.1）主要对应 core 榜单的综合评分语义，而 hot 榜单则更强调短期涨幅与资金承接，并对暂时缺失的估值、盈利和成长字段采用中性处理。",
        "从工程实现看，完整评分链路会在统一的 raw_data 结构上聚合估值、财务和行情字段，并按六维权重体系生成总分。由于盈利能力和成长性权重相对较高，该模型更适合作为行业内核心标的筛选依据；与单纯按涨幅排序相比，它能够减少短期情绪对结果的扰动。":
            "从工程实现看，公式（4.4.1）对应的是完整评分链路下的六维加权总分，其中 s1 至 s6 分别表示市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度的归一化得分。完整评分链路会在统一的 raw_data 结构上聚合估值、财务和行情字段，并按该权重体系生成总分。因此，前端 core 榜单主要反映这一综合评分语义；与单纯按涨幅排序相比，它能够减少短期情绪对结果的扰动，更适合作为行业内核心标的筛选依据。",
        "在行业热度页面中，页面响应速度非常重要":
            "快速评分是从页面响应路径倒推出来的。行业热度页一次可能要同时给出多个候选股，如果每只股票都临时补抓完整财务数据，列表刷新会明显变慢。为此，系统在行业列表场景里优先复用成分股快照中已经带回的价格、成交额和资金承接等字段先生成一轮结果，只有用户进入详情页或继续深挖时再走完整评分链路。",
        "快速评分模式适合用于行业热度页中的候选股预筛选":
            "因此，快速评分和完整评分不是谁替代谁，而是分别服务于不同界面。前者解决的是“先把候选集合稳定列出来”，后者解决的是“把一只股票为什么排在前面解释清楚”。在结果组织上，前端继续把输出拆成 core（核心资产）和 hot（热点先锋）两类榜单：core 更接近行业内长期代表性筛选，hot 则专门保留短期涨幅和资金承接的冲击信息。这样用户先在行业页缩小范围，再到详情页看拆解，整个交互会更顺。",
        "从工程角度看，本系统具有以下几个明显特征":
            "从工程角度看，本系统具有以下几个明显特征。第一，系统以 THS-first 适配器为主入口，并通过 AKShare、新浪和腾讯形成多源补齐链路，保证了行业数据的连通性。第二，行业评分与龙头股评分同时引入快速评分链路与多级缓存机制，在保证响应速度的同时尽量维持评分口径一致。第三，系统输出不仅提供综合得分，还提供资金、波动率、市值来源和评分拆解等辅助信息，因此具有较好的结果解释性。",
        "在本次毕业设计与论文写作过程中":
            "在本次毕业设计与论文写作过程中，感谢指导教师沈文辉老师在选题论证、系统设计、论文结构与写作规范等方面给予的耐心指导与帮助。感谢项目开发和调试过程中提供建议与支持的同学和朋友，他们的交流与反馈为系统完善和论文修改提供了重要参考。",
        "同时，也感谢开源社区提供的 Python、FastAPI、React、AKShare 等工具和框架":
            "另外，论文能够完成，也离不开 Python、FastAPI、React、AKShare、scikit-learn 等开源工具在开发阶段提供的直接帮助。正是因为这些工具足够成熟，我才能把更多精力放在行业分析流程、页面交互和论文整理本身；也感谢学校和学院为毕业设计提供的学习环境与资源支持。"
    }

    for old_text, new_text in replacements.items():
        replace_first_paragraph_starting_with(doc, old_text, new_text)

    heatmap_snapshot = load_thesis_heatmap_snapshot(days=5)
    leader_case_payload = load_thesis_leader_case()
    heatmap_intro_text, heatmap_analysis_text, heatmap_conclusion_text = build_heatmap_result_section_text(heatmap_snapshot)
    leader_case_text = build_leader_case_analysis_text(leader_case_payload)
    replace_first_paragraph_starting_with(doc, "为了增强论文结果分析的真实性", heatmap_intro_text)
    replace_first_paragraph_starting_with(doc, "表 6.2 项目历史快照中的前十行业结果样例", "表 6.2 项目当前保留五日快照中的前十行业结果样例")
    replace_first_paragraph_starting_with(doc, "从表 6.2 可以看出", heatmap_analysis_text)
    replace_first_paragraph_starting_with(doc, "这一结果与系统代码中的模型设计是一致的。", heatmap_conclusion_text)

    replace_section_body(
        doc,
        "1.1 研究背景与意义",
        "1.2 国内外研究现状",
        [
            "在证券市场里，真正难做的往往不是盯住一两只股票，而是同时判断热点行业怎么切换、行业内部哪些公司更值得继续跟踪。一个行业持续走强，通常意味着资金、预期和景气度正在往同一方向聚集；而行业里的代表性公司又会因为规模、盈利和市场关注度，被更频繁地拿来比较。于是，如何从海量数据中尽快识别当期热点，并在行业内部继续缩小到更有代表性的股票，就成了一个既有研究价值也有实际用途的问题。",
            "过去这类工作很大程度依赖研究员人工切行情终端、翻行业报告和积累经验。行情节奏慢的时候，这种方式还勉强可行；但市场数据量和更新频率上来之后，单靠人工已经很难同时完成横截面比较、历史跟踪和多维交叉验证。也正因为如此，把大数据处理、量化评分和可视化页面真正接成一条可持续运行的研究链路，比单独写一套纸面分析方法更有工程意义。[1]",
            "本文依托的项目并不是单一的课程设计程序，而是一个较为完整的量化研究平台。当前公开仓库对外聚焦策略回测、实时行情与行业热度三个主工作区，同时仍保留部分相关研究能力代码与接口。在这些能力中，行业热度页面与毕业设计任务书的目标最为一致，因此本文不再泛化讨论平台全部模块，而是从现有工程实现中抽取行业热度子系统作为核心研究对象，在真实代码、真实接口与真实历史快照的基础上完成毕业论文撰写。",
            "对本文而言，这项工作的价值不只在于“做出一个页面”。一方面，系统把多源金融数据整理成可比较的行业研究结果，能够明显减少人工筛选成本；另一方面，评分过程保持了较强可解释性，便于把工程实现和理论依据放在一起说明。更现实的一点是，这个原型系统已经可以直接支撑毕业设计展示、运行截图获取和后续迭代，而不是停留在方案层面。",
        ],
    )

    replace_section_body(
        doc,
        "1.2 国内外研究现状",
        "1.3 研究内容与技术路线",
        [
            "1.2.1 金融大数据分析研究现状",
            "把现有文献和实际系统实现放在一起看，最先暴露出来的问题往往不是算法本身，而是数据怎么接、怎么对齐。行业快照、资金流、估值和财务表来自不同接口时，更新时间、单位和字段名经常对不上；同一个行业在不同平台里也可能对应不同叫法。孟小峰等对大数据管理技术的梳理说明，容量、速度、多样性和真实性这些约束在金融场景里一直存在。[1] 张晓琳等关于风险监测的研究也说明，多源金融数据融合已经成了很多金融分析工作的前提。[2]",
            "因此，本文借鉴这些研究时，更看重它们对数据组织方式的启发，而不是再额外设计一套抽象工具链。具体到本课题，后文采用 pandas、NumPy 和 scikit-learn，并不是因为这些工具本身更“先进”，而是因为它们已经足够支撑项目里真实存在的字段清洗、横截面比较和评分计算流程。[9-10]",
            "1.2.2 行业轮动与热点识别研究现状",
            "在行业轮动与热点识别方面，国内外研究普遍不主张只看单一指标。无论是从宏观周期、风格切换出发，还是从价格动量、资金流向和交易活跃度出发，最终都要回到多因子综合判断上来。[4][7-8] 这也解释了为什么本文没有把“热门行业”简单理解成短期涨幅排名，而是希望在横截面上同时观察价格、资金和活跃度这些信号。",
            "1.2.3 龙头股筛选与系统实现研究现状",
            "把文献里的龙头股识别方法和当前项目放在一起看，有一个很直接的共识：真正能长期代表行业状态的股票，通常不会只靠某一天的涨幅来判断。规模、盈利、成长和估值提供的是相对稳定的基本面线索，成交额、换手率和近期动量则更像市场有没有继续确认这一判断。[3][5-6]",
            "但不少既有研究停留在静态评价、因子验证或回测层面，对一套系统真正落地后要面对的缓存、回退、详情解释和页面展示写得比较少。本文之所以单列这一节，就是因为毕业设计依托的是已经在运行的行业子系统，后文需要把这些工程细节和评分逻辑一起说清楚，而不是只给出一套纸面指标体系。",
            "总体来看，现有研究已经为热门行业识别和龙头股筛选提供了较充分的理论基础，但一落到工程实现，就还会遇到两个明显空档：一是很少有人细讲多源适配、字段回退和页面展示这些真正影响可用性的细节；二是行业识别和龙头股遴选经常被拆开讨论，缺少统一的研究闭环。本文的重点并不是再堆一个复杂模型，而是在真实项目基础上，把多源数据获取、行业评分、龙头股筛选和前端展示这几部分真正接起来。",
        ],
    )

    replace_section_body(
        doc,
        "1.3 研究内容与技术路线",
        "1.4 论文结构安排",
        [
            "写作时，本文没有把研究内容拆成几块互不相干的模块，而是尽量按系统真实使用的顺序来梳理。更接近日常研究场景的路径，是先打开行业页，看热力图和排行榜怎样把横截面结果摆出来；接着沿着某个行业继续查看趋势、成分股和龙头股详情；最后再回头解释这些结果是怎样由适配器、缓存和评分器一步步算出来的。",
            "对应到实现链路，前端页面先请求热门行业、热力图和详情相关数据，路由层再把请求分发给 IndustryAnalyzer 和 LeaderStockScorer，分析层继续向 THS-first 主路径和补充数据源取回字段并整理结果。本文之所以按这条顺序展开，是因为这样最容易让论文里的截图、表格和项目代码彼此对上，也更方便答辩时解释页面里每一步是怎么来的。",
        ],
    )

    replace_section_body(
        doc,
        "1.4 论文结构安排",
        "相关技术与理论基础",
        [
            "全文按“提出问题、解释方法、落到系统、再看结果”的顺序展开。第一章说明选题背景、研究现状以及本文聚焦的行业子系统；第二章交代金融数据特征、多源采集和理论基础；第三章从需求、架构、数据流程和存储设计角度说明系统整体方案；第四章展开热门行业识别与龙头股遴选模型；第五章回到具体实现；第六章结合测试结果和固定快照样例分析系统输出，最后给出结论与后续改进方向。",
        ],
    )

    replace_section_body(
        doc,
        "2.1 金融大数据的特征",
        "2.2 多源数据采集与清洗",
        [
            "落到这个课题里，金融大数据最突出的麻烦并不只是“量大”这一个词，而是不同数据一起进来时口径很不整齐。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率不同，结构也不同，可信程度还会跟着来源变化。比如行业快照可以按日更新，财务指标却按季度披露；同一个字段在不同接口里又可能出现名称不一致、单位不同或缺失值写法不同的情况。真要把这些数据送进同一条分析链路，先解决标准化和可用性问题，比直接套模型更重要。",
            "对于热门行业识别尤其如此。只盯涨跌幅，很容易把短期情绪波动当成热点；只看资金流，又可能忽略趋势已经转弱的行业。行业研究真正需要的是把价格、资金、活跃度乃至风险约束放在同一横截面里比较，这也是本文后续采用多维评分而不是单指标排序的原因。",
        ],
    )

    replace_section_body(
        doc,
        "2.2 多源数据采集与清洗",
        "2.3 热门行业识别的理论基础",
        [
            "多源采集在本项目里不是可有可无的补充，而是整条行业分析主链路的一部分。行业子系统按 THS-first 组织数据接口：同花顺负责提供行业目录、行业摘要、资金流和领涨股这类主数据，AKShare 补行业元数据、成分股、估值与财务字段，新浪和腾讯主要承担回退与补缺。这样做并不是为了“多接几个接口”，而是因为单一来源很难同时覆盖行业研究真正需要的全部字段。",
            "数据拿到之后，马上要处理的不是模型，而是口径统一。项目里会先统一行业名称、字段命名和数值类型，再根据字段特征做缺失值处理、异常值裁剪和重复剔除。例如行业名称需要映射到统一的 industry_name，涨跌幅和资金流要变成可运算的浮点数，个股快照里暂时缺失的财务字段则用中性值参与快速评分。只有经过这一步，不同来源的数据才能进入后面的行业评分和龙头股筛选链路。[9]",
        ],
    )

    replace_section_body(
        doc,
        "2.4 龙头股评价的理论基础",
        "2.5 系统关键技术",
        [
            "在这个课题里，“龙头股”并不等于某一天涨得最快的股票。更常见的情况是，短期涨幅靠前的个股未必真的能代表行业，有时只是情绪推动；真正更能代表行业状态的，往往是那些规模、盈利、成长和市场关注度都更稳的公司。传统龙头企业识别也大多沿着这个思路，从市值规模、产业地位、盈利质量、成长能力、估值合理性和市场表现等多个维度综合判断。[3][5-6]",
            "项目里的做法也是如此：先把规模、盈利、成长、估值和交易信号放到一起，再看它们在同一行业内的相对位置。这样筛出来的结果不一定是短期最活跃的股票，但通常更容易解释，也更接近研究场景下“代表性公司”的含义。换句话说，本文希望筛出来的是能够代表行业结构和市场确认状态的标的，而不是一次性情绪冲高的样本。",
        ],
    )

    replace_section_body(
        doc,
        "3.3 非功能需求分析",
        "3.4 系统总体架构设计",
        [
            "从实际使用感受看，行业模块的非功能问题首先体现在等待时间上。用户切到行业页时，如果每次都重新抓数据、重算排行、再补详情，页面几秒内就会变得很拖沓，所以系统才把缓存拆到分析层和接口层，用来挡掉重复计算和重复请求。",
            "另一个很现实的问题是外部接口并不稳定。行业目录、资金流和个股字段只要有一个来源短时抖动，页面就可能出现空值或局部缺口，因此系统必须允许主路径失败后继续回退，并让热力图、排行榜和详情页尽量保持可用。再往后，如果还要继续接新的指标或研究视图，这套结构也不能每扩一块就大动一次，所以可扩展性同样是非功能需求的一部分。",
        ],
    )

    replace_section_body(
        doc,
        "3.4 系统总体架构设计",
        "3.5 数据流程设计",
        [
            "如图 3.1 所示，系统总体上采用四层结构：表现层、服务层、分析层和数据层。这样拆分不是为了把架构图画得更复杂，而是因为行业研究页面从请求发出到结果落地，确实会依次经过这四层。表现层主要是前端行业主页面、热力图、排行榜以及行业详情和龙头股详情弹窗，回放、观察列表和偏好配置等功能则作为辅助研究交互存在；服务层由 FastAPI 行业接口构成，负责请求接收、参数校验、缓存与响应封装；分析层由行业分析器和龙头股评分器组成，真正完成行业综合评分、龙头股筛选和评分拆解；数据层则以 THS-first 适配器为核心，统一协调同花顺、AKShare、新浪和腾讯等数据源。",
            "这样分层之后，每一层要解决的问题会清楚很多。表现层只需要关心图表交互和结果展示，服务层负责把分析结果组织成统一响应，分析层专注评分计算与结果解释，数据层则通过名称映射、节点回退、符号缓存和过期缓存保障等机制提高数据可用性。后面如果要替换局部数据源，或者微调评分权重，改动也大多集中在分析层和数据层，不必把前端页面整体重写一遍。",
            "和那种只依赖单一接口的数据抓取脚本相比，这套架构更适合支撑准实时行业研究。一方面，多源适配提高了行业目录、资金流、成分股、估值和历史走势等字段的覆盖率；另一方面，缓存和分层封装让页面响应速度与结果口径更容易保持一致。对论文写作来说，这一点也很重要，因为只有系统本身运行得足够稳定，后面拿到的快照样本和案例分析才有说服力。",
        ],
    )

    replace_section_body(
        doc,
        "3.6 存储设计",
        "热门行业识别与龙头股遴选模型设计",
        [
            "行业子系统现在采用的是“文件快照加浏览器本地状态”的持久化组合。热力图历史直接写入 data/industry/heatmap_history.json，观察行业与阈值配置按 profile 写入 data/industry_preferences/<profile>.json，龙头股财务缓存保存在 cache/financial_cache.json；前端页面再把回放所需的快照片段和当前选中状态留在 localStorage。",
            "这样处理首先是因为当前数据规模和使用方式都比较明确。答辩展示、论文复核和本地调试经常需要回看某个固定样本时点，直接读取这些文件会比额外再搭一层存储服务更省步骤；同时文件路径也更方便和脚本、截图、附录材料保持对应关系。对这套毕业设计原型来说，先把样例可复核和页面可回放做好，比把存储层做成复杂部署更重要。",
        ],
    )

    replace_section_body(
        doc,
        "6.5 对毕业设计目标的达成情况",
        "结 论",
        [
            "因此，本文已把任务书中的理论梳理、模型设计、系统实现和结果验证落到论文与项目证据中，主线完整。",
        ],
    )

    replace_section_body(
        doc,
        "5.1 数据提供器与回退机制实现",
        "5.2 行业分析模块实现",
        [
            "行业子系统的数据入口虽然封装在 SinaIndustryAdapter 中，但它在项目里的职责已经不止是调用单一新浪接口，更像是一个行业数据适配中枢。实际运行时，系统优先从同花顺接口拿行业目录、行业摘要、资金流、领涨股和行业指数，再由 AKShare 补齐行业元数据、成分股、估值、财务与历史行情；新浪和腾讯则更多承担回退和字段补缺任务。通过这种分工，系统逐渐形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行方式。[1-2]",
            "适配层真正麻烦的地方，在于同一行业在不同来源里的命名和节点并不一致。项目里为此维护了名称映射、节点映射、代理节点、符号缓存和反向映射等机制；当主路径失败时，再配合过期缓存把结果尽量补齐。这样做的目的不是多加一道兜底本身，而是让行业目录和成分股匹配尽量不断链。",
            "从论文角度看，这套适配设计最大的价值，是把行业摘要、资金流、成分股、估值和财务字段收拢到统一入口下，后面的行业评分和龙头股筛选就不需要分别适配多套来源。对工程实现来说，它也明显降低了外部数据源短时波动对系统可用性的影响。",
        ],
    )

    replace_section_body(
        doc,
        "5.2 行业分析模块实现",
        "5.3 龙头股评分模块实现",
        [
            "IndustryAnalyzer 更像是一个把原始行业字段整理成研究结果的中间层。它先利用行业摘要、资金流和少量横截面指标把热力图与排行榜所需的数据跑出来，尽量不在首屏就逐行业拉全量成分股；等用户继续下钻时，再补趋势统计、覆盖率和解释字段。",
            "在这条链路里，change_pct 或 weighted_change 负责描述近期强弱，flow_strength 表示资金承接，avg_volume 拿不到时再退到 turnover_rate，industry_volatility 则交给历史波动率或代理值处理。最后输出给前端的，已经不是零散原始列，而是一组带 total_score、资金流、换手率和说明字段的结构化结果，所以热力图、排行榜和趋势面板才能共享同一套分析口径。",
        ],
    )

    replace_section_body(
        doc,
        "5.3 龙头股评分模块实现",
        "5.4 后端接口实现",
        [
            "LeaderStockScorer 负责的也不只是“给股票打个分”。在同一行业里，它既要输出列表排序，又要给详情页提供评分拆解和原始字段，所以模块同时保留了完整评分与快照快速评分两条链路。前者更适合龙头股详情和深度研究，后者则是为了让行业列表页先把候选股票尽快筛出来。",
            "模块还会把结果拆成 core 和 hot 两类榜单。core 更偏向六维综合质量筛选，hot 更强调短期涨幅和资金承接；两者共用同一套 raw_data 结构，但在无法拿到 ROE 或增长字段时，快速链路会用中性分处理缺口。这样做的好处是，页面响应速度和评分口径之间可以取得相对平衡，而不是二选一。",
        ],
    )

    replace_section_body(
        doc,
        "5.4 后端接口实现",
        "5.5 前端可视化实现",
        [
            "后端接口层在项目里承担的是“把分析能力组织成可稳定调用的服务”。热门行业列表、成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情和偏好配置这些接口，基本把前端行业工作区需要的数据都覆盖了；聚类和轮动接口则更多承担辅助分析作用。[1-2][9-10]",
            "真正影响使用体验的还有缓存和兜底逻辑。路由层维护了 3 分钟端点缓存、30 分钟 parity 缓存，以及完整版成分股的异步构建机制；必要时还会返回过期缓存，避免外部数据源短时波动时页面直接空掉。热力图历史数据也设置了记录数和文件大小上限，保证快照文件长期可控。",
        ],
    )

    replace_section_body(
        doc,
        "5.5 前端可视化实现",
        "5.6 偏好配置与持续跟踪实现",
        [
            "前端行业工作区是按研究动作顺序排的，而不是先做若干孤立图表再拼到一起。页面顶部先给出当前市场概况和筛选条件，中间区域同时摆放热力图与排行榜，右侧保留龙头股榜单，用户一般会先看板块分布，再点进某个行业，最后查看个股详情。IndustryDashboard 这一页承担的，就是把这条浏览路径压缩到一次切换里完成。",
            "具体交互也尽量贴近项目实际使用方式。热力图切换颜色维度时，用户能直接看到板块冷热变化；尺寸维度和市值来源切换更多是为了判断排序背后的口径差异；排行榜里的回看窗口、波动过滤和排序项则帮助用户快速缩小范围。等用户打开详情弹窗后，评分拆解、价格走势和财务字段会接着出现，所以这里更像一条连续的研究路径，而不是几张静态图并排展示。",
            "如图 5.1 所示，行业热度总览页面将热力图、行业排行榜、趋势面板与龙头股入口整合在同一研究界面中，便于研究者快速完成行业筛选、排序比较与详情联动。",
            "图 5.1 行业热度总览界面",
            "图片来源：系统运行截图（作者自制）",
            "如图 5.2 所示，行业热力图页面支持按不同维度观察行业强弱与结构分布，既可以用颜色突出短期热度，也可以通过尺寸和市值来源信息帮助用户理解当前排序结果的背景。",
            "图 5.2 行业热力图界面",
            "图片来源：系统运行截图（作者自制）",
            "如图 5.3 所示，龙头股详情页面展示综合得分、维度拆解、价格走势与关键财务字段，能够帮助用户理解候选股票被遴选出的具体依据。",
            "图 5.3 龙头股详情界面",
            "图片来源：系统运行截图（作者自制）",
        ],
    )

    replace_section_body(
        doc,
        "6.4 系统特征与不足分析",
        "6.5 对毕业设计目标的达成情况",
        [
            "如果只看当前仓库真正保留下来的行业模块，它最有用的地方在于同一套数据入口已经能支撑从热力图到个股详情的连续分析。用户先在页面里看到行业排序，再点进趋势、龙头股和评分拆解时，背后用的仍是同一条适配、评分和缓存链路，因此结果解释不会前后脱节。",
            "但它的边界也很明确。权重目前还是工程经验优先，部分行业在成分股映射不完整时仍会走降级路径；模块刷新主要依赖 REST、缓存和历史快照，不是持续流式的实时系统；行业识别结果和后续收益之间的量化验证也还可以继续补强。这些问题不会影响毕业设计展示，但确实是后续迭代最该继续打磨的地方。",
        ],
    )

    replace_section_body(
        doc,
        "结 论",
        "参考文献",
        [
            "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
            "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率代理为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型。行业热度部分采用横截面标准化与加权合成的方法，并进一步压缩到便于展示的相对得分区间；龙头股部分则同时支持完整评分与快照快速评分两条链路，在保证口径一致的前提下兼顾分析完整性与页面响应效率。",
            "在工程实现方面，本文说明了 THS-first 多源数据回退、分析层与路由层缓存、热力图历史快照、财务缓存以及按 profile 偏好持久化的协同工作方式。第六章结合项目当前保留的五日热力图快照、龙头股双榜单样例和测试结果说明，该系统能够识别样本时点的阶段性强势行业，并进一步给出行业内部具有代表性的龙头股候选结果，具有较好的结果解释性与展示性。",
            "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果与后续收益表现之间的回测验证；第四，继续提高成分股映射完整度并优化核心行业分析链路。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。",
        ],
    )

    test_heading = next((p for p in doc.paragraphs if p.text.strip() == "6.2 功能测试"), None)
    if test_heading is not None:
        next_element = test_heading._element.getnext()
        intro_text = (
            "当前仓库保留的验证方式并不只有人工点页面这一种。分析层有行业评分与快路径回退相关单测，接口层有"
            "名称映射、偏好服务和龙头股列表/详情的返回结构检查，前端还保留了热力图切换、行业搜索和详情弹窗"
            "闭环的端到端脚本。表 6.1 汇总的是这些验证里与毕业设计主链路最相关的部分，因此更关注评分是否可"
            "解释、接口是否连通、页面操作是否能闭环，而不是生产环境压测。"
        )
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            next_para = Paragraph(next_element, test_heading._parent)
            if next_para.text.strip().startswith("功能测试主要分为三类"):
                replace_paragraph_text(next_para, intro_text)
            elif next_para.text.strip() != "表 6.1 主要功能测试结果":
                replace_paragraph_text(next_para, intro_text)
            else:
                intro_para = insert_paragraph_after(test_heading, intro_text)
                intro_para.paragraph_format.keep_with_next = True

    function_table = find_table_containing_text(doc, "模块")
    if function_table is not None and len(function_table.rows) >= 6 and len(function_table.columns) >= 3:
        function_table.cell(1, 1).text = "获取 THS 行业目录、行业摘要、成分股与财务等数据"
        function_table.cell(1, 2).text = "THS-first 适配器主导，AKShare/Sina/腾讯补齐"
        function_table.cell(2, 2).text = "横截面评分、波动率代理、聚类与轮动分析"
        function_table.cell(4, 2).text = "React 仪表盘、热力图与详情联动"
        function_table.cell(5, 2).text = "按 profile 的 JSON 持久化与浏览器缓存"

    sample_table = find_table_containing_text(doc, "资金流(元)")
    if sample_table is None:
        sample_table = find_table_containing_text(doc, "资金流(亿元)")
    update_heatmap_sample_table(sample_table, heatmap_snapshot)

    test_table = find_table_containing_text(doc, "测试项目")
    if test_table is None:
        test_table = find_table_containing_text(doc, "输入或操作")
    if test_table is not None and len(test_table.rows) >= 6 and len(test_table.columns) >= 4:
        test_table.cell(0, 0).text = "测试项目"
        row_values = [
            ("行业评分计算与回退", "运行行业评分逻辑并验证快路径与字段回退", "快速评分链路可输出 20-95 展示得分，代理波动率与字段回退生效", "满足预期"),
            ("THS-first 适配器映射", "验证行业名称映射和多源回退", "THS 主源可用，AKShare/Sina/腾讯补齐链路与缓存兜底正常", "满足预期"),
            ("龙头股列表与详情接口", "调用龙头股列表与详情接口", "core/hot 榜单与详情接口均返回综合得分、维度拆解和原始字段", "满足预期"),
            ("偏好配置持久化", "保存并重置观察行业与阈值配置", "JSON 与浏览器状态保持一致", "满足预期"),
            ("前端行业页面闭环", "切换热力图周期、搜索行业并打开详情", "热力图、排行榜与详情弹窗联动正常", "满足预期"),
        ]
        for row_index, values in enumerate(row_values, start=1):
            for col_index, value in enumerate(values):
                test_table.cell(row_index, col_index).text = value

    ensure_section_trailing_paragraph(
        doc,
        "6.4 系统特征与不足分析",
        "为保证龙头股案例",
        leader_case_text,
    )

    append_thesis_expansion_content(doc)
    apply_project_focused_thesis_revision(doc)
    append_project_depth_content(doc)
    apply_final_project_consistency_cleanup(doc)
    insert_project_depth_tables(doc)
    insert_architecture_figure(doc)
    insert_task_completion_table(doc)
    normalize_numbered_table_captions(doc)


def find_paragraph_by_text(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_last_paragraph_by_text(doc: Document, text: str):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_previous_drawing_paragraph(doc: Document, anchor_paragraph):
    paragraphs = doc.paragraphs
    anchor_index = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph._element is anchor_paragraph._element:
            anchor_index = idx
            break
    if anchor_index is None:
        raise RuntimeError(f"Anchor paragraph not found in document: {anchor_paragraph.text}")
    for idx in range(anchor_index - 1, -1, -1):
        paragraph = paragraphs[idx]
        if "w:drawing" in paragraph._element.xml:
            return paragraph
    raise RuntimeError(f"Drawing paragraph not found before: {anchor_paragraph.text}")


def ensure_figure_source_paragraph(caption):
    next_element = caption._element.getnext()
    if next_element is not None and next_element.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        next_para = Paragraph(next_element, caption._parent)
        if next_para.text.strip().startswith("图片来源："):
            replace_paragraph_text(next_para, SOURCE_NOTE)
            return next_para
    return insert_paragraph_after(caption, SOURCE_NOTE)


def ensure_figure_intro_paragraph(image_para, intro_text: str):
    from docx.text.paragraph import Paragraph

    figure_match = re.match(r"^(如图\s*\d+\.\d+)\s*所示", intro_text)
    figure_prefix = figure_match.group(1) if figure_match else None
    previous = image_para._element.getprevious()
    while previous is not None:
        if previous.tag != qn("w:p"):
            previous = previous.getprevious()
            continue
        previous_para = Paragraph(previous, image_para._parent)
        previous_text = previous_para.text.strip()
        if not previous_text:
            previous = previous.getprevious()
            continue
        if previous_text == intro_text or (figure_prefix and previous_text.startswith(figure_prefix)):
            replace_paragraph_text(previous_para, intro_text)
            return previous_para
        break

    intro_para = insert_paragraph_before_element(image_para._element, image_para._parent)
    intro_para.add_run(intro_text)
    return intro_para


def relayout_figures(doc: Document) -> None:
    body_section = doc.sections[-1]
    max_body_width = body_section.page_width - body_section.left_margin - body_section.right_margin
    TMP_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for title, config in FIGURE_IMAGES.items():
        caption = find_paragraph_by_text(doc, title)
        previous_element = caption._element.getprevious()
        image_para = None
        if previous_element is not None and previous_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            previous_para = Paragraph(previous_element, caption._parent)
            if "w:drawing" in previous_para._element.xml and not previous_para.text.strip():
                image_para = previous_para
        if image_para is None:
            image_para = insert_paragraph_before_element(caption._element, caption._parent)

        previous = image_para._element.getprevious()
        if previous is not None and previous.tag == qn("w:p") and 'w:type="page"' in previous.xml:
            previous.getparent().remove(previous)
        source_path = Path(config["path"])
        crop_box = config["crop"]
        output_path = TMP_ASSET_DIR / f"{config['slug']}.png"
        with Image.open(source_path) as img:
            cropped = img.crop(crop_box)
            bordered = ImageOps.expand(cropped, border=4, fill="#C9D4E0")
            bordered.save(output_path)
        clear_paragraph(image_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        image_para.paragraph_format.line_spacing = 1.0
        image_para.paragraph_format.space_before = Pt(0)
        image_para.paragraph_format.space_after = Pt(0)
        image_para.paragraph_format.first_line_indent = Pt(0)
        remove_page_break_before_flag(image_para)
        width_scale = config.get("width_scale", 0.97)
        target_width = Emu(int(max_body_width * width_scale))
        image_para.add_run().add_picture(str(output_path), width=target_width)
        intro_text = FIGURE_INTRO_PARAGRAPHS.get(title)
        if intro_text:
            intro_para = ensure_figure_intro_paragraph(image_para, intro_text)
            intro_para.paragraph_format.keep_with_next = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(0)
        caption.paragraph_format.keep_with_next = True
        image_para.paragraph_format.keep_with_next = True
        source_para = ensure_figure_source_paragraph(caption)
        source_para.paragraph_format.space_before = Pt(0)
        source_para.paragraph_format.space_after = Pt(3)
        source_para.paragraph_format.keep_with_next = False


def _set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_borders(table, top: bool, header_bottom: bool, bottom: bool) -> None:
    rows = table.rows
    for row_idx, row in enumerate(rows):
        _set_row_no_split(row)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                node = tc_borders.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_borders.append(node)
                enabled = False
                if edge == "top" and top and row_idx == 0:
                    enabled = True
                elif edge == "bottom" and header_bottom and row_idx == 0:
                    enabled = True
                elif edge == "bottom" and bottom and row_idx == len(rows) - 1:
                    enabled = True
                if enabled:
                    node.set(qn("w:val"), "single")
                    node.set(qn("w:sz"), "8")
                    node.set(qn("w:space"), "0")
                    node.set(qn("w:color"), "000000")
                else:
                    node.set(qn("w:val"), "nil")


def set_table_column_widths(table, widths: tuple[Emu | Cm, ...]) -> None:
    table.autofit = False
    for col_idx, width in enumerate(widths):
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width


def format_data_tables(doc: Document) -> None:
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "姓    名：" in text or "论文题目：" in text:
            continue
        if "行业热力图 / 行业排行榜" in text and "THS 主数据 / AKShare 增强" in text:
            continue
        if "Sindustry" in text or "Sleader" in text or "m:oMath" in table._tbl.xml:
            continue
        if not table.rows:
            continue
        is_main_test_table = "行业评分计算与回退" in text and "前端行业页面闭环" in text
        if is_main_test_table:
            set_table_column_widths(table, (Cm(2.9), Cm(3.4), Cm(6.1), Cm(2.2)))
        set_table_borders(table, top=True, header_bottom=True, bottom=True)
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    if row_index == 0 or (is_main_test_table and cell_index == 3):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing = Pt(15) if is_main_test_table else Pt(20)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = row_index == 0
                        font_size = Pt(9) if is_main_test_table and row_index != 0 else Pt(10.5)
                        set_run_text_font(run, "宋体", font_size)


def append_omml_equation(paragraph, equation_text: str, font_half_points: int = 20) -> None:
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_run_props = OxmlElement("m:rPr")
    for tag in ("w:sz", "w:szCs"):
        size_node = OxmlElement(tag)
        size_node.set(qn("w:val"), str(font_half_points))
        math_run_props.append(size_node)
    math_run.append(math_run_props)
    math_text = OxmlElement("m:t")
    math_text.text = equation_text
    math_run.append(math_text)
    omath.append(math_run)
    omath_para.append(omath)
    paragraph._p.append(omath_para)


def replace_formula_block(doc: Document, keyword: str, equation_text: str, equation_number: str) -> None:
    stale_formula_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith(keyword) or paragraph.text.strip() == equation_number
    ]
    anchor_element = None
    anchor_parent = None
    target_table = find_table_containing_text(doc, keyword)
    if target_table is None:
        target_table = next((table for table in doc.tables if keyword in table._tbl.xml), None)
    if target_table is not None:
        anchor_element = target_table._element
        anchor_parent = target_table._parent
    else:
        first_stale_paragraph = stale_formula_paragraphs[0] if stale_formula_paragraphs else None
        if first_stale_paragraph is not None:
            anchor_element = first_stale_paragraph._element
            anchor_parent = first_stale_paragraph._parent

    if anchor_element is None or anchor_parent is None:
        return

    equation_para = insert_paragraph_before_element(anchor_element, anchor_parent)
    equation_para.style = "Normal"
    equation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_para.paragraph_format.first_line_indent = Pt(0)
    equation_para.paragraph_format.space_before = Pt(6)
    equation_para.paragraph_format.space_after = Pt(0)
    equation_para.paragraph_format.line_spacing = BODY_LINE_SPACING
    equation_run = equation_para.add_run(equation_text)
    equation_run.font.name = "Times New Roman"
    equation_run.font.size = Pt(10.5)
    set_east_asia_font(equation_run, "Times New Roman")

    number_para = insert_paragraph_after(equation_para)
    number_para.style = "Normal"
    number_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_para.paragraph_format.first_line_indent = Pt(0)
    number_para.paragraph_format.space_before = Pt(0)
    number_para.paragraph_format.space_after = Pt(6)
    number_para.paragraph_format.line_spacing = BODY_LINE_SPACING
    number_run = number_para.add_run(equation_number)
    number_run.font.name = "Times New Roman"
    number_run.font.size = Pt(10.5)
    set_east_asia_font(number_run, "Times New Roman")

    anchor_element.getparent().remove(anchor_element)
    keep_elements = {equation_para._element, number_para._element}
    for paragraph in list(doc.paragraphs):
        stripped = paragraph.text.strip()
        if paragraph._element in keep_elements:
            continue
        if stripped.startswith(keyword) or stripped == equation_number:
            paragraph._element.getparent().remove(paragraph._element)


def polish_academic_tone(doc: Document) -> None:
    """Replace overly colloquial wording with submission-ready academic prose."""
    exact_replacements = {
        "论文写作时，本文直接使用项目保留下来的热力图历史快照和本地运行结果做样例验证。按五日窗口观察，电子化学品、通信设备、半导体、消费电子和能源金属等阶段性强势行业能够被系统稳定排到前列，行业内部也能继续给出代表性龙头候选。换句话说，这套原型系统至少已经能够支撑“先识别热点，再往行业内部找代表股票”这一最核心的研究流程。":
            "论文写作时，本文直接使用项目保留的热力图历史快照和本地运行结果进行样例验证。按五日窗口观察，能源金属、元件、非金属材料、军工电子、军工装备和小金属等阶段性强势行业位于前列，行业内部也能继续给出代表性龙头候选。结果表明，这套原型系统能够支撑“先识别热点行业，再在行业内部遴选代表股票”的核心研究流程。",
        "从实际使用感受看，行业模块的非功能问题首先体现在等待时间上。用户切到行业页时，如果每次都重新抓数据、重算排行、再补详情，页面几秒内就会变得很拖沓，所以系统才把缓存拆到分析层和接口层，用来挡掉重复计算和重复请求。":
            "从实际使用过程看，行业模块的非功能需求首先体现在响应时间上。用户切换到行业页时，如果每次都重新获取数据、计算排行榜并补充详情，页面响应会明显变慢。因此，系统在分析层和接口层分别设置缓存机制，用于减少重复计算和重复请求。",
        "LeaderStockScorer 负责的也不只是“给股票打个分”。在同一行业里，它既要输出列表排序，又要给详情页提供评分拆解和原始字段，所以模块同时保留了完整评分与快照快速评分两条链路。前者更适合龙头股详情和深度研究，后者则是为了让行业列表页先把候选股票尽快筛出来。":
            "LeaderStockScorer 的职责并不只是生成股票得分。在同一行业内，该模块既需要输出列表排序，也需要向详情页提供评分拆解和原始字段，因此同时保留了完整评分与快照快速评分两条链路。前者更适合龙头股详情和深度研究，后者则用于在行业列表页中快速形成候选股票集合。",
        "如果只看当前仓库真正保留下来的行业模块，它最有用的地方在于同一套数据入口已经能支撑从热力图到个股详情的连续分析。用户先在页面里看到行业排序，再点进趋势、龙头股和评分拆解时，背后用的仍是同一条适配、评分和缓存链路，因此结果解释不会前后脱节。":
            "从当前仓库实际保留的行业模块来看，其主要价值在于同一套数据入口已经能够支撑从热力图到个股详情的连续分析。用户先在页面中查看行业排序，再进入趋势、龙头股和评分拆解时，背后仍然使用同一条适配、评分和缓存链路，因此结果解释具有较好的连贯性。",
        "但它的边界也很明确。权重目前还是工程经验优先，部分行业在成分股映射不完整时仍会走降级路径；模块刷新主要依赖 REST、缓存和历史快照，不是持续流式的实时系统；行业识别结果和后续收益之间的量化验证也还可以继续补强。这些问题不会影响毕业设计展示，但确实是后续迭代最该继续打磨的地方。":
            "但系统边界也比较明确。权重目前仍以工程经验为主，部分行业在成分股映射不完整时仍会进入降级路径；模块刷新主要依赖 REST、缓存和历史快照，并不是持续流式的实时系统；行业识别结果与后续收益之间的量化验证也仍需进一步补强。这些问题不会影响毕业设计展示，但属于后续迭代中需要继续优化的重点。",
        "另一个很现实的问题是外部接口并不稳定。行业目录、资金流和个股字段只要有一个来源短时抖动，页面就可能出现空值或局部缺口，因此系统必须允许主路径失败后继续回退，并让热力图、排行榜和详情页尽量保持可用。再往后，如果还要继续接新的指标或研究视图，这套结构也不能每扩一块就大动一次，所以可扩展性同样是非功能需求的一部分。":
            "另一个需要考虑的问题是外部接口稳定性。行业目录、资金流和个股字段只要有一个来源出现短时波动，页面就可能出现空值或局部缺口，因此系统必须允许主路径失败后继续回退，并让热力图、排行榜和详情页尽量保持可用。后续若继续接入新的指标或研究视图，系统结构也应避免每次扩展都进行大范围改造，因此可扩展性同样是非功能需求的一部分。",
        "行业子系统的数据入口虽然封装在 SinaIndustryAdapter 中，但它在项目里的职责已经不止是调用单一新浪接口，更接近是一个行业数据适配中枢。实际运行时，系统优先从同花顺接口拿行业目录、行业摘要、资金流、领涨股和行业指数，再由 AKShare 补充行业元数据、成分股、估值、财务与历史行情；新浪和腾讯则更多承担回退和字段补缺任务。通过这种分工，系统逐渐形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行方式。[1-2]":
            "行业子系统的数据入口虽然封装在 SinaIndustryAdapter 中，但其职责已经不再局限于调用单一新浪接口，而是承担行业数据适配中枢的作用。实际运行时，系统优先从同花顺接口获取行业目录、行业摘要、资金流、领涨股和行业指数，再由 AKShare 补充行业元数据、成分股、估值、财务与历史行情；新浪和腾讯则更多承担回退和字段补缺任务。通过这种分工，系统逐渐形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行方式。[1-2]",
        "适配层真正麻烦的地方，在于同一行业在不同来源里的命名和节点并不一致。项目里为此维护了名称映射、节点映射、代理节点、符号缓存和反向映射等机制；当主路径失败时，再配合过期缓存把结果尽量补充。这样做的目的不是增加回退机制本身，而是让行业目录和成分股匹配尽量不断链。":
            "适配层的核心难点在于同一行业在不同来源中的命名和节点并不一致。项目为此维护了名称映射、节点映射、代理节点、符号缓存和反向映射等机制；当主路径失败时，再配合过期缓存尽量补充结果。这样做的目的不是单纯增加回退机制，而是尽量保持行业目录和成分股匹配链路的连续性。",
        "在代码组织上，适配层承担了很多看起来琐碎但非常关键的工作。例如行业节点需要映射，股票名称需要转换成代码，部分字段需要从字符串转成数值，某些接口失败后还要尝试代理节点或缓存结果。这些工作如果散落到前端或评分器里，系统会很难维护。":
            "在代码组织上，适配层承担了大量细节性但关键的工作。例如行业节点需要映射，股票名称需要转换成代码，部分字段需要从字符串转成数值，某些接口失败后还要尝试代理节点或缓存结果。若这些工作分散到前端或评分器中，系统维护成本会明显上升。",
        "IndustryAnalyzer 更接近是一个把原始行业字段整理成研究结果的中间层。它先利用行业摘要、资金流和少量横截面指标把热力图与排行榜所需的数据生成，尽量不在首屏就逐行业拉全量成分股；等用户继续下钻时，再补趋势统计、覆盖率和解释字段。":
            "IndustryAnalyzer 的作用是把原始行业字段整理为研究结果。它先利用行业摘要、资金流和少量横截面指标生成热力图与排行榜所需的数据，尽量避免在首屏阶段逐行业获取全量成分股；当用户继续下钻时，再补充趋势统计、覆盖率和解释字段。",
        "在证券市场里，真正难做的往往不是盯住一两只股票，而是同时判断热点行业怎么切换、行业内部哪些公司更值得继续跟踪。一个行业持续走强，通常意味着资金、预期和景气度正在往同一方向聚集；而行业里的代表性公司又会因为规模、盈利和市场关注度，被更频繁地拿来比较。于是，如何从海量数据中尽快识别当期热点，并在行业内部继续缩小到更有代表性的股票，就成了一个既有研究价值也有实际用途的问题。":
            "在证券市场研究中，难点往往不在于跟踪少数个股，而在于同时判断热点行业的切换方向以及行业内部代表性公司的变化。行业持续走强通常意味着资金、预期和景气度在一定阶段内形成共振，而行业内代表性公司又会因规模、盈利能力和市场关注度差异而呈现不同表现。因此，如何从海量数据中识别当期热点行业，并在行业内部进一步遴选具有代表性的股票，成为兼具研究价值和应用价值的问题。",
        "过去这类工作很大程度依赖研究员人工切行情终端、翻行业报告和积累经验。行情节奏慢的时候，这种方式还勉强可行；但市场数据量和更新频率上来之后，单靠人工已经很难同时完成横截面比较、历史跟踪和多维交叉验证。也正因为如此，把大数据处理、量化评分和可视化页面真正接成一条可持续运行的研究链路，比单独写一套纸面分析方法更有工程意义。[1-2]":
            "过去，这类工作较多依赖研究人员切换行情终端、阅读行业报告和积累经验。在市场节奏较慢时，该方式尚可满足部分研究需求；但随着市场数据规模和更新频率提高，仅依赖人工方式已经难以同时完成横截面比较、历史跟踪和多维交叉验证。因此，将大数据处理、量化评分和可视化展示连接为可持续运行的研究链路，相比单纯提出纸面分析方法更具有工程实践意义。[1-2]",
        "从工程角度看，这项工作的价值不只在于“做出一个页面”。一方面，系统把多源金融数据整理成可比较的行业研究结果，能够减少人工筛选成本；另一方面，评分过程保持较强可解释性，便于把工程实现、模型口径和样例结果放在一起说明。更现实的一点是，这个原型已经可以支撑毕业设计展示、运行截图获取和后续迭代，而不是停留在方案层面。":
            "从工程角度看，这项工作的价值不仅在于实现一个展示页面。一方面，系统将多源金融数据整理为可比较的行业研究结果，能够降低人工筛选成本；另一方面，评分过程保持较强可解释性，便于将工程实现、模型口径和样例结果统一说明。从项目完成度看，该原型已经能够支撑毕业设计展示、运行截图获取和后续迭代，而不是停留在方案层面。",
        "本文没有把研究内容组织为几块互不相干的模块，而是尽量按系统真实使用的顺序来梳理。更接近日常研究场景的路径，是先打开行业页，看热力图和排行榜怎样把横截面结果呈现出来；接着沿着某个行业继续查看趋势、成分股和龙头股详情；最后再回头解释这些结果是怎样由适配器、缓存和评分器一步步算出来的。":
            "本文没有将研究内容组织为若干相互割裂的模块，而是按照系统实际使用顺序进行梳理。较符合常规行业研究场景的路径是：首先进入行业页面，观察热力图和排行榜呈现的横截面结果；随后围绕某一行业继续查看趋势、成分股和龙头股详情；最后再解释这些结果由适配器、缓存和评分器逐步生成的过程。",
        "对应到实现链路，前端页面先请求热门行业、热力图和详情相关数据，路由层再把请求分发给 IndustryAnalyzer 和 LeaderStockScorer，分析层继续向 THS-first 主路径和补充数据源取回字段并整理结果。本文之所以按这条顺序展开，是因为这样最容易让论文里的截图、表格和项目代码彼此对上，也更方便答辩时解释页面里每一步是怎么来的。":
            "对应到实现链路，前端页面先请求热门行业、热力图和详情相关数据，路由层再将请求分发给 IndustryAnalyzer 和 LeaderStockScorer，分析层继续向 THS-first 主路径和补充数据源获取字段并整理结果。本文按照这一路径展开，能够使论文中的截图、表格和项目代码形成对应关系，也便于在答辩说明中解释页面展示结果的生成过程。",
        "当前仓库保留的验证方式并不只有人工点页面这一种。分析层有行业评分与快路径回退相关单测，接口层有名称映射、偏好服务和龙头股列表/详情的返回结构检查，前端还保留了热力图切换、行业搜索和详情弹窗闭环的端到端脚本。表 6.1 汇总的是这些验证里与毕业设计主链路最相关的部分，因此更关注评分是否可解释、接口是否连通、页面操作是否能闭环，而不是生产环境压测。":
            "当前系统的验证方式并不限于人工页面检查。分析层包括行业评分与快速路径回退相关单元测试，接口层包括名称映射、偏好服务和龙头股列表/详情返回结构检查，前端则保留了热力图切换、行业搜索和详情弹窗闭环的端到端脚本。表 6.1 汇总了这些验证中与毕业设计核心链路最相关的部分，因此重点关注评分是否可解释、接口是否连通以及页面操作是否形成闭环，而不是生产环境压力测试。",
        "接口测试主要检查返回结构。热门行业、热力图、趋势、龙头股列表和详情接口，都需要返回前端能够直接消费的字段。如果后端某次改动导致字段名称变化，前端页面即使能加载，也可能出现空图、空榜单或详情缺失。":
            "接口测试主要检查返回结构。热门行业、热力图、趋势、龙头股列表和详情接口，都需要返回前端能够直接使用的字段。若后端接口调整导致字段名称变化，前端页面即使能够加载，也可能出现图表为空、榜单为空或详情缺失等问题。",
        "前端测试的重点不是逐个按钮截图，而是验证用户路径是否连贯。一个完整路径应当包括进入行业页、查看热力图、切换统计窗口、选择行业、打开龙头股详情、查看评分拆解，并在必要时回到历史快照。只要这条路径稳定，毕业设计展示的可信度就会明显提高。":
            "前端测试的重点不是逐项记录按钮截图，而是验证使用路径是否连贯。一个完整路径应当包括进入行业页、查看热力图、切换统计窗口、选择行业、打开龙头股详情、查看评分拆解，并在必要时回到历史快照。只要该路径保持稳定，毕业设计展示的可信度就会明显提高。",
        "固定快照的意义在于把动态市场状态固定下来。若论文每次生成都重新抓取最新行情，表 6.2 的行业排序会不断变化，老师和读者很难复核。使用固定快照后，样例分析就可以围绕同一时点展开，结论也不会随着行情刷新而漂移。":
            "固定快照的意义在于将动态市场状态固化为可复核样本。若论文每次生成都重新获取最新行情，表 6.2 的行业排序会不断变化，评阅者和读者将难以复核。使用固定快照后，样例分析可以围绕同一时点展开，结论也不会随着行情刷新而漂移。",
        "系统的第二个特征是结果解释性较强。行业得分可以拆回动量、资金、活跃度和波动代理，个股得分可以拆回规模、估值、盈利、成长、动量和活跃度。即使用户不同意当前权重，也能理解系统为什么给出这样的排序。":
            "系统的第二个特征是结果解释性较强。行业得分可以拆解为动量、资金、活跃度和波动代理，个股得分可以拆解为规模、估值、盈利、成长、动量和活跃度。即使研究者对当前权重设置存在不同判断，也能够理解系统生成该排序的依据。",
        "另外，论文能够完成，也离不开 Python、FastAPI、React、AKShare、scikit-learn 等开源工具在开发阶段提供的直接帮助。正是因为这些工具足够成熟，我才能把更多精力放在行业分析流程、页面交互和论文整理本身；也感谢学校和学院为毕业设计提供的学习环境与资源支持。":
            "另外，论文能够完成，也离不开 Python、FastAPI、React、AKShare、scikit-learn 等开源工具在开发阶段提供的支持。正是因为这些工具较为成熟，本人得以将更多精力放在行业分析流程、页面交互和论文整理本身；也感谢学校和学院为毕业设计提供的学习环境与资源支持。",
        "把现有文献和实际系统实现放在一起看，最先暴露出来的问题往往不是算法本身，而是数据怎么接、怎么对齐。行业快照、资金流、估值和财务表来自不同接口时，更新时间、单位和字段名经常对不上；同一个行业在不同平台里也可能对应不同叫法。大数据管理、商业智能和数据密集型系统相关研究都强调，容量、速度、多样性、真实性和可维护性会共同影响金融数据系统能否稳定运行。[1][24-29]":
            "结合现有文献和实际系统实现可以发现，首先需要解决的问题往往不是算法本身，而是数据接入与字段对齐。行业快照、资金流、估值和财务表来自不同接口时，更新时间、单位和字段名称经常不一致；同一行业在不同平台中也可能对应不同名称。大数据管理、商业智能和数据密集型系统相关研究都强调，容量、速度、多样性、真实性和可维护性会共同影响金融数据系统能否稳定运行。[1][24-29]",
        "落到这个课题里，金融大数据最突出的问题并不只是“量大”这一个词，而是多类数据进入同一链路时口径差异较大。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率不同，结构也不同，可信程度还会跟着来源变化。比如行业快照可以按日更新，财务指标却按季度披露；同一个字段在不同接口里又可能出现名称不一致、单位不同或缺失值写法不同的情况。若要将这些数据纳入同一条分析链路，先解决标准化和可用性问题，比直接套模型更重要。[24-29]":
            "在本课题中，金融大数据的突出问题并不只是数据规模较大，而是多类数据进入同一链路时存在较明显的口径差异。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率、数据结构和可信程度均可能随来源变化。例如，行业快照可以按日更新，财务指标通常按季度披露；同一字段在不同接口中又可能出现名称不一致、单位不同或缺失值表示方式不同等情况。因此，在建模之前先解决标准化和可用性问题，是系统能够稳定运行的前提。[24-29]",
        "这样处理首先是因为当前数据规模和使用方式都比较明确。答辩展示、论文复核和本地调试经常需要回看某个固定样本时点，直接读取这些文件会比额外再搭一层存储服务更省步骤；同时文件路径也更方便和脚本、截图、附录材料保持对应关系。对这套毕业设计原型来说，先把样例可复核和页面可回放做好，比把存储层做成复杂部署更重要。":
            "这样处理首先是因为当前数据规模和使用方式较为明确。答辩演示、论文复核和本地调试经常需要回看某一固定样本时点，直接读取这些文件比额外引入存储服务更加简洁；同时，文件路径也便于与脚本、截图和附录材料保持对应关系。对于本毕业设计原型而言，优先保证样例可复核和页面可回放，比构建复杂存储部署更符合当前阶段需求。",
        "从工程角度看，本地 JSON、浏览器 localStorage 和前端状态并不是彼此孤立的。后端负责保存可复核的历史数据，前端负责保留使用者当前研究状态，二者配合后，页面既能支持答辩展示，也能支持日常复盘。":
            "从工程角度看，本地 JSON、浏览器 localStorage 和前端状态并不是彼此孤立的。后端负责保存可复核的历史数据，前端负责保留使用者当前研究状态，二者配合后，系统既能支持答辩演示中的样例复核，也能支持后续复盘分析。",
        "本文的不足也比较明确。首先，行业热度和龙头股评分权重主要依据工程经验设定，尚未经过系统参数优化；其次，样例分析使用的是固定快照，不能替代长周期回测；再次，政策文本、新闻舆情和产业链上下游关系尚未进入核心得分。把这些边界写清楚，有助于让论文结论更稳，也有助于后续迭代继续推进。":
            "本文的不足也比较明确。首先，行业热度和龙头股评分权重主要依据工程经验设定，尚未经过系统参数优化；其次，样例分析使用的是固定快照，不能替代长周期回测；再次，政策文本、新闻舆情和产业链上下游关系尚未进入核心得分。明确上述边界，有助于提高论文结论的稳健性，也为后续迭代提供了方向。",
        "落到这个课题里，金融大数据最突出的麻烦并不只是“量大”这一个词，而是不同数据一起进来时口径很不整齐。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率不同，结构也不同，可信程度还会跟着来源变化。比如行业快照可以按日更新，财务指标却按季度披露；同一个字段在不同接口里又可能出现名称不一致、单位不同或缺失值写法不同的情况。真要把这些数据送进同一条分析链路，先解决标准化和可用性问题，比直接套模型更重要。[24-29]":
            "在本课题中，金融大数据的突出问题并不只是数据规模较大，而是多类数据进入同一链路时存在较明显的口径差异。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率、数据结构和可信程度均可能随来源变化。例如，行业快照可以按日更新，财务指标通常按季度披露；同一字段在不同接口中又可能出现名称不一致、单位不同或缺失值表示方式不同等情况。因此，在建模之前先解决标准化和可用性问题，是系统能够稳定运行的前提。[24-29]",
        "落到这个课题里，金融大数据最突出的麻烦并不只是“量大”这一个词，而是不同数据一起进来时口径很不整齐。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率不同，结构也不同，可信程度还会跟着来源变化。比如行业快照可以按日更新，财务指标却按季度披露；同一个字段在不同接口里又可能出现名称不一致、单位不同或缺失值写法不同的情况。真要把这些数据送进同一条分析链路，先解决标准化和可用性问题，比直接套模型更重要。":
            "在本课题中，金融大数据的突出问题并不只是数据规模较大，而是多类数据进入同一链路时存在较明显的口径差异。股票日线、行业资金流、财务报表、估值指标和新闻文本的更新频率、数据结构和可信程度均可能随来源变化。例如，行业快照可以按日更新，财务指标通常按季度披露；同一字段在不同接口中又可能出现名称不一致、单位不同或缺失值表示方式不同等情况。因此，在建模之前先解决标准化和可用性问题，是系统能够稳定运行的前提。",
    }
    phrase_replacements = {
        "A 股市场的行业轮动很快，真正做研究时，光靠人工盯盘、翻资讯和手工比对往往跟不上节奏。":
            "A 股市场行业轮动速度较快，仅依赖人工盯盘、资讯阅读和手工比对，难以及时形成稳定的横截面判断。",
        "所以本文没有另外搭一套脱离项目背景的新系统，而是把现有量化研究平台里已经在运行的行业热度子系统拿出来，按真实代码、真实接口和真实历史快照重新梳理一遍，把数据接入、行业评分、个股筛选和可视化展示这几条主链路讲清楚。":
            "因此，本文没有另行构建脱离项目背景的新系统，而是以现有量化研究平台中已经实现的行业热度子系统为研究对象，基于真实代码、真实接口和真实历史快照，系统梳理数据接入、行业评分、个股筛选和可视化展示等主链路。",
        "没有追求大而全": "没有追求覆盖所有量化研究功能",
        "拿行业目录和摘要": "获取行业目录和摘要",
        "从同花顺接口拿": "从同花顺接口获取",
        "再按字段缺口补接": "再根据字段缺口补充接入",
        "串起来": "连接起来",
        "拆成": "组织为",
        "这样的写法虽然不炫技": "这种处理方式没有追求模型复杂度",
        "看的主要是": "主要依据",
        "更像市场有没有继续确认这一判断": "更侧重反映市场对该判断的短期确认程度",
        "摆出来": "呈现出来",
        "最突出的麻烦": "最突出的问题",
        "不同数据一起进来时口径很不整齐": "多类数据进入同一链路时口径差异较大",
        "真要把这些数据送进": "若要将这些数据纳入",
        "只盯": "仅关注",
        "只看": "仅关注",
        "数据拿到之后，马上要处理的不是模型": "数据获取之后，优先处理的不是模型选择",
        "拿不到": "无法获得",
        "拿到": "获得",
        "跑起来": "完成输出",
        "跑出来": "生成",
        "页面空掉": "页面出现空白或不可用状态",
        "空掉": "空白或不可用",
        "兜底逻辑": "回退逻辑",
        "多加一道兜底": "增加回退机制",
        "补源": "调用补充数据源",
        "补齐": "补充",
        "更像": "更接近",
        "前台": "前端视野",
        "几张静态图并排展示": "若干静态图表的简单并列",
        "继续打磨": "继续优化",
        "性能小技巧": "性能优化辅助手段",
        "更接近是一个行业数据适配中枢": "而是承担行业数据适配中枢的作用",
        "适配层真正麻烦的地方": "适配层的核心难点",
        "项目里为此维护了": "项目为此维护了",
        "再配合过期缓存把结果尽量补充": "再配合过期缓存尽量补充结果",
        "不是增加回退机制本身": "不是单纯增加回退机制",
        "让行业目录和成分股匹配尽量不断链": "尽量保持行业目录和成分股匹配链路的连续性",
        "IndustryAnalyzer 更接近是一个把原始行业字段整理成研究结果的中间层":
            "IndustryAnalyzer 的作用是把原始行业字段整理为研究结果",
        "把热力图与排行榜所需的数据生成": "生成热力图与排行榜所需的数据",
        "逐行业拉全量成分股": "逐行业获取全量成分股",
        "系统没有把缓存理解为性能优化辅助手段，而是把它当作行业研究原型稳定运行的一部分。":
            "系统不仅将缓存视为性能优化手段，也将其作为行业研究原型稳定运行的一部分。",
        "短期被市场推到前端视野的公司": "短期市场关注度快速上升的公司",
        "谁正在被市场短期推到前端视野": "谁在短期内获得更高市场关注",
        "这个子系统": "该系统",
        "真实代码、真实接口": "实际代码、实际接口",
        "主链路": "核心链路",
        "快路径": "快速路径",
        "给用户一个透明的排序依据": "为使用者提供可解释的排序依据",
        "用户至少要知道": "使用者应能够辨析",
        "用户在详情页中会关注一只股票为什么排在前面": "使用者在详情页中会关注某只股票排序靠前的原因",
        "便于老师检查": "便于评阅教师复核",
        "老师和读者": "评阅者和读者",
        "用户看到的是": "使用者看到的是",
        "用户可以": "使用者可以",
        "用户能": "使用者能够",
        "用户从": "使用者从",
        "用户每次": "使用者每次",
        "用户当前": "使用者当前",
        "用户才能": "使用者才能",
        "用户先": "使用者先",
        "用户打开": "使用者打开",
        "用户切换": "使用者切换",
        "用户一般会": "使用者通常会",
        "给用户": "为使用者",
        "帮助用户": "帮助使用者",
        "让用户": "使使用者",
        "等用户": "当使用者",
        "某个行业": "某一行业",
        "只要这条路径稳定": "只要该路径稳定",
        "本文没有把研究内容组织为几块互不相干的模块，而是尽量按系统真实使用的顺序来梳理。":
            "本文没有将研究内容组织为若干相互割裂的模块，而是按照系统实际使用顺序进行梳理。",
        "更接近日常研究场景的路径，是先打开行业页，看热力图和排行榜怎样把横截面结果呈现出来；接着沿着某一行业继续查看趋势、成分股和龙头股详情；最后再回头解释这些结果是怎样由适配器、缓存和评分器一步步算出来的。":
            "较符合常规行业研究场景的路径是：首先进入行业页面，观察热力图和排行榜呈现的横截面结果；随后围绕某一行业继续查看趋势、成分股和龙头股详情；最后再解释这些结果由适配器、缓存和评分器逐步生成的过程。",
        "逐项记录按钮截图": "逐项检查界面元素状态",
        "老师或后续读者": "评阅教师或后续读者",
        "老师和后续读者": "评阅教师和后续读者",
        "为使用者一个可理解的反馈": "为使用者提供可理解的反馈",
        "页面既能支持答辩展示，也能支持日常复盘": "系统既能支持答辩演示中的样例复核，也能支持后续复盘分析",
        "对于用户来说": "对于使用者而言",
        "更稳的公司": "更加稳定的公司",
        "更稳定": "更加稳定",
        "很难": "难以",
    }

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        original = paragraph.text.strip()
        if not original:
            continue
        polished = exact_replacements.get(original, original)
        for old, new in phrase_replacements.items():
            polished = polished.replace(old, new)
        if polished != original:
            replace_paragraph_text(paragraph, polished)


def normalize_standalone_project_scope(doc: Document) -> None:
    """Present the industry-heat work as the complete thesis project."""
    exact_replacements = {
        "A 股市场行业轮动速度较快，仅依赖人工盯盘、资讯阅读和手工比对，难以及时形成稳定的横截面判断。毕业设计任务书希望解决的正是这一类问题，因此，本文没有另行构建脱离项目背景的新系统，而是以现有量化研究平台中已经实现的行业热度子系统为研究对象，基于实际代码、实际接口和真实历史快照，系统梳理数据接入、行业评分、个股筛选和可视化展示等核心链路。":
            "A 股市场行业轮动速度较快，仅依赖人工盯盘、资讯阅读和手工比对，难以及时形成稳定的横截面判断。围绕这一问题，本文设计并实现了一个面向 A 股行业研究的行业热度识别与龙头股遴选系统，基于实际代码、实际接口和真实历史快照，系统梳理数据接入、行业评分、个股筛选和可视化展示等核心链路。",
        "从工程实现看，该子系统没有追求覆盖所有量化研究功能，而是围绕日常行业研究最常用的界面来组织。前端主要保留热力图、排行榜、趋势面板和龙头股详情这些入口；后端用 FastAPI 把热门行业、成分股、热力图、趋势和龙头股相关接口连接起来；评分和整理逻辑则放在 Python 分析层中完成。数据侧也不是单一来源，而是先由同花顺优先（Tonghuashun-first，以下简称 THS-first）适配器获取行业目录和摘要，再根据字段缺口补充接入 AKShare、新浪财经和腾讯接口。":
            "从工程实现看，本系统围绕行业研究中的核心流程来组织。前端提供热力图、排行榜、趋势面板和龙头股详情等界面；后端用 FastAPI 连接热门行业、成分股、热力图、趋势和龙头股相关接口；评分和整理逻辑则放在 Python 分析层中完成。数据侧也不是单一来源，而是先由同花顺优先（Tonghuashun-first，以下简称 THS-first）适配器获取行业目录和摘要，再根据字段缺口补充接入 AKShare、新浪财经和腾讯接口。",
        "模型部分没有额外引入难以解释的黑箱方法，而是直接沿用项目里已经落地的两个分析类。IndustryAnalyzer 负责判断行业热度，主要依据价格变化、资金承接、交易活跃度和波动率代理；LeaderStockScorer 负责在行业内部比较个股，把规模、估值、盈利、成长、动量和活跃度六个维度合成综合分，并保留快照场景下的快速评分链路。这种处理方式没有追求模型复杂度，但更容易和现有工程实现一一对应。":
            "模型部分没有额外引入难以解释的黑箱方法，而是围绕本系统中的两个核心分析类展开。IndustryAnalyzer 负责判断行业热度，主要依据价格变化、资金承接、交易活跃度和波动率代理；LeaderStockScorer 负责在行业内部比较个股，把规模、估值、盈利、成长、动量和活跃度六个维度合成综合分，并保留快照场景下的快速评分链路。这种处理方式没有追求模型复杂度，但更容易与系统实现一一对应。",
        "关键词：金融大数据；热门行业识别；龙头股遴选；多源数据；量化研究平台":
            "关键词：金融大数据；热门行业识别；龙头股遴选；多源数据；行业热度分析系统",
        "Against the backdrop of rapidly expanding financial market datasets, approaches that rely primarily on manual experience are no longer sufficient for identifying hot industries and representative leader stocks in a timely and interpretable manner. Focusing on the graduation-project topic of hot-industry identification and leader-stock selection, this thesis takes the industry-heat subsystem of an existing quantitative research platform as its research object and develops an A-share-oriented prototype for financial big-data analysis that integrates data acquisition, industry scoring, stock screening, and visual presentation.":
            "Against the backdrop of rapidly expanding financial market datasets, approaches that rely primarily on manual experience are no longer sufficient for identifying hot industries and representative leader stocks in a timely and interpretable manner. Focusing on the graduation-project topic of hot-industry identification and leader-stock selection, this thesis designs and implements an A-share-oriented industry-heat analysis system that integrates multi-source data acquisition, industry scoring, leader-stock screening, and visual presentation.",
        "Keywords: financial big data; hot-industry identification; leading stock selection; multi-source data; quantitative research platform":
            "Keywords: financial big data; hot-industry identification; leading stock selection; multi-source data; industry-heat analysis system",
        "本文依托的项目并不是单一的课程设计程序，而是一个较为完整的量化研究平台。当前公开仓库对外保留今日研究、策略回测、实时行情和行业热度四个公开入口，其中今日研究用于汇总当天研究档案，策略回测和实时行情用于验证与监控，行业热度页面则与毕业设计任务书的目标最为一致。因此，本文不再泛化讨论平台全部模块，而是从现有工程实现中抽取行业热度子系统作为核心研究对象，在实际代码、实际接口与真实历史快照的基础上完成毕业论文撰写。":
            "本文围绕热门行业识别与龙头股遴选任务，设计并实现了一个面向 A 股行业研究的行业热度分析系统。该系统以多源行情、资金流、估值和财务字段为基础，形成从数据采集、字段清洗、行业评分、龙头股遴选到可视化展示的完整流程。论文后续章节均以该系统作为完整毕业设计项目进行论述，并在实际代码、实际接口与真实历史快照的基础上完成分析。",
        "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。":
            "本文围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了行业热度分析系统的需求分析、模型设计、工程实现和结果总结。通过梳理系统中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
    }
    phrase_replacements = {
        "现有量化研究平台中已经实现的行业热度子系统": "本文设计并实现的行业热度分析系统",
        "现有量化研究平台中的行业热度子系统": "本文设计并实现的行业热度分析系统",
        "现有量化研究平台里已经在运行的行业热度子系统": "本文设计并实现的行业热度分析系统",
        "量化研究平台": "行业热度分析系统",
        "行业热度子系统": "行业热度分析系统",
        "行业子系统": "行业热度分析系统",
        "该子系统": "该系统",
        "子系统": "系统",
        "从现有工程实现中抽取": "围绕",
        "现有工程实现": "系统实现",
        "当前公开仓库对外保留今日研究、策略回测、实时行情和行业热度四个公开入口，": "",
        "其中今日研究用于汇总当天研究档案，策略回测和实时行情用于验证与监控，": "",
        "因此，本文不再泛化讨论平台全部模块，而是": "因此，本文",
        "全文按“提出问题、解释方法、落到系统、再看结果”的顺序展开。第一章说明选题背景、研究现状以及本文聚焦的行业热度分析系统；":
            "全文按“提出问题、解释方法、落到系统、再看结果”的顺序展开。第一章说明选题背景、研究现状以及本文所设计的行业热度分析系统；",
        "若要把行业信号进一步接入回测或研究报告生成流程，可以把固定快照转成更规范的数据表，再与今日研究或研究档案的任务记录关联。":
            "若要将行业信号进一步扩展为回测验证或研究报告生成流程，可以把固定快照转成更规范的数据表，再与后续研究记录关联。",
        "也能为后续扩展到今日研究、报告生成或回测模块留下入口":
            "也能为后续扩展到报告生成或回测验证留下入口",
        "明确行业热度分析系统的数据来源": "明确本系统的数据来源",
        "表 3.2 行业热度分析系统核心模块职责与证据来源": "表 3.2 行业热度分析系统核心模块职责与证据来源",
        "结合系统总体架构，行业热度分析系统各层职责和证据来源": "结合系统总体架构，各层职责和证据来源",
        "本项目里": "本系统中",
        "项目里": "系统中",
        "项目中的": "系统中的",
        "项目当前保留的": "系统保存的",
        "项目保留": "系统保存",
    }

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        original = paragraph.text.strip()
        if not original:
            continue
        normalized = exact_replacements.get(original, original)
        for old, new in phrase_replacements.items():
            normalized = normalized.replace(old, new)
        if normalized != original:
            replace_paragraph_text(paragraph, normalized)


def polish_final_submission_copy(doc: Document) -> None:
    """Tighten final wording so the thesis reads like a formal submission."""
    exact_replacements = {
        "把文献里的龙头股识别方法和当前项目放在一起看，有一个很直接的共识：真正能长期代表行业状态的股票，通常不会只靠某一天的涨幅来判断。规模、盈利、成长和估值提供的是相对稳定的基本面线索，成交额、换手率和近期动量则更侧重反映市场对该判断的短期确认程度。[4-5][8-12]":
            "综合相关文献与本系统实现可以发现，能够长期代表行业状态的股票，通常不能只依据单日涨幅判断。规模、盈利、成长和估值提供的是相对稳定的基本面线索，成交额、换手率和近期动量则更侧重反映市场对该判断的短期确认程度。[4-5][8-12]",
        "总体来看，现有研究已经为热门行业识别和龙头股筛选提供了较充分的理论基础，但一落到工程实现，就还会遇到两个明显空档：一是多源适配、字段回退、缓存和页面展示这些真正影响可用性的细节经常被简化；二是行业识别和龙头股遴选常被拆开讨论，缺少统一的研究闭环。本文的重点不是再堆一个复杂模型，而是在真实项目基础上，把多源数据获取、行业评分、龙头股筛选和前端展示真正接起来。":
            "总体来看，现有研究已经为热门行业识别和龙头股筛选提供了较充分的理论基础，但落实到工程实现层面，仍存在两个明显不足：一是多源适配、字段回退、缓存和页面展示等直接影响可用性的细节经常被简化；二是行业识别和龙头股遴选常被拆开讨论，缺少统一的研究闭环。本文的重点不是继续堆叠复杂模型，而是在系统实现基础上，将多源数据获取、行业评分、龙头股筛选和前端展示有效连接起来。",
        "在这个课题里，“龙头股”并不等于某一天涨得最快的股票。更常见的情况是，短期涨幅靠前的个股未必真的能代表行业，有时只是情绪推动；真正更能代表行业状态的，往往是那些规模、盈利、成长和市场关注度都更加稳定的公司。传统龙头企业识别也大多沿着这个思路，从市值规模、产业地位、盈利质量、成长能力、估值合理性和市场表现等多个维度综合判断。[4-5][8-12]":
            "在本课题中，“龙头股”并不等同于单日涨幅最高的股票。短期涨幅靠前的个股未必能够代表行业状态，有时更多反映市场情绪推动；更能体现行业代表性的，通常是规模、盈利、成长和市场关注度较稳定的公司。传统龙头企业识别也大多沿着这一思路，从市值规模、产业地位、盈利质量、成长能力、估值合理性和市场表现等多个维度综合判断。[4-5][8-12]",
        "这样分层之后，每一层要解决的问题会清楚很多。表现层只需要关心图表交互和结果展示，服务层负责把分析结果组织成统一响应，分析层专注评分计算与结果解释，数据层则通过名称映射、节点回退、符号缓存和过期缓存保障等机制提高数据可用性。后面如果要替换局部数据源，或者微调评分权重，改动也大多集中在分析层和数据层，不必把前端页面整体重写一遍。":
            "完成分层后，各层职责更加清晰。表现层主要负责图表交互和结果展示，服务层负责将分析结果组织为统一响应，分析层专注于评分计算与结果解释，数据层则通过名称映射、节点回退、符号缓存和过期缓存保障等机制提高数据可用性。后续若需要替换局部数据源或微调评分权重，改动也大多集中在分析层和数据层，无需整体重写前端页面。",
        "和那种只依赖单一接口的数据抓取脚本相比，这套架构更适合支撑准实时行业研究。一方面，多源适配提高了行业目录、资金流、成分股、估值和历史走势等字段的覆盖率；另一方面，缓存和分层封装让页面响应速度与结果口径更容易保持一致。对论文写作来说，这一点也很重要，因为只有系统本身运行得足够稳定，后面获得的快照样本和案例分析才有说服力。":
            "与仅依赖单一接口的数据抓取脚本相比，这套架构更适合支撑准实时行业研究。一方面，多源适配提高了行业目录、资金流、成分股、估值和历史走势等字段的覆盖率；另一方面，缓存和分层封装使页面响应速度与结果口径更容易保持一致。对论文写作而言，这一点具有重要意义，因为只有系统本身运行足够稳定，后续获得的快照样本和案例分析才具有说服力。",
        "从毕业设计展示角度看，四层架构还有一个好处：它能把论文中的模型说明和真实页面运行过程对应起来。答辩时如果从前端热力图开始演示，页面请求会先进入服务层，再调用分析层和数据层；如果从代码角度说明，则可以反过来从数据适配器、行业分析器、评分器一路讲到前端展示。两种讲法对应的是同一条链路，能够减少论文和项目脱节的问题。":
            "从毕业设计展示角度看，四层架构还有一项优势：它能够将论文中的模型说明与页面运行过程对应起来。答辩时若从前端热力图开始演示，页面请求会先进入服务层，再调用分析层和数据层；若从代码角度说明，则可以反向从数据适配器、行业分析器和评分器说明至前端展示。两种说明路径对应同一条链路，能够减少论文表述与系统实现脱节的问题。",
        "表现层并不是简单负责“好看”。行业热度页面需要同时承载总览、排序、筛选、下钻和回放等动作，因此前端组件必须把热力图、排行榜、详情弹窗和偏好配置组织在同一个状态流里。如果前端只展示静态表格，用户就难以从行业热度继续推进到龙头股遴选。":
            "表现层并不仅承担视觉呈现功能。行业热度页面需要同时承载总览、排序、筛选、下钻和回放等动作，因此前端组件必须把热力图、排行榜、详情弹窗和偏好配置组织在同一个状态流中。如果前端只展示静态表格，使用者就难以从行业热度继续推进到龙头股遴选。",
        "文件化存储还有一个好处，是非常适合毕业设计阶段的复核。老师或答辩委员如果追问第六章表格里的数据来源，论文可以直接说明其来自固定快照文件，而不是当天重新计算出的临时结果。对一个本地运行的研究原型来说，这种透明度比复杂数据库更重要。":
            "文件化存储还有一项优势，即较适合毕业设计阶段的复核。评阅教师或答辩委员如果追问第六章表格中的数据来源，论文可以直接说明其来自固定快照文件，而不是当天重新计算出的临时结果。对于本地运行的研究原型而言，这种透明度比复杂数据库更具现实意义。",
        "这种拆分可以降低单次请求压力。若把所有字段都塞进一个接口，首屏加载会被详情数据拖慢；若每个小字段都组织为单独接口，前端又会出现过多请求。当前设计介于两者之间，把常用汇总结果和详情结果分开，既能保证首屏速度，也能支持用户继续下钻。":
            "这种拆分可以降低单次请求压力。若将所有字段集中在单一接口中，首屏加载会被详情数据拖慢；若每个细粒度字段都组织为单独接口，前端又会出现过多请求。当前设计介于两者之间，将常用汇总结果和详情结果分开，既能保证首屏速度，也能支持使用者继续下钻。",
    }
    phrase_replacements = {
        "真正的完整投资验证": "更完整的投资有效性验证",
        "很容易把短期情绪波动当成热点": "容易把短期情绪波动当成热点",
        "行业研究真正需要的是": "行业研究更需要的是",
        "很强的时序性": "较强的时序性",
        "用户通常不会只接受一个最终分数。研究者会继续追问": "研究者通常不会只接受一个最终分数，而会继续追问",
        "真正需要的全部字段": "需要的全部字段",
        "用户操作": "使用者操作",
        "用户选中": "使用者选中",
        "才真正能转化": "才能转化",
        "使用户不必": "使使用者不必",
        "方便用户判断": "便于使用者判断",
        "真正完成行业综合评分": "承担行业综合评分",
        "很容易和项目当前状态脱节": "容易与系统当前状态脱节",
        "项目当前状态": "系统当前状态",
        "真实项目实现": "系统实现",
        "本项目当前": "本系统当前",
        "本项目中": "本系统中",
        "用户状态": "使用者状态",
        "用户习惯": "研究习惯",
        "尺度差异很大": "尺度差异较大",
        "用户可能看到": "使用者可能看到",
        "短期涨幅很高但波动也很剧烈": "短期涨幅较高但波动也较剧烈",
        "短期涨幅很高但估值": "短期涨幅较高但估值",
        "减少用户把短期涨幅等同于长期龙头的误解": "减少使用者将短期涨幅等同于长期龙头的误解",
        "提醒用户区分": "提醒使用者区分",
        "用户实际浏览页面的路径": "使用者实际浏览页面的路径",
        "真正影响使用体验的还有": "直接影响使用体验的还有",
        "帮助用户理解": "帮助使用者理解",
        "用户关心的行业": "使用者关注的行业",
        "当前项目": "本系统",
        "项目代码": "系统代码",
        "真实项目基础": "系统实现基础",
        "数据层则承担了系统可靠性的底座作用": "数据层承担系统可靠性的基础保障作用",
        "模型公式即使写得很清楚": "模型公式即使表述清晰",
        "还有一个好处": "还有一项优势",
        "这样做的好处是": "这样做的优势在于",
    }

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        original = paragraph.text.strip()
        if not original:
            continue
        polished = exact_replacements.get(original, original)
        for old, new in phrase_replacements.items():
            polished = polished.replace(old, new)
        if polished != original:
            replace_paragraph_text(paragraph, polished)


def remove_repeated_long_body_paragraphs(doc: Document) -> None:
    """Remove accidental repeated body paragraphs introduced by iterative regeneration."""
    seen: set[str] = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if len(text) < 60:
            continue
        if re.match(r"^\[\d+\]", text):
            continue
        if re.match(r"^[表图]\s*\d+\.\d+", text):
            continue
        if text in seen:
            delete_paragraph(paragraph)
            continue
        seen.add(text)


def replace_formula_tables(doc: Document) -> None:
    replace_formula_block(
        doc,
        "Sindustry",
        "Sindustry=0.35×Zm+0.35×Zf+0.15×Zv-0.15×Zr",
        "（4.2.1）",
    )
    replace_formula_block(
        doc,
        "Sleader",
        "Sleader=100×(0.20×s1+0.15×s2+0.25×s3+0.20×s4+0.10×s5+0.10×s6)",
        "（4.4.1）",
    )


def normalize_reference_entries(doc: Document) -> None:
    entries = [
        "[1] 孟小峰, 慈祥. 大数据管理：概念、技术与挑战[J]. 计算机研究与发展, 2013, 50(1): 146-169.",
        "[2] 张寒冰, 李智鑫, 荆一楠, 等. 数字化转型背景下金融风险监测与预警体系研究[J]. 中国工程科学, 2024, 26(3). DOI: 10.15302/J-SSCAE-2024.03.020.",
        "[3] 武文超. 中国A股市场的行业轮动现象分析：基于动量和反转交易策略的检验[J]. 金融理论与实践, 2014(9): 111-114.",
        "[4] Koller T, Goedhart M, Wessels D. Valuation: Measuring and Managing the Value of Companies[M]. 7th ed. Hoboken: Wiley, 2020.",
        "[5] Damodaran A. Investment Valuation: Tools and Techniques for Determining the Value of Any Asset[M]. 3rd ed. Hoboken: Wiley, 2012.",
        [
            "[6] McKinney W. Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython[M].",
            "3rd ed. Sebastopol: O'Reilly Media, 2022.",
        ],
        [
            "[7] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python[J].",
            "Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        ],
        "[8] Fama E F. Efficient Capital Markets: A Review of Theory and Empirical Work[J]. The Journal of Finance, 1970, 25(2): 383-417.",
        "[9] Markowitz H. Portfolio Selection[J]. The Journal of Finance, 1952, 7(1): 77-91.",
        "[10] Sharpe W F. Capital Asset Prices: A Theory of Market Equilibrium under Conditions of Risk[J]. The Journal of Finance, 1964, 19(3): 425-442.",
        "[11] Fama E F, French K R. Common Risk Factors in the Returns on Stocks and Bonds[J]. Journal of Financial Economics, 1993, 33(1): 3-56.",
        "[12] Fama E F, French K R. A Five-Factor Asset Pricing Model[J]. Journal of Financial Economics, 2015, 116(1): 1-22.",
        "[13] Jegadeesh N, Titman S. Returns to Buying Winners and Selling Losers: Implications for Stock Market Efficiency[J]. The Journal of Finance, 1993, 48(1): 65-91.",
        "[14] Chan L K C, Jegadeesh N, Lakonishok J. Momentum Strategies[J]. The Journal of Finance, 1996, 51(5): 1681-1713.",
        "[15] Moskowitz T J, Grinblatt M. Do Industries Explain Momentum?[J]. The Journal of Finance, 1999, 54(4): 1249-1290.",
        "[16] Hou K. Industry Information Diffusion and the Lead-Lag Effect in Stock Returns[J]. The Review of Financial Studies, 2007, 20(4): 1113-1138.",
        "[17] Hong H, Stein J C. A Unified Theory of Underreaction, Momentum Trading, and Overreaction in Asset Markets[J]. The Journal of Finance, 1999, 54(6): 2143-2184.",
        "[18] Barberis N, Shleifer A, Vishny R. A Model of Investor Sentiment[J]. Journal of Financial Economics, 1998, 49(3): 307-343.",
        "[19] Lo A W, MacKinlay A C. Data-Snooping Biases in Tests of Financial Asset Pricing Models[J]. The Review of Financial Studies, 1990, 3(3): 431-467.",
        "[20] Campbell J Y, Lo A W, MacKinlay A C. The Econometrics of Financial Markets[M]. Princeton: Princeton University Press, 1997.",
        "[21] Gu S, Kelly B, Xiu D. Empirical Asset Pricing via Machine Learning[J]. The Review of Financial Studies, 2020, 33(5): 2223-2273.",
        "[22] Kelly B, Pruitt S, Su Y. Characteristics Are Covariances: A Unified Model of Risk and Return[J]. Journal of Financial Economics, 2019, 134(3): 501-524.",
        "[23] Henrique B M, Sobreiro V A, Kimura H. Literature Review: Machine Learning Techniques Applied to Financial Market Prediction[J]. Expert Systems with Applications, 2019, 124: 226-251.",
        [
            "[24] Cavalcante R C, Brasileiro R C, Souza V L, et al. Computational Intelligence and Financial Markets: A Survey and Future Directions[J].",
            "Expert Systems with Applications, 2016, 55: 194-211.",
        ],
        [
            "[25] Hasan M M, Popp J, Olah J. Current landscape and influence of big data on finance[J].",
            "Journal of Big Data, 2020, 7: 21.",
        ],
        "[26] Chen M, Mao S, Liu Y. Big Data: A Survey[J]. Mobile Networks and Applications, 2014, 19(2): 171-209.",
        "[27] Chen H, Chiang R H L, Storey V C. Business Intelligence and Analytics: From Big Data to Big Impact[J]. MIS Quarterly, 2012, 36(4): 1165-1188.",
        "[28] Jagadish H V, Gehrke J, Labrinidis A, et al. Big Data and Its Technical Challenges[J]. Communications of the ACM, 2014, 57(7): 86-94.",
        "[29] Kleppmann M. Designing Data-Intensive Applications[M]. Sebastopol: O'Reilly Media, 2017.",
        "[30] MacQueen J. Some Methods for Classification and Analysis of Multivariate Observations[C]//Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability. Berkeley: University of California Press, 1967: 281-297.",
        "[31] Arthur D, Vassilvitskii S. k-means++: The Advantages of Careful Seeding[C]//Proceedings of the Eighteenth Annual ACM-SIAM Symposium on Discrete Algorithms. Philadelphia: SIAM, 2007: 1027-1035.",
        "[32] Rousseeuw P J. Silhouettes: A Graphical Aid to the Interpretation and Validation of Cluster Analysis[J]. Journal of Computational and Applied Mathematics, 1987, 20: 53-65.",
        "[33] Hastie T, Tibshirani R, Friedman J. The Elements of Statistical Learning: Data Mining, Inference, and Prediction[M]. 2nd ed. New York: Springer, 2009.",
        "[34] McKinney W. Data Structures for Statistical Computing in Python[C]//Proceedings of the 9th Python in Science Conference. Austin: SciPy, 2010: 56-61.",
        "[35] Harris C R, Millman K J, van der Walt S J, et al. Array Programming with NumPy[J]. Nature, 2020, 585(7825): 357-362.",
        "[36] Virtanen P, Gommers R, Oliphant T E, et al. SciPy 1.0: Fundamental Algorithms for Scientific Computing in Python[J]. Nature Methods, 2020, 17(3): 261-272.",
        "[37] Hunter J D. Matplotlib: A 2D Graphics Environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.",
        "[38] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. Irvine: University of California, Irvine, 2000.",
        "[39] FastAPI. FastAPI Documentation[EB/OL]. [2026-04-24]. https://fastapi.tiangolo.com/.",
        "[40] Meta Open Source. React Documentation[EB/OL]. [2026-04-24]. https://react.dev/.",
        "[41] AKShare. AKShare Documentation[EB/OL]. [2026-04-24]. https://akshare.akfamily.xyz/.",
        "[42] Loughran T, McDonald B. When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks[J]. The Journal of Finance, 2011, 66(1): 35-65.",
    ]
    ref_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "参考文献"),
        None,
    )
    if ref_index is None:
        raise RuntimeError("Reference section anchor not found.")
    ack_index = next(
        (
            idx
            for idx in range(ref_index + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip() == "致 谢"
        ),
        None,
    )
    if ack_index is None:
        raise RuntimeError("Acknowledgement section anchor not found.")

    ref_heading = doc.paragraphs[ref_index]
    ack_element = doc.paragraphs[ack_index]._element
    current_element = ref_heading._element.getnext()
    while current_element is not None and current_element is not ack_element:
        next_element = current_element.getnext()
        current_element.getparent().remove(current_element)
        current_element = next_element

    current = ref_heading
    for entry in entries:
        paragraph = insert_paragraph_after(current)
        clear_paragraph(paragraph)
        if isinstance(entry, str):
            paragraph.add_run(entry)
            current = paragraph
            continue
        for index, line in enumerate(entry):
            run = paragraph.add_run(line)
            if index < len(entry) - 1:
                run.add_break(WD_BREAK.LINE)
        current = paragraph


def apply_minor_layout_overrides(doc: Document) -> None:
    target_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "5.6 偏好配置与持续跟踪实现"),
        None,
    )
    if target_heading is not None:
        target_heading.paragraph_format.page_break_before = True
        target_heading.paragraph_format.keep_with_next = True
        next_element = target_heading._element.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            first_para = Paragraph(next_element, target_heading._parent)
            first_para.paragraph_format.keep_with_next = True
            second_element = next_element.getnext()
            if second_element is not None and second_element.tag == qn("w:p"):
                second_para = Paragraph(second_element, target_heading._parent)
                second_para.paragraph_format.keep_together = True

    completion_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "6.5 对毕业设计目标的达成情况"),
        None,
    )
    if completion_heading is not None:
        completion_heading.paragraph_format.page_break_before = False
        completion_heading.paragraph_format.keep_with_next = True
        next_element = completion_heading._element.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            intro_para = Paragraph(next_element, completion_heading._parent)
            intro_para.paragraph_format.keep_with_next = True


def apply_superscript_citations(paragraph) -> None:
    text = paragraph.text
    if not text or "[" not in text or paragraph.text.strip().startswith("["):
        return
    matches = list(re.finditer(r"\[[0-9,\-，]+\]", text))
    if not matches:
        return
    clear_paragraph(paragraph)
    cursor = 0
    for match in matches:
        if match.start() > cursor:
            normal_run = paragraph.add_run(text[cursor:match.start()])
            format_body_run(normal_run)
        citation_run = paragraph.add_run(match.group(0))
        format_body_run(citation_run)
        citation_run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        tail_run = paragraph.add_run(text[cursor:])
        format_body_run(tail_run)


def format_paragraphs(doc: Document) -> None:
    abstract_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "摘  要"),
        0,
    )
    english_abstract_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "ABSTRACT"),
        abstract_index,
    )
    toc_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "目  录"),
        english_abstract_index,
    )
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < abstract_index:
            continue
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""

        if not text:
            continue
        if text == "摘  要":
            clear_paragraph(paragraph)
            run = paragraph.add_run("摘  要")
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if text == "ABSTRACT":
            clear_paragraph(paragraph)
            run = paragraph.add_run("ABSTRACT")
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            set_run_text_font(run, "Times New Roman", Pt(18), bold=True)
            continue
        if text == "目  录":
            format_toc_title(paragraph)
            continue
        if style_name in {"toc 1", "toc 2"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(12), bold=False)
            continue
        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if style_name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(14), bold=True)
            continue
        if style_name == "Heading 3":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=True)
            continue
        if text in {"结 论", "参考文献", "致 谢"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if text.startswith("关键词："):
            content = text.split("：", 1)[1]
            clear_paragraph(paragraph)
            label_run = paragraph.add_run("关键词：")
            value_run = paragraph.add_run(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            set_run_text_font(label_run, "宋体", Pt(12), bold=False)
            set_run_text_font(value_run, "宋体", Pt(12), bold=False)
            continue
        if text.startswith("Keywords:"):
            content = text[len("Keywords:"):].strip()
            clear_paragraph(paragraph)
            label_run = paragraph.add_run("Keywords:")
            spacer_run = paragraph.add_run(" ")
            value_run = paragraph.add_run(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            set_run_text_font(label_run, "Times New Roman", Pt(12), bold=True)
            set_run_text_font(spacer_run, "Times New Roman", Pt(12), bold=False)
            set_run_text_font(value_run, "Times New Roman", Pt(12), bold=False)
            continue
        if re.match(r"^图\s*\d+\.\d+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=False)
            continue
        if text.startswith("图片来源："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(18)
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(10.5), bold=False)
            continue
        if re.match(r"^表\s*\d+\.\d+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=False)
            continue
        if text.startswith("Sindustry") or text.startswith("Sleader"):
            equation_text = "Sindustry=0.35×Zm+0.35×Zf+0.15×Zv-0.15×Zr"
            if text.startswith("Sleader"):
                equation_text = "Sleader=100×(0.20×s1+0.15×s2+0.25×s3+0.20×s4+0.10×s5+0.10×s6)"
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            run = paragraph.add_run(equation_text)
            set_run_text_font(run, "Times New Roman", Pt(10.5), bold=False)
            continue
        if text in {"（4.2.1）", "（4.4.1）"}:
            clear_paragraph(paragraph)
            run = paragraph.add_run(text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            set_run_text_font(run, "Times New Roman", Pt(10.5), bold=False)
            continue
        if re.match(r"^\[\d+\]", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(12), bold=False)
            continue
        if text.startswith("题   目：") or text.startswith("学    院：") or text.startswith("专    业："):
            continue

        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if abstract_index < idx < english_abstract_index or english_abstract_index < idx < toc_index:
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
        else:
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        for run in paragraph.runs:
            if english_abstract_index < idx < toc_index:
                set_run_text_font(run, "Times New Roman", Pt(12), bold=False)
            else:
                format_body_run(run)
        apply_superscript_citations(paragraph)


def configure_headers_and_footers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        if index < 2:
            configure_section_header(section, None)
            configure_section_footer(section, False)
            continue
        configure_section_header(section, HEADER_TEXT)
        configure_section_footer(section, True)

    # The composed official template keeps the front matter split as:
    # cover/declaration, Chinese abstract, English abstract + TOC, then body.
    # Keep all front-matter pages on Roman numerals and restart正文 at 1.
    if len(doc.sections) >= 5:
        set_section_page_number_format(doc.sections[2], fmt="upperRoman", start=1)
        set_section_page_number_format(doc.sections[3], fmt="upperRoman", start=2)
        set_section_page_number_format(doc.sections[4], fmt="decimal", start=1)
        return

    if len(doc.sections) >= 3:
        set_section_page_number_format(doc.sections[2], fmt="upperRoman", start=1)
    if len(doc.sections) >= 4:
        set_section_page_number_format(doc.sections[3], fmt="decimal", start=1)


def compute_toc_pages_from_pdf(pdf_path: Path) -> dict[str, str]:
    reader = PdfReader(str(pdf_path))
    raw_page_texts = [page.extract_text() or "" for page in reader.pages]
    page_texts = [normalize_search_text(text) for text in raw_page_texts]
    page_lines = [
        [normalize_search_text(line) for line in text.splitlines() if line.strip()]
        for text in raw_page_texts
    ]

    toc_start_index = next(
        (i for i, text in enumerate(page_texts) if text.count(".") > 100 and "1绪论" in text),
        None,
    )
    chapter_start_index = None
    if toc_start_index is not None:
        chapter_start_index = next(
            (
                i
                for i in range(toc_start_index + 1, len(page_texts))
                if page_texts[i].count(".") < 50 and "1绪论" in page_texts[i]
            ),
            None,
        )
    if chapter_start_index is None:
        chapter_start_index = next(
            (i for i, text in enumerate(page_texts) if text.count(".") < 50 and "1绪论" in text),
            None,
        )
    if chapter_start_index is None:
        chapter_start_index = next(
            (
                i
                for i, text in enumerate(page_texts)
                if text.count(".") < 50 and "绪论" in text and "研究背景与意义" in text
            ),
            None,
        )
    if chapter_start_index is None:
        raise RuntimeError("Failed to determine the first body page for TOC rebuilding.")

    toc_pages = {
        "摘  要": "I",
        "ABSTRACT": "II",
    }
    for _, title in TOC_BLUEPRINT:
        if title in toc_pages:
            continue
        normalized_title = normalize_search_text(title)
        title_without_chapter_number = re.sub(r"^\d+\s+", "", title)
        normalized_title_variants = {normalized_title}
        if title_without_chapter_number != title:
            normalized_title_variants.add(normalize_search_text(title_without_chapter_number))
        for page_index in range(chapter_start_index, len(page_texts)):
            if any(variant in page_lines[page_index] for variant in normalized_title_variants):
                toc_pages[title] = str(page_index - chapter_start_index + 1)
                break
    return toc_pages


def remove_page_break_paragraph_before(doc: Document, title: str) -> None:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != title:
            continue
        if idx == 0:
            return
        previous = paragraphs[idx - 1]
        if previous.text.strip() == "" and 'w:type="page"' in previous._element.xml:
            delete_paragraph(previous)
        paragraph.paragraph_format.page_break_before = True
        return


def normalize_major_breaks(doc: Document) -> None:
    heading_1_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Heading 1"]
    for index, paragraph in enumerate(heading_1_paragraphs):
        previous = paragraph._element.getprevious()
        if previous is not None and previous.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            previous_para = Paragraph(previous, paragraph._parent)
            if previous_para.text.strip() == "" and 'w:type="page"' in previous_para._element.xml:
                delete_paragraph(previous_para)
        if index == 0:
            paragraph.paragraph_format.page_break_before = False
            remove_page_break_before_flag(paragraph)
        else:
            paragraph.paragraph_format.page_break_before = True

    for title in ["结 论", "参考文献", "致 谢"]:
        remove_page_break_paragraph_before(doc, title)


def harmonize_section_breaks(target_doc: Document, source_doc: Document) -> None:
    source_chapter_break = deepcopy(find_heading_break_paragraph(source_doc, "绪论")._element.pPr.find(qn("w:sectPr")))
    target_chapter_break_paragraph = find_heading_break_paragraph(target_doc, "绪论")
    target_chapter_break_ppr = target_chapter_break_paragraph._element.get_or_add_pPr()
    existing_target_break = target_chapter_break_ppr.find(qn("w:sectPr"))
    if existing_target_break is not None:
        target_chapter_break_ppr.remove(existing_target_break)
    target_chapter_break_ppr.append(source_chapter_break)

    heading_index = None
    for index, paragraph in enumerate(target_doc.paragraphs):
        if paragraph.text.strip() == "绪论":
            heading_index = index
            break
    if heading_index is None:
        raise RuntimeError("Failed to locate chapter heading for section cleanup.")

    for index in range(max(0, heading_index - 4), heading_index):
        paragraph = target_doc.paragraphs[index]
        if paragraph._element is target_chapter_break_paragraph._element:
            continue
        p_pr = paragraph._element.pPr
        if p_pr is None:
            continue
        sect_pr = p_pr.find(qn("w:sectPr"))
        if sect_pr is not None:
            p_pr.remove(sect_pr)

    target_body = target_doc._element.body
    target_body_sect = target_body.find(qn("w:sectPr"))
    if target_body_sect is not None:
        target_body.remove(target_body_sect)
    target_body.append(deepcopy(source_doc._element.body.find(qn("w:sectPr"))))


def compose_official_template(source_doc: Document, source_layout_doc: Document) -> Document:
    TMP_DOC_DIR.mkdir(parents=True, exist_ok=True)

    front_doc = Document(str(TEMPLATE_PATH))
    fill_cover(front_doc)
    fill_declaration_table(front_doc)
    remove_elements_from_paragraph(front_doc, "摘  要")

    remove_elements_before_paragraph(source_doc, "摘  要")

    composer = Composer(front_doc)
    composer.append(source_doc)
    composer.save(str(COMPOSED_TMP_PATH))

    composed_doc = Document(str(COMPOSED_TMP_PATH))
    harmonize_section_breaks(composed_doc, source_layout_doc)
    set_core_properties(composed_doc)
    return composed_doc


def export_submission_artifacts(doc_path: Path) -> Path:
    template_page_1, template_page_2 = ensure_template_pdf_pages()
    front_cover_png = TMP_FRONT_ASSET_DIR / "front_cover.png"
    declaration_png = TMP_FRONT_ASSET_DIR / "front_declaration.png"
    front_pdf = TMP_FRONT_ASSET_DIR / "front_pages.pdf"

    cover_render = render_cover_page(template_page_1, front_cover_png)
    declaration_asset, declaration_render = prepare_declaration_page_asset(template_page_2, declaration_png)
    build_front_pdf(front_cover_png, declaration_asset, front_pdf)

    body_pdf = export_docx_pdf(doc_path, TMP_DOCX_RENDER_DIR)
    merge_submission_pdf(front_pdf, body_pdf, OUTPUT_PDF_PATH)
    if LEGACY_DUPLICATE_PDF_PATH.exists():
        LEGACY_DUPLICATE_PDF_PATH.unlink()
    render_pdf_preview(OUTPUT_PDF_PATH, TMP_SUBMISSION_RENDER_DIR)
    validate_front_page_alignment(
        get_rendered_preview_page(TMP_SUBMISSION_RENDER_DIR, 1),
        cover_render["placements"],
        cover_render["image_size"],
    )
    if declaration_render is not None:
        validate_declaration_page_alignment(
            get_rendered_preview_page(TMP_SUBMISSION_RENDER_DIR, 2),
            declaration_render["placements"],
            declaration_render["label_boxes"],
            declaration_render["image_size"],
        )
    return OUTPUT_PDF_PATH


def main() -> None:
    source_layout_doc = Document(str(DOC_PATH))
    normalize_chapter_headings(source_layout_doc)
    update_body_content(source_layout_doc)
    polish_academic_tone(source_layout_doc)
    normalize_standalone_project_scope(source_layout_doc)
    polish_final_submission_copy(source_layout_doc)
    remove_repeated_long_body_paragraphs(source_layout_doc)
    replace_formula_tables(source_layout_doc)
    relayout_figures(source_layout_doc)
    polish_final_submission_copy(source_layout_doc)
    normalize_major_breaks(source_layout_doc)

    composed_doc = compose_official_template(source_layout_doc, Document(str(DOC_PATH)))
    normalize_reference_entries(composed_doc)
    polish_final_submission_copy(composed_doc)
    format_data_tables(composed_doc)
    configure_headers_and_footers(composed_doc)
    format_paragraphs(composed_doc)
    apply_minor_layout_overrides(composed_doc)
    composed_doc.save(str(DOC_PATH))

    pdf_path = export_submission_artifacts(DOC_PATH)
    toc_pages = compute_toc_pages_from_pdf(pdf_path)
    rebuild_toc(composed_doc, toc_pages)
    polish_final_submission_copy(composed_doc)
    format_paragraphs(composed_doc)
    apply_minor_layout_overrides(composed_doc)
    composed_doc.save(str(DOC_PATH))

    pdf_path = export_submission_artifacts(DOC_PATH)
    verified_toc_pages = compute_toc_pages_from_pdf(pdf_path)
    if verified_toc_pages != toc_pages:
        rebuild_toc(composed_doc, verified_toc_pages)
        polish_final_submission_copy(composed_doc)
        format_paragraphs(composed_doc)
        apply_minor_layout_overrides(composed_doc)
        composed_doc.save(str(DOC_PATH))
        export_submission_artifacts(DOC_PATH)


if __name__ == "__main__":
    main()
